"""FastAPI-сервер: загрузка .docx, правка запросом (полный цикл edit.run_edit), скачивание."""

import asyncio
import json
import time
import uuid
from pathlib import Path

from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse

from docx_editor import edit as edit_mod
from docx_editor import log
from docx_editor.llm import _env
from docx_editor.parse import doc_map, index

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "out"
OUT.mkdir(exist_ok=True)

app = FastAPI()


def _session_path(session: str) -> Path:
    path = OUT / f"{session}.docx"
    if not path.exists():
        raise HTTPException(404, f"сессия {session} не найдена")
    return path


@app.get("/")
def root():
    return HTMLResponse((ROOT / "web" / "index.html").read_text())


@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    """Заводит новую сессию из .docx, отдаёт карту документа для немедленного рендера."""
    session = uuid.uuid4().hex
    path = OUT / f"{session}.docx"
    path.write_bytes(await file.read())
    doc = Document(str(path))
    return {"session": session, "blocks": doc_map(doc, index(doc))}


@app.post("/edit")
async def edit(prompt: str = Form(...), session: str = Form(...)):
    """Правка поверх out/{session}.docx: запрос делится на атомарные правки
    (edit.split), каждая проходит полный цикл edit.run_edit по очереди —
    правки видят результат предыдущих (см. BUILD_PLAN: без параллельного
    исполнения). Итог сохраняется на диск, чтобы /preview и /download
    отдавали то, что реально получилось."""
    path = _session_path(session)
    doc = Document(str(path))
    idx = index(doc)
    tasks = edit_mod.split(prompt)
    model = _env()["LLM_MODEL"]

    async def stream():
        nonlocal doc, idx
        started = prev = time.perf_counter()
        yield json.dumps({"type": "planning"}) + "\n"

        done, failed = [], []
        for task_n, task in enumerate(tasks, start=1):
            # run_edit синхронный и может считать до минуты (кавычки/типографика по
            # многим блокам) — уводим его с event loop в поток, чтобы между байтами
            # успевать слать пинг: облачный туннель рвёт соединение по тишине, не по
            # длительности (Ф10). Пинг игнорируется клиентом (неизвестный type) —
            # web/index.html не трогаем.
            edit_task = asyncio.create_task(asyncio.to_thread(edit_mod.run_edit, doc, idx, task))
            while not edit_task.done():
                try:
                    await asyncio.wait_for(asyncio.shield(edit_task), timeout=5)
                except asyncio.TimeoutError:
                    yield json.dumps({"type": "ping"}) + "\n"
            result, doc, idx = edit_task.result()
            doc.save(str(path))  # сразу после правки — файл и событие не должны расходиться
            log.append(session, {
                "task": task_n,
                "task_text": task,
                "model": model,
                "iter": result["iter"],
                "verdict": result["verdict"],
                "reason": result["reason"],
                "ids": result["ids"],
                # renderLogRecord кладёт reply прямо в <pre> и ждёт строку —
                # это РАЗОБРАННЫЙ JSON-ответ Навигатора/Редактора, не сырое тело HTTP.
                "reply": json.dumps(result["reply"], ensure_ascii=False),
            })
            now = time.perf_counter()
            # "already" — тоже честный успех (item A из Ф8: rule, ничего не поменявший,
            # потому что требуемое состояние уже в документе), а не провал: статус и
            # список "done" его не отличают от "done", "verdict" в событии несёт разницу.
            success = result["verdict"] in ("done", "already")
            text = "; ".join(result["applied"]) if result["verdict"] == "done" else result["reason"]
            (done if success else failed).append(text)
            yield json.dumps({
                "type": "op",
                "status": "done" if success else "failed",
                "text": text,
                "task": task_n,
                "task_text": task,
                "model": model,
                "at": round((now - started) * 1000),
                "dt": round((now - prev) * 1000),
                "verdict": result["verdict"],
                "blocks": doc_map(doc, idx),
            }) + "\n"
            prev = now

        yield json.dumps({
            "type": "result",
            "session": session,
            "done": done,
            "failed": failed,
            "at": round((time.perf_counter() - started) * 1000),
        }) + "\n"

    return StreamingResponse(stream(), media_type="application/x-ndjson")


@app.get("/logs")
def logs(session: str, n: int = 50):
    """Хвост журнала правок (docx_editor/log.py) — только заданной сессии.
    session обязателен: без него ручка отдавала бы журнал всех сессий (Ф10)."""
    return log.tail(session, n)


@app.get("/preview/{session}")
def preview(session: str):
    path = _session_path(session)
    doc = Document(str(path))
    return doc_map(doc, index(doc))


@app.get("/download/{session}")
def download(session: str):
    path = _session_path(session)
    return FileResponse(
        str(path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{session}.docx",
    )
