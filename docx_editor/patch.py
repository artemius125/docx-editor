"""Валидация патча (до применения) и применение операций к живому документу.

Патч — список операций из контракта BUILD_PLAN.md: replace_text, set_text,
insert_after, delete, move_after, set_style, create_table, set_cell, normalize,
replace_all, плюс из Ф15 — set_format (начертание рана — абзаца ИЛИ, с В2,
ячейки таблицы), set_list_level (уровень вложенности уже существующего
списка) и footnote (настоящая сноска Word: footnotes.xml + связь в .rels +
w:footnoteReference в теле), плюс из Ф19-бис — set_list (обычный абзац
становится элементом списка через уже существующую в документе нумерацию).

В2 (дыры контракта, найденные на первом же обычном офисном документе):
insert_row/delete_row/insert_col/delete_col — строка и колонка СУЩЕСТВУЮЩЕЙ
таблицы (create_table остаётся только для новой); insert_paragraphs —
несколько абзацев разными стилями одной операцией («раздел» = заголовок +
текст); set_format научился адресоваться ячейкой таблицы (id таблицы + row +
col); validate отбивает буквальный «\\n» в тексте любой операции — этим
подделывали то, чего нет в контракте (см. диагноз в BUILD_PLAN_V2.md).

В2-бис (поля Word и колонтитулы) — field (одна операция на все инструкции
поля: PAGE/NUMPAGES/DATE простым w:fldSimple, TOC/REF составным
w:fldChar begin/separate/end + w:instrText — механизм выбирает код по
инструкции, не вызывающий) и set_header_footer (текст колонтитула и,
опционально, поле внутри него — через doc.sections[-1].header/footer,
которые python-docx уже умеет заводить как отдельные части пакета).
"""

from copy import deepcopy
from difflib import SequenceMatcher

from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.table import Table
from docx.text.font import Font
from docx.text.paragraph import Paragraph

from docx_editor.find import _flex_span

_OPS = {
    "replace_text", "set_text", "insert_after", "delete", "move_after",
    "set_style", "create_table", "set_cell", "normalize", "replace_all",
    "set_format", "set_list_level", "footnote", "set_list",
    "insert_row", "delete_row", "insert_col", "delete_col", "insert_paragraphs",
    "field", "set_header_footer",
}
_PARAGRAPH_ONLY = {
    "replace_text", "set_text", "set_style", "set_list_level", "footnote", "set_list", "field",
}
# set_format не входит в _PARAGRAPH_ONLY (В2): адресуется и абзацем, и ячейкой
# таблицы, поэтому проверку его id/kind ведёт собственная ветка validate.
_TABLE_ONLY = {"set_cell", "insert_row", "delete_row", "insert_col", "delete_col"}
_NORMALIZE_RULES = {"typography", "quotes"}
# В2-бис: замкнутый словарь инструкций поля — "нет парсера выражений поля"
# означает ровно это: код не разбирает switches/аргументы инструкции, только
# проверяет, что она начинается с одного из пяти поддерживаемых кодов.
_FIELD_VERBS = {"PAGE", "NUMPAGES", "DATE", "TOC", "REF"}
_SIMPLE_FIELD_VERBS = {"PAGE", "NUMPAGES", "DATE"}


def _by_id(blocks):
    return {b["id"]: b for b in blocks}


def _table_texts(block):
    return [cell for row in block["rows"] for cell in row]


def _paragraph_style_names(doc):
    return [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]


def _style_error(style, doc):
    """None — стиль есть в документе; иначе текст ошибки со списком доступных."""
    names = _paragraph_style_names(doc)
    if style in names:
        return None
    return f"стиль {style!r} не найден; доступны стили абзацев: {', '.join(names)}"


def _numbering_ids(doc):
    """[numId, ...] нумераций, реально определённых в word/numbering.xml
    документа, [] если часть отсутствует (Ф19-бис: set_list превращает
    обычный абзац в элемент списка, ссылаясь на УЖЕ существующее определение
    нумерации — придумывать своё нельзя, а создание numbering.xml с нуля не
    поддержано python-docx, см. NumberingPart.new())."""
    try:
        part = doc.part.part_related_by(RT.NUMBERING)
    except KeyError:
        return []
    return [int(n.get(qn("w:numId"))) for n in part.element.findall(qn("w:num"))]


_ELLIPSIS = ("…", "...")


def _ends_with_ellipsis(text):
    return (text or "").rstrip().endswith(_ELLIPSIS)


def _ellipsis_truncation(old_text, new_text):
    """True — new/text обрывается многоточием, а исходный текст им не
    оканчивался: похоже на обрезанный хвост абзаца, выданный за полный текст
    (находка Ф8 — модель укоротила абзац 456→332 знака, дописав «…»,
    и Проверяющий это пропустил, потому что смысловая правка была на месте)."""
    return _ends_with_ellipsis(new_text) and not _ends_with_ellipsis(old_text)


def _ellipsis_error(block_id):
    return (
        f"{block_id}: текст обрывается многоточием («…» или «...»), хотя исходный текст "
        f"им не оканчивался — похоже, конец абзаца потерян при сокращении. Пришли текст "
        f"ЦЕЛИКОМ, без обрезания хвоста; многоточие в конце допустимо только тогда, когда "
        f"оно было в исходном тексте."
    )


def _mid_word_cut(haystack, span):
    """True — правая граница совпадения обрывается посреди слова: символ,
    которым заканчивается найденный кусок, и символ сразу за ним — оба
    словарные, то есть между ними нет границы слова (Ф12: `render` резал
    абзац на 300-м знаке, модель добросовестно копировала обрубок в «old»,
    `_flex_span` его находил — он ДЕЙСТВИТЕЛЬНО есть в документе, — а хвост
    слова оставался приклеен к новому тексту: «...требуют с» + «ерьёзных...»
    → «...новых мощностей, херьёзных...»).

    Это НЕ бьёт по легитимной правке суффикса внутри более длинного слова
    («энкодеры»→«энкодера» в «кросс-энкодеры»): там ЛЕВАЯ граница совпадения
    приходится на середину слова, а правая — на настоящий конец слова, и
    проверяется только правая. Отличить «модель обрубила цитату» от «модель
    намеренно правит часть слова» по одной лишь левой границе нельзя (оба
    случая режут слово слева одинаково) — только по правой, поэтому гвард
    смотрит исключительно на неё."""
    _, end = span
    return end < len(haystack) and haystack[end - 1].isalnum() and haystack[end].isalnum()


def _mid_word_error(block_id, old):
    where = f"{block_id}: " if block_id else ""
    return (
        f"{where}текст «{old}» обрывается посреди слова — символ сразу за совпадением "
        f"продолжает то же слово. Похоже, это обрезанный кусок (например, конец абзаца "
        f"был обрублен где-то в подсказке), а не целая фраза. Процитируй «old» целыми "
        f"словами, не разрезая последнее слово пополам."
    )


def _find_newline(value):
    """Первая строка с буквальным «\\n», найденная внутри value (рекурсивно —
    поля вроде rows/cells/items несут списки строк или списков строк), или None."""
    if isinstance(value, str):
        return value if "\n" in value else None
    if isinstance(value, list):
        for v in value:
            found = _find_newline(v)
            if found is not None:
                return found
    if isinstance(value, dict):
        for v in value.values():
            found = _find_newline(v)
            if found is not None:
                return found
    return None


def _newline_error(op):
    """В2, п.4: буквальный перевод строки внутри текста операции — не новый
    абзац, а именно так модель подделывает то, чего нет в контракте (найдено
    на «Добавь раздел „6. Пересмотр“» — Heading 1 с «6. Пересмотр\\nРегламент…»
    внутри одного абзаца). Ошибка называет это прямо, чтобы модель поправилась
    на insert_paragraphs вместо \\n."""
    for key, value in op.items():
        found = _find_newline(value)
        if found is not None:
            return (
                f"поле «{key}» операции {op.get('op')!r} содержит перевод строки (\\n) в тексте "
                f"«{found}» — перевод строки внутри абзаца не создаёт новый абзац, это ровно то, "
                f"чем подделывают операцию, которой нет в контракте. Раздели текст на отдельные "
                f"абзацы через insert_paragraphs (список {{\"text\",\"style\"}}), а не через \\n "
                f"внутри одной строки."
            )
    return None


_REQUIRED_TEXT = {
    # Поля-строки, без которых операция не просто бессмысленна, а роняет код:
    # _flex_span(text, None) и _replace_span(..., None) дают TypeError. Модель
    # присылает такое редко, но одна операция без old убивала ВЕСЬ прогон из
    # 20 правок — bench/run.py ловит только обрыв транспорта.
    "replace_text": ("old", "new"), "replace_all": ("old", "new"),
    "set_format": ("old",), "footnote": ("old", "text"),
    "set_text": ("text",), "insert_after": ("text",),
    "field": ("instr",), "set_header_footer": ("text",),
}


def _field_verb_error(instr):
    """None — instr начинается с одного из пяти поддерживаемых кодов поля,
    иначе текст ошибки. Код НЕ разбирает switches/аргументы инструкции
    («нет парсера выражений поля», В2-бис) — только первое слово."""
    verb = instr.split()[0].upper() if instr.split() else ""
    if verb in _FIELD_VERBS:
        return None
    return f"инструкция поля {instr!r} не начинается с одного из поддерживаемых кодов: {sorted(_FIELD_VERBS)}"


def validate(blocks, op, doc):
    """None — патч валиден; иначе русский текст ошибки для повторной попытки модели.

    doc нужен только для set_style: полный список стилей абзацев, реально
    определённых в документе, берётся из doc.styles — в blocks есть только
    стили, УЖЕ применённые к каким-то абзацам, а этого недостаточно (в
    Архитектура_ColBERT.docx все абзацы Normal, но List Paragraph в
    документе определён и доступен).
    """
    name = op.get("op")
    if name not in _OPS:
        return f"неизвестная операция {name!r}; допустимые: {sorted(_OPS)}"

    err = _newline_error(op)
    if err:
        return err

    by_id = _by_id(blocks)

    if name in _PARAGRAPH_ONLY or name in ("insert_after", "delete", "move_after", "insert_paragraphs"):
        block_id = op.get("id")
        if block_id not in by_id:
            return f"блок {block_id!r} не найден в документе"
        if name in _PARAGRAPH_ONLY and by_id[block_id]["kind"] != "p":
            return f"{block_id} — это таблица, {name} работает только с абзацами"

    if name in _TABLE_ONLY:
        block_id = op.get("id")
        if block_id not in by_id:
            return f"блок {block_id!r} не найден в документе"
        if by_id[block_id]["kind"] != "t":
            return f"{block_id} — это абзац, {name} работает только с таблицами"

    if name == "set_text" and _ellipsis_truncation(by_id[block_id]["text"], op.get("text")):
        return _ellipsis_error(block_id)

    if name == "insert_after" and op.get("style"):
        err = _style_error(op["style"], doc)
        if err:
            return err

    if name == "insert_paragraphs":
        items = op.get("items")
        if not isinstance(items, list) or not items:
            return "insert_paragraphs: items не может быть пустым списком"
        for i, item in enumerate(items):
            if not isinstance(item, dict) or not isinstance(item.get("text"), str):
                return f"insert_paragraphs: items[{i}] должен быть объектом со строковым полем text"
            style = item.get("style")
            if style:
                err = _style_error(style, doc)
                if err:
                    return err

    if name == "move_after":
        after = op.get("after")
        if after not in by_id:
            return f"блок {after!r} (after) не найден в документе"

    if name == "create_table":
        after = op.get("after")
        if after not in by_id:
            return f"блок {after!r} (after) не найден в документе"
        rows = op.get("rows")
        if not rows or not rows[0]:
            return "create_table: rows не может быть пустым"
        ncols = len(rows[0])
        if any(len(r) != ncols for r in rows):
            return f"create_table: строки разной длины: {[len(r) for r in rows]}"

    if name == "set_cell":
        block = by_id[op["id"]]
        nrows = len(block["rows"])
        ncols = len(block["rows"][0]) if nrows else 0
        row, col = op.get("row"), op.get("col")
        if not (isinstance(row, int) and 0 <= row < nrows) or not (isinstance(col, int) and 0 <= col < ncols):
            return f"ячейка ({row},{col}) вне таблицы {op['id']}: размер {nrows}x{ncols}"

    if name in ("insert_row", "delete_row"):
        block = by_id[op["id"]]
        nrows = len(block["rows"])
        ncols = len(block["rows"][0]) if nrows else 0
        if name == "insert_row":
            at = op.get("at")
            if not (isinstance(at, int) and not isinstance(at, bool) and 0 <= at <= nrows):
                return f"at должен быть целым числом от 0 до {nrows} (строк в {op['id']}: {nrows}), получено {at!r}"
            cells = op.get("cells")
            if not isinstance(cells, list) or len(cells) != ncols or not all(isinstance(v, str) for v in cells):
                return f"cells должен быть списком из {ncols} строк (по числу колонок {op['id']}), получено {cells!r}"
        else:
            row = op.get("row")
            if not (isinstance(row, int) and not isinstance(row, bool) and 0 <= row < nrows):
                return f"row вне таблицы {op['id']}: строк {nrows}, получено {row!r}"
            if nrows <= 1:
                return f"{op['id']}: нельзя удалить последнюю оставшуюся строку таблицы"

    if name in ("insert_col", "delete_col"):
        block = by_id[op["id"]]
        nrows = len(block["rows"])
        ncols = len(block["rows"][0]) if nrows else 0
        if name == "insert_col":
            at = op.get("at")
            if not (isinstance(at, int) and not isinstance(at, bool) and 0 <= at <= ncols):
                return f"at должен быть целым числом от 0 до {ncols} (колонок в {op['id']}: {ncols}), получено {at!r}"
            cells = op.get("cells")
            if not isinstance(cells, list) or len(cells) != nrows or not all(isinstance(v, str) for v in cells):
                return f"cells должен быть списком из {nrows} строк (по числу строк {op['id']}), получено {cells!r}"
        else:
            col = op.get("col")
            if not (isinstance(col, int) and not isinstance(col, bool) and 0 <= col < ncols):
                return f"col вне таблицы {op['id']}: колонок {ncols}, получено {col!r}"
            if ncols <= 1:
                return f"{op['id']}: нельзя удалить последнюю оставшуюся колонку таблицы"

    for field in _REQUIRED_TEXT.get(name, ()):
        if not isinstance(op.get(field), str):
            return f"в операции {name} нет поля «{field}» со строкой"

    if name in ("replace_text", "replace_all"):
        old, new = op.get("old"), op.get("new")
        if old == new:
            return f"old и new совпадают («{old}») — пустая операция, нечего менять"
        if name == "replace_text":
            text = by_id[op["id"]]["text"]
            span = _flex_span(text, old)
            if span is None:
                return f"в {op['id']} нет текста «{old}»"
            if _mid_word_cut(text, span):
                return _mid_word_error(op["id"], old)
            if _ellipsis_truncation(old, new):
                return _ellipsis_error(op["id"])
        else:
            found = cut = False
            for b in blocks:
                for t in ([b["text"]] if b["kind"] == "p" else _table_texts(b)):
                    span = _flex_span(t, old)
                    if span is not None:
                        found = True
                        cut = cut or _mid_word_cut(t, span)
            if not found:
                return f"текст «{old}» не найден нигде в документе"
            if cut:
                return _mid_word_error(None, old)

    if name == "set_format":
        block_id = op.get("id")
        if block_id not in by_id:
            return f"блок {block_id!r} не найден в документе"
        if op.get("b") is None and op.get("i") is None and op.get("u") is None:
            return f"{block_id}: не указано ни одного из b/i/u — операция ничего не меняет"
        block = by_id[block_id]
        if block["kind"] == "t":
            # В2, п.3: set_format адресуется ячейкой таблицы (row+col) — то же,
            # чем set_cell находит ячейку для замены текста.
            nrows = len(block["rows"])
            ncols = len(block["rows"][0]) if nrows else 0
            row, col = op.get("row"), op.get("col")
            if not (isinstance(row, int) and 0 <= row < nrows) or not (isinstance(col, int) and 0 <= col < ncols):
                return f"ячейка ({row},{col}) вне таблицы {block_id}: размер {nrows}x{ncols}"
            text = block["rows"][row][col]
        else:
            text = block["text"]
        old = op.get("old")
        span = _flex_span(text, old)
        if span is None:
            return f"в {block_id} нет текста «{old}»"
        if _mid_word_cut(text, span):
            return _mid_word_error(block_id, old)

    if name == "footnote":
        if not op.get("text"):
            return f"{op['id']}: не указан текст сноски"
        old = op.get("old")
        text = by_id[op["id"]]["text"]
        span = _flex_span(text, old)
        if span is None:
            return f"в {op['id']} нет текста «{old}»"
        if _mid_word_cut(text, span):
            return _mid_word_error(op["id"], old)

    if name == "field":
        err = _field_verb_error(op["instr"])
        if err:
            return err
        old = op.get("old")
        if old:
            text = by_id[op["id"]]["text"]
            span = _flex_span(text, old)
            if span is None:
                return f"в {op['id']} нет текста «{old}»"
            if _mid_word_cut(text, span):
                return _mid_word_error(op["id"], old)

    if name == "set_header_footer":
        which = op.get("which")
        if which not in ("header", "footer"):
            return f"which должен быть «header» или «footer», получено {which!r}"
        field = op.get("field")
        if field:
            err = _field_verb_error(field)
            if err:
                return err
        if not op.get("text") and not field:
            return "ни text, ни field не указаны — операция ничего не меняет"

    if name == "set_list_level":
        ilvl = op.get("ilvl")
        if not (isinstance(ilvl, int) and not isinstance(ilvl, bool) and ilvl >= 0):
            return f"ilvl должен быть неотрицательным целым числом, получено {ilvl!r}"
        if by_id[op["id"]]["list"] is None:
            return f"{op['id']} не является элементом списка (нет w:numPr) — сменить уровень нельзя"

    if name == "set_list":
        if by_id[op["id"]]["list"] is not None:
            return f"{op['id']} уже является элементом списка — используй set_list_level, чтобы сменить уровень"
        ilvl = op.get("ilvl", 0)
        if not (isinstance(ilvl, int) and not isinstance(ilvl, bool) and ilvl >= 0):
            return f"ilvl должен быть неотрицательным целым числом, получено {ilvl!r}"
        if not _numbering_ids(doc):
            return ("в документе нет ни одной определённой нумерации (word/numbering.xml отсутствует "
                     "или пуст) — сделать абзац элементом списка нечем")

    if name == "set_style":
        err = _style_error(op.get("style"), doc)
        if err:
            return err

    if name == "normalize":
        rule = op.get("rule")
        if rule not in _NORMALIZE_RULES:
            return f"неизвестное правило {rule!r}; допустимые: {sorted(_NORMALIZE_RULES)}"

    return None


def _runs(p_el):
    """Все раны абзаца в порядке документа, включая раны внутри w:hyperlink.

    p_el.r_lst — это только раны верхнего уровня; текст гиперссылки в него
    не попадает, а p.text (то, что видит validate и модель) гиперссылки
    включает. Расхождение офсетов между _ptext и p.text — гарантированная
    порча документа (find возвращает -1 или не то смещение). iter() обходит
    поддерево в порядке документа, поэтому раны внутри w:hyperlink встают
    на своё естественное место и офсеты снова совпадают с p.text.
    """
    return list(p_el.iter(qn("w:r")))


def _ptext(p_el):
    """Текст абзаца, собранный ровно из тех же ранов, что и _replace_span."""
    return "".join(r.text for r in _runs(p_el))


def _run_spans(p_el):
    """Раны абзаца и их смещения [start,end) в тексте — общая арифметика,
    которую делят _replace_span и set_format (_split_span_runs)."""
    runs = _runs(p_el)
    spans, pos = [], 0
    for r in runs:
        t = r.text
        spans.append((pos, pos + len(t)))
        pos += len(t)
    return runs, spans


def _touched_runs(spans, start, end):
    return [i for i, (s, e) in enumerate(spans) if e > start and s < end]


def _replace_span(p_el, start, end, new_text):
    """Заменяет [start:end) в тексте абзаца на new_text, не трогая чужие раны.

    Граница замены может резать чужой ран пополам (фраза лежит на стыке
    ранов с разным начертанием) — тогда затронутые крайние раны режутся по
    границе, а раны целиком внутри диапазона стираются, чужие раны снаружи
    не трогаются вовсе.
    """
    if start < 0 or end < start:
        raise ValueError(f"недопустимый диапазон замены [{start}:{end})")
    runs, spans = _run_spans(p_el)
    if not runs:
        p_el.add_r().text = new_text
        return
    touched = _touched_runs(spans, start, end)
    if not touched:
        runs[-1].text = runs[-1].text + new_text
        return
    fi, li = touched[0], touched[-1]
    prefix = runs[fi].text[: start - spans[fi][0]]
    suffix = runs[li].text[end - spans[li][0]:]
    if fi == li:
        runs[fi].text = prefix + new_text + suffix
    else:
        runs[fi].text = prefix + new_text
        runs[li].text = suffix
        for i in range(fi + 1, li):
            runs[i].text = ""


def _split_run_at(r, cut):
    """Физически режет ран r на два по смещению cut внутри его текста: левая
    часть остаётся в r, правая уезжает в новый ран с той же rPr (форматирование
    копируется), вставленный сразу за r. cut<=0 или cut>=len(text) — резать
    нечего, ран не трогается. Нужен set_format: в отличие от _replace_span
    (который просто переписывает текст рана, не создавая новых элементов),
    здесь префикс/спан/суффикс должны получить РАЗНОЕ форматирование, а у
    одного <w:r> одна rPr на всех — значит, ран необходимо реально разделить."""
    text = r.text
    if cut <= 0 or cut >= len(text):
        return
    r.text = text[:cut]
    new = OxmlElement("w:r")
    rPr = r.find(qn("w:rPr"))
    if rPr is not None:
        new.append(deepcopy(rPr))
    r.addnext(new)
    new.text = text[cut:]


def _split_span_runs(p_el, start, end):
    """Возвращает раны, физически покрывающие ровно [start:end) — крайние
    затронутые раны при необходимости режутся по границам спана
    (_split_run_at), раны целиком внутри диапазона отдаются как есть. Правая
    граница режется первой: её смещение не зависит от левого разреза, а
    вставка нового рана после li не сдвигает индексы левее li."""
    runs, spans = _run_spans(p_el)
    touched = _touched_runs(spans, start, end)
    fi, li = touched[0], touched[-1]
    _split_run_at(runs[li], end - spans[li][0])
    _split_run_at(runs[fi], start - spans[fi][0])
    runs, spans = _run_spans(p_el)
    return [r for r, (s, e) in zip(runs, spans) if s >= start and e <= end]


def _register(idx, prefix, el):
    n = 0
    while f"{prefix}{n}" in idx:
        n += 1
    new_id = f"{prefix}{n}"
    idx[new_id] = el
    return new_id


def _op_replace_text(doc, idx, op):
    el = idx[op["id"]]
    full = _ptext(el)
    span = _flex_span(full, op["old"])
    if span is None:
        # validate должен был отсечь это раньше; если добрались сюда — громко падаем,
        # а не режем абзац по отрицательному смещению (см. находку про гиперссылки)
        raise ValueError(f"текст {op['old']!r} не найден в {op['id']} на момент применения")
    # end берётся из реального совпадения (_flex_span), а не из len(op["old"]):
    # из-за разночтений в пробелах (обычный ↔ U+00A0) найденный в документе
    # фрагмент может быть длиннее или короче исходного needle
    _replace_span(el, span[0], span[1], op["new"])
    return f"В {op['id']} заменено «{op['old']}» на «{op['new']}»"


def _op_set_text(doc, idx, op):
    # В3: Ф19 велит переписывать абзац целиком, но старый/новый текст обычно
    # разделяют длинные неизменившиеся куски (общий сосед фразы), а не только
    # первый/последний символ — старый код заменял [0:len(old)) одним куском,
    # из-за чего ВЕСЬ новый текст оказывался в ранe САМОГО ПЕРВОГО совпавшего
    # рана (какое бы начертание у него ни было), а раны с остальным
    # начертанием абзаца становились пустыми — начертание физически
    # оставалось в документе, но на пустом ранe, невидимо (архитектор поймал
    # это на живом жирном «0.5%», которое "переехало" в пустой ран).
    #
    # Правка: не трогаем совпавшие (SequenceMatcher, tag=="equal") куски
    # вообще — их раны остаются как были, с текстом и начертанием. Только
    # реально изменившиеся диапазоны переписываются _replace_span, той же
    # функцией и с тем же ограничением, что и replace_text: если диапазон
    # целиком перекрывает несколько по-разному оформленных ранов, новый текст
    # наследует начертание ПЕРВОГО из них (это старое, уже проверенное
    # поведение _replace_span, не новое допущение). Диапазоны идут в
    # обратном порядке (по убыванию начала), чтобы более ранние офсеты не
    # съехали от более поздней замены — тот же приём, что в _op_replace_all.
    el = idx[op["id"]]
    old_text, new_text = _ptext(el), op["text"]
    opcodes = SequenceMatcher(None, old_text, new_text, autojunk=False).get_opcodes()
    for tag, i1, i2, j1, j2 in reversed(opcodes):
        if tag != "equal":
            _replace_span(el, i1, i2, new_text[j1:j2])
    return f"Текст {op['id']} заменён целиком"


def _op_insert_after(doc, idx, op):
    ref = idx[op["id"]]
    style = op.get("style") or Paragraph(ref, doc).style
    new_p = doc.add_paragraph(op["text"], style=style)
    ref.addnext(new_p._p)
    new_id = _register(idx, "p", new_p._p)
    return f"После {op['id']} вставлен новый абзац {new_id}"


def _op_insert_paragraphs(doc, idx, op):
    # В2, п.2: «раздел» — заголовок плюс текст, минимум два абзаца двух
    # стилей одной операцией; без style у элемента — обычный стиль по
    # умолчанию (в отличие от insert_after, здесь стиль каждого абзаца несёт
    # смысл операции, наследовать стиль соседа было бы неверно).
    ref = idx[op["id"]]
    new_ids = []
    for item in op["items"]:
        new_p = doc.add_paragraph(item["text"], style=item.get("style"))
        ref.addnext(new_p._p)
        ref = new_p._p
        new_ids.append(_register(idx, "p", new_p._p))
    return f"После {op['id']} вставлено {len(new_ids)} абзацев: {', '.join(new_ids)}"


def _op_delete(doc, idx, op):
    el = idx[op["id"]]
    el.getparent().remove(el)
    return f"Блок {op['id']} удалён"


def _op_move_after(doc, idx, op):
    src, dst = idx[op["id"]], idx[op["after"]]
    dst.addnext(src)
    return f"Блок {op['id']} перемещён после {op['after']}"


def _op_set_style(doc, idx, op):
    Paragraph(idx[op["id"]], doc).style = op["style"]
    return f"Стиль {op['id']} изменён на {op['style']!r}"


def _op_set_format(doc, idx, op):
    el = idx[op["id"]]
    if el.tag == qn("w:tbl"):
        # В2, п.3: ячейка таблицы — тот же способ найти её абзац, что у set_cell
        el = Table(el, doc).cell(op["row"], op["col"]).paragraphs[0]._p
    span = _flex_span(_ptext(el), op["old"])
    if span is None:
        # validate должен был отсечь это раньше (см. _op_replace_text) — громко падаем.
        raise ValueError(f"текст {op['old']!r} не найден в {op['id']} на момент применения")
    for r in _split_span_runs(el, *span):
        font = Font(r)
        if op.get("b") is not None:
            font.bold = op["b"]
        if op.get("i") is not None:
            font.italic = op["i"]
        if op.get("u") is not None:
            font.underline = op["u"]
    flags = ", ".join(f"{k}={op[k]}" for k in ("b", "i", "u") if op.get(k) is not None)
    return f"В {op['id']} для «{op['old']}» установлено форматирование: {flags}"


def _op_set_list_level(doc, idx, op):
    numPr = idx[op["id"]].find(qn("w:pPr")).find(qn("w:numPr"))
    ilvl_el = numPr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.insert(0, ilvl_el)  # w:numId уже стоит первым или единственным — не трогаем
    ilvl_el.set(qn("w:val"), str(op["ilvl"]))
    return f"Уровень списка {op['id']} изменён на {op['ilvl']}"


def _op_set_list(doc, idx, op):
    # get_or_add_numPr (в отличие от pPr.append) ставит w:numPr в схемную
    # позицию pPr сам — важно, если у абзаца уже есть другие свойства
    # (w:jc, w:spacing и т.п.), которые обязаны идти ПОСЛЕ w:numPr.
    numPr = idx[op["id"]].get_or_add_pPr().get_or_add_numPr()
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(op.get("ilvl", 0)))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(_numbering_ids(doc)[0]))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    return f"Абзац {op['id']} сделан элементом списка (уровень {op.get('ilvl', 0)})"


_FOOTNOTES_PARTNAME = PackURI("/word/footnotes.xml")
# id -1 и 0 — зарезервированы OOXML под разделитель/разделитель-продолжение
# (17.3.1.11); настоящие сноски нумеруются с 1, чтобы никогда с ними не
# столкнуться. Word сам создаёт такую же пару при первой сноске в документе.
_FOOTNOTES_XML = (
    b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    b'<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
    b'<w:footnote w:type="continuationSeparator" w:id="0">'
    b'<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
    b'</w:footnotes>'
)


def _footnotes_part(doc):
    """Часть footnotes.xml документа: существующая или только что созданная.

    [Content_Types].xml python-docx строит сам при save() из part.content_type
    каждой части (см. docx.opc.pkgwriter._ContentTypesItem) — отдельно его
    редактировать не нужно. Связь в document.xml.rels создаёт relate_to.

    Если документ УЖЕ содержал footnotes.xml (открыт заново после чужой правки
    или пришёл с сносками из Word), PartFactory не регистрирует под
    WML_FOOTNOTES собственный класс — в python-docx нет API сносок, — и часть
    приходит как обобщённый Part с сырыми байтами, без живого lxml-дерева.
    Оборачиваем её в XmlPart и подменяем цель УЖЕ существующей связи (тот же
    rId, add_relationship с тем же rId просто переписывает запись), чтобы не
    завести вторую связь на ту же часть.
    """
    try:
        part = doc.part.part_related_by(RT.FOOTNOTES)
    except KeyError:
        part = XmlPart(_FOOTNOTES_PARTNAME, CT.WML_FOOTNOTES, parse_xml(_FOOTNOTES_XML), doc.part.package)
        doc.part.relate_to(part, RT.FOOTNOTES)
        return part
    if isinstance(part, XmlPart):
        return part
    xml_part = XmlPart(part.partname, part.content_type, parse_xml(part.blob), doc.part.package)
    rId = next(rid for rid, rel in doc.part.rels.items() if rel.reltype == RT.FOOTNOTES)
    doc.part.rels.add_relationship(RT.FOOTNOTES, xml_part, rId)
    return xml_part


def _add_footnote(doc, text):
    """Добавляет новую w:footnote в footnotes.xml (создав часть при первой
    сноске), возвращает её id. Существующие сноски не трогает — только append."""
    part = _footnotes_part(doc)
    used_ids = [int(fn.get(qn("w:id"))) for fn in part.element.findall(qn("w:footnote"))]
    fid = max(used_ids, default=0) + 1

    fn = OxmlElement("w:footnote")
    fn.set(qn("w:id"), str(fid))
    p = OxmlElement("w:p")
    ref_run = OxmlElement("w:r")
    ref_rPr = OxmlElement("w:rPr")
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    ref_rPr.append(vert)
    ref_run.append(ref_rPr)
    ref_run.append(OxmlElement("w:footnoteRef"))
    p.append(ref_run)
    text_run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    body = f" {text}"
    if len(body.strip()) < len(body):
        t.set(qn("xml:space"), "preserve")
    t.text = body
    text_run.append(t)
    p.append(text_run)
    fn.append(p)
    part.element.append(fn)
    return fid


def _op_footnote(doc, idx, op):
    el = idx[op["id"]]
    span = _flex_span(_ptext(el), op["old"])
    if span is None:
        # validate должен был отсечь это раньше (см. _op_replace_text) — громко падаем.
        raise ValueError(f"текст {op['old']!r} не найден в {op['id']} на момент применения")
    anchor = _split_span_runs(el, *span)[-1]

    fid = _add_footnote(doc, op["text"])
    ref = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    rPr.append(vert)
    ref.append(rPr)
    fref = OxmlElement("w:footnoteReference")
    fref.set(qn("w:id"), str(fid))
    ref.append(fref)
    anchor.addnext(ref)
    return f"В {op['id']} после «{op['old']}» добавлена сноска {fid}: «{op['text']}»"


def _fldchar_run(kind):
    r = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), kind)
    r.append(fld)
    return r


_FIELD_CAVEAT = ("значения не будет, пока документ не открыт в Word/LibreOffice и поля не "
                 "обновлены (выделить всё — Ctrl+A, затем F9)")


def _field_elements(instr):
    """Элементы OOXML для поля Word с инструкцией instr: простое w:fldSimple
    для PAGE/NUMPAGES/DATE (схема допускает их без вложенной структуры),
    составное w:fldChar begin/separate/end + w:instrText — для остальных
    (TOC, REF), где простой формы недостаточно (В2-бис, BUILD_PLAN_V2.md).

    Кэш-значение внутри поля сознательно не кладём: реальный номер страницы
    или запись оглавления посчитать нечем без движка вёрстки (см. решение
    честности в В2-бис) — Word/LibreOffice вычисляют его сами при открытии
    или обновлении полей."""
    verb = instr.split()[0].upper()
    if verb in _SIMPLE_FIELD_VERBS:
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        return [fld]
    instr_r = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {instr} "
    instr_r.append(instrText)
    return [_fldchar_run("begin"), instr_r, _fldchar_run("separate"), _fldchar_run("end")]


def _op_field(doc, idx, op):
    el, instr = idx[op["id"]], op["instr"]
    old = op.get("old")
    if old:
        span = _flex_span(_ptext(el), old)
        if span is None:
            # validate должен был отсечь это раньше (см. _op_replace_text) — громко падаем.
            raise ValueError(f"текст {old!r} не найден в {op['id']} на момент применения")
        ref = _split_span_runs(el, *span)[-1]
        for e in _field_elements(instr):
            ref.addnext(e)
            ref = e
        where = f"после «{old}» в {op['id']}"
    else:
        for e in _field_elements(instr):
            el.append(e)
        where = f"в конец {op['id']}"
    return f"Поле «{instr}» вставлено {where} — {_FIELD_CAVEAT}"


def _op_set_header_footer(doc, idx, op):
    which, text, field = op["which"], op.get("text", ""), op.get("field")
    target = doc.sections[-1].header if which == "header" else doc.sections[-1].footer
    target.is_linked_to_previous = False  # заводит отдельную часть пакета и ссылку в w:sectPr
    p = target.paragraphs[0]
    p.text = text
    if field:
        for e in _field_elements(field):
            p._p.append(e)
    ru_which = "Верхний" if which == "header" else "Нижний"
    note = f'; поле «{field}» — {_FIELD_CAVEAT}' if field else ""
    return (f"{ru_which} колонтитул изменён{note} — в предпросмотре он не отображается "
            f"(предпросмотр показывает только тело документа), правку видно в скачанном файле")


def _add_borders(tbl):
    """Одиночные границы прямо в tblPr. В контракте create_table нет поля
    style, а из именованных табличных стилей в документе часто есть только
    Normal Table (без рамок) — рассчитывать на style="Table Grid" нельзя,
    его может не быть (KeyError). Ставим границы всегда, без опций."""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl.tblPr.append(borders)


def _op_create_table(doc, idx, op):
    ref = idx[op["after"]]
    rows = op["rows"]
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    _add_borders(table._tbl)
    for r, row_vals in enumerate(rows):
        for c, val in enumerate(row_vals):
            table.cell(r, c).text = val
    if op.get("header"):
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    ref.addnext(table._tbl)
    new_id = _register(idx, "t", table._tbl)
    return f"После {op['after']} создана таблица {new_id} {len(rows)}x{ncols}"


def _op_set_cell(doc, idx, op):
    table = Table(idx[op["id"]], doc)
    p_el = table.cell(op["row"], op["col"]).paragraphs[0]._p
    _replace_span(p_el, 0, len(_ptext(p_el)), op["text"])
    return f"Ячейка {op['id']}[{op['row']}][{op['col']}] заменена на «{op['text']}»"


def _op_insert_row(doc, idx, op):
    # В2, п.1: строка СУЩЕСТВУЮЩЕЙ таблицы — то, чем модель раньше подделывала
    # правку абзацем со своей нотацией "r4c0:… | r4c1:…". table.add_row()
    # добавляет строку в конец (как и рамки — см. _add_borders, общие на
    # всю таблицу, per-cell ничего добавлять не нужно); если целевая позиция
    # не последняя, physически переносим новый <w:tr> перед нужным соседом —
    # тот же приём addprevious/addnext, что и у move_after/insert_after.
    table = Table(idx[op["id"]], doc)
    at, cells = op["at"], op["cells"]
    nrows_before = len(table.rows)
    new_row = table.add_row()
    if at < nrows_before:
        table.rows[at]._tr.addprevious(new_row._tr)
    for c, val in enumerate(cells):
        table.cell(at, c).text = val
    return f"В {op['id']} вставлена строка на позицию {at}: {cells}"


def _op_delete_row(doc, idx, op):
    tr = Table(idx[op["id"]], doc).rows[op["row"]]._tr
    tr.getparent().remove(tr)
    return f"Из {op['id']} удалена строка {op['row']}"


def _op_insert_col(doc, idx, op):
    # Аналог insert_row по колонке: своего "insert в середину" метода у
    # python-docx нет (Table.add_column только дописывает справа), поэтому
    # gridCol и по одной ячейке в каждой строке добавляются напрямую в XML
    # (add_gridCol/add_tc — то же автогенерируемое API, которым пользуется
    # сам python-docx в Table.add_row/add_column) и при необходимости
    # переносятся на нужную позицию тем же addprevious.
    tbl = idx[op["id"]]
    at, cells = op["at"], op["cells"]
    grid_cols = tbl.tblGrid.gridCol_lst
    ncols_before = len(grid_cols)
    ref_col = grid_cols[min(at, ncols_before - 1)]
    new_col = tbl.tblGrid.add_gridCol()
    if ref_col.w is not None:
        new_col.w = ref_col.w
    if at < ncols_before:
        ref_col.addprevious(new_col)
    for tr in tbl.tr_lst:
        tcs = tr.tc_lst
        ref_tc = tcs[min(at, len(tcs) - 1)]
        new_tc = tr.add_tc()
        if at < len(tcs):
            ref_tc.addprevious(new_tc)
    table = Table(tbl, doc)
    for r, val in enumerate(cells):
        table.cell(r, at).text = val
    return f"В {op['id']} вставлена колонка на позицию {at}: {cells}"


def _op_delete_col(doc, idx, op):
    tbl = idx[op["id"]]
    col = op["col"]
    grid_col = tbl.tblGrid.gridCol_lst[col]
    grid_col.getparent().remove(grid_col)
    for tr in tbl.tr_lst:
        tc = tr.tc_lst[col]
        tc.getparent().remove(tc)
    return f"Из {op['id']} удалена колонка {col}"


def _op_replace_all(doc, idx, op):
    old, new = op["old"], op["new"]
    total = blocks_touched = 0
    for p_el in doc.element.body.iter(qn("w:p")):
        full = _ptext(p_el)
        matches, pos = [], 0
        while True:
            span = _flex_span(full[pos:], old)
            if span is None:
                break
            m_start, m_end = pos + span[0], pos + span[1]
            pos += span[1]
            matches.append((m_start, m_end))
        if not matches:
            continue
        for m_start, m_end in reversed(matches):
            _replace_span(p_el, m_start, m_end, new)
        if _ptext(p_el) == full:
            # найдено, но применение не изменило текст (например, old/new
            # расходятся только в пробелах, которые _flex_span сам нормализует
            # при поиске) — отчёт не должен приписывать себе правку, которой
            # не произошло (инвариант 5, «отчёт не врёт»)
            continue
        total += len(matches)
        blocks_touched += 1
    return f"Заменено {total} вхождений «{old}» на «{new}» в {blocks_touched} блоках (включая ячейки таблиц)"


def _op_normalize(doc, idx, op):
    # Ленивый импорт: rules.py импортирует _ptext/_replace_span из этого модуля,
    # импорт на верхнем уровне создал бы цикл.
    from docx_editor.rules import quotes, typography

    return typography(doc) if op["rule"] == "typography" else quotes(doc)


_HANDLERS = {
    "replace_text": _op_replace_text,
    "set_text": _op_set_text,
    "insert_after": _op_insert_after,
    "delete": _op_delete,
    "move_after": _op_move_after,
    "set_style": _op_set_style,
    "create_table": _op_create_table,
    "set_cell": _op_set_cell,
    "normalize": _op_normalize,
    "replace_all": _op_replace_all,
    "set_format": _op_set_format,
    "set_list_level": _op_set_list_level,
    "footnote": _op_footnote,
    "set_list": _op_set_list,
    "insert_row": _op_insert_row,
    "delete_row": _op_delete_row,
    "insert_col": _op_insert_col,
    "delete_col": _op_delete_col,
    "insert_paragraphs": _op_insert_paragraphs,
    "field": _op_field,
    "set_header_footer": _op_set_header_footer,
}


def apply(doc, idx, op):
    """Применяет одну из операций контракта (патч уже прошёл validate), возвращает описание."""
    return _HANDLERS[op["op"]](doc, idx, op)
