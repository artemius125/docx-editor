"""Приёмка bench/oracle.py (В8): m19 (настоящая сноска Word) и c7 (одно
написание термина) были сломаны так, что НИ ОДИН документ не мог их
пройти — сама проверка была ложью оценщика. Строим фикстуры в памяти
(без сети, без живого прогона) и требуем PASS на верном результате и
FAIL на неверном — так, что каждая проверка реально способна и пройти,
и провалиться."""

import io
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "bench"))
from oracle import c7, m19  # noqa: E402

from docx_editor.parse import index  # noqa: E402
from docx_editor.patch import apply  # noqa: E402


def _footnote_docx():
    """Настоящая сноска Word — тем же путём, что и в проде (docx_editor.patch,
    op "footnote"): footnotes.xml + связь в .rels + w:footnoteReference."""
    doc = Document()
    doc.add_paragraph("Ландауэр вывел предел рассеиваемой энергии на бит информации.")
    idx = index(doc)
    op = {"op": "footnote", "id": "p0", "old": "Ландауэр вывел предел", "text": "Landauer, R. (1961)."}
    apply(doc, idx, op)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _fake_footnote_docx():
    """Старая болезнь: «сноска» — просто абзац стиля Footnote в теле, без
    footnotes.xml вообще. Ровно то, что настоящая проверка обязана ловить."""
    doc = Document()
    doc.add_paragraph("Ландауэр вывел предел рассеиваемой энергии на бит информации.")
    style = doc.styles.add_style("Footnote Text", WD_STYLE_TYPE.PARAGRAPH)
    p = doc.add_paragraph("Landauer, R. (1961).")
    p.style = style
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_m19_footnote_check():
    good, bad = _footnote_docx(), _fake_footnote_docx()
    assert m19([], [], [], [], docx_path=good) is True, "настоящая сноска обязана проходить проверку"
    assert m19([], [], [], [], docx_path=bad) is False, "абзац со стилем Footnote в теле — не настоящая сноска"
    assert m19([], [], [], [], docx_path=None) is None, "без пакета (seq-прогон) — честное n/a, а не провал"
    print("oracle_demo: m19 отличает настоящую сноску Word (footnotes.xml) от абзаца, притворяющегося ею")


def test_c7_single_spelling_check():
    consistent = "Здесь используется Single-Vector Bi-encoder. Позже снова Single-Vector Bi-encoder."
    inconsistent = "Здесь Single-Vector Bi-encoder. А тут отдельно Bi-Encoders без приставки."
    assert c7([], [consistent], [], []) is True, "единственное написание составного термина обязано проходить"
    assert c7([], [inconsistent], [], []) is False, "два разных написания обязаны провалить проверку"
    print("oracle_demo: c7 больше не проваливается на составном каноническом термине «Single-Vector Bi-encoder»")


if __name__ == "__main__":
    test_m19_footnote_check()
    test_c7_single_spelling_check()
