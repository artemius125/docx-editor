"""Приёмка edit.py: split() на настоящем файле правок, отказ/запасной поиск,
откат и правка 12 на настоящем документе — везде фейковые роли (без вызова
модели, см. CLAUDE.md: демо, закрываемое фикстурой, закрываем фикстурой).
Живая приёмка того же корпуса — examples/colbert_run.py, не *_demo.py.
"""

import os
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_editor import edit as edit_mod
from docx_editor import find, patch
from docx_editor.edit import run_edit, split
from docx_editor.parse import doc_map, index, render

REAL_DOC = "/home/artem/Загрузки/Архитектура_ColBERT.docx"
EDITS_FILE = "/home/artem/Загрузки/Правки_ColBERT_20.md"
REGLAMENT_DOC = "/home/artem/Документы/artemius125/docx-editor/bench/fixtures/Регламент.docx"


def test_split_real_file():
    text = open(EDITS_FILE, encoding="utf-8").read()
    items = split(text)
    assert len(items) == 20, f"ожидали 20 правок, получили {len(items)}"
    assert items[11].startswith("Написано «снижает требования")
    print("edit_demo: split() даёт ровно 20 правок из реального файла")


def _fake_doc():
    doc = Document()
    doc.add_paragraph("Обычный абзац без ничего примечательного.")
    idx = index(doc)
    return doc, idx


def test_fallback_search_then_honest_refusal():
    # Навигатор промахнулся: ids содержит несуществующий блок, anchor не
    # находится — код обязан САМ попробовать запасной поиск по дословной
    # цитате в «ёлочках» из текста правки. Цитата тоже отсутствует в
    # документе, поэтому итог — честный отказ, а не тихое "сделано".
    doc, idx = _fake_doc()
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p99"], "anchors": ["совсем другой текст"]}

    def fake_editor(*a, **kw):
        raise AssertionError("Редактор не должен вызываться — резолв не прошёл ни на одном этапе")

    request = "Правка по цитате «текст, которого точно нет в этом документе»."
    result, doc, idx = run_edit(doc, idx, request, navigator=fake_navigator, editor=fake_editor)

    assert result["verdict"] == "failed", result
    assert result["reason"], "у отказа должна быть причина"
    assert result["applied"] == [] and result["ids"] == []
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "документ не должен измениться при отказе"
    print(f"edit_demo: запасной поиск отработал, честный отказ: {result['reason']!r}")


def test_rollback_on_checker_reject():
    # Редактор и Навигатор — валидные фейки (предлагают реальную, применимую
    # операцию), а Проверяющий говорит "не то" — цикл обязан откатить документ
    # к снимку и вернуть rolled_back, а не оставить недоделанную правку.
    doc, idx = _fake_doc()
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "Обычный", "new": "Особенный"}]}

    def fake_checker(request, diff):
        return {"ok": False, "reason": "тестовый отказ проверяющего"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'Обычный' на 'Особенный'",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "rolled_back", result
    assert result["reason"] == "тестовый отказ проверяющего"
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "после отката текст обязан совпадать с исходным побайтово"
    print("edit_demo: откат по вердикту Проверяющего вернул документ как был")


def test_resolve_normalizes_integer_id():
    # Находка Ф5 (реопен, зеркало уже описанной в BUILD_PLAN "id внутрь
    # цитаты"): живой Навигатор вернул ids: [112] — ЧИСЛОМ, а не строкой
    # "p112". Старый _resolve сравнивал напрямую со множеством строк и тихо
    # терял такой id, резолв выживал только на якоре и уезжал не туда.
    # _resolve обязан нормализовать форму сам, а не полагаться на промпт.
    doc, idx = _fake_doc()
    reached = []

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": [0], "anchors": []}  # int, не "p0"

    def fake_editor(fragment_text, request, feedback=None):
        reached.append(fragment_text)
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "Обычный", "new": "Особенный"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'Обычный' на 'Особенный'",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert reached, "целочисленный id (0) от Навигатора обязан резолвиться в p0 и дойти до Редактора"
    assert result["verdict"] == "done", result
    print("edit_demo: целочисленный id от Навигатора (0) нормализован в p0 и резолвится")


def test_ellipsis_truncation_retries_to_full_text():
    # Фикстура реального провала приёмки (Ф8, правка 18/p35): живой Редактор
    # ответил set_text, укоротив абзац и оборвав его буквальным «…» —
    # смысловая правка была на месте, поэтому Проверяющий сказал ok, а
    # машинная сверка "текст изменился" такое по устройству не ловит.
    # validate (Layer 1) обязан отбить обрезанный ответ ДО применения и
    # вернуть Редактору текст ошибки; получив обратную связь, Редактор
    # обязан прислать текст ЦЕЛИКОМ — и цикл обязан ПРОЙТИ ретраем, а не
    # просто честно отказать.
    doc = Document()
    original = "Первое предложение абзаца. Второе предложение абзаца, длиннее первого для правдоподобия."
    doc.add_paragraph(original)
    idx = index(doc)

    truncated = "Первое предложение абзаца. Второе предложение…"
    full = "Первое предложение абзаца полностью исправлено. Второе предложение абзаца, длиннее первого для правдоподобия."

    feedbacks = []

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        feedbacks.append(feedback)
        if feedback is None:
            return {"ops": [{"op": "set_text", "id": "p0", "text": truncated}]}
        return {"ops": [{"op": "set_text", "id": "p0", "text": full}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "исправь первое предложение",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    assert result["iter"] == 2, "guard обязан отбить первый ответ и провести настоящий ретрай, а не убить правку"
    after_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p0")
    assert after_text == full, after_text
    assert not after_text.rstrip().endswith(("…", "...")), "результат не должен обрываться многоточием"
    assert len(after_text) >= len(original), "текст не должен схлопнуться короче исходного"
    assert len(feedbacks) == 2 and feedbacks[0] is None and feedbacks[1], \
        "второй вызов редактора обязан нести текст ошибки валидатора"
    print(f"edit_demo: обрезанный set_text отбит validate, ретрай дал полный текст, iter={result['iter']}")


def test_check_diff_carries_lengths():
    # Layer 2 (Ф8): _check обязан показывать Проверяющему длины "было"/"стало",
    # иначе числовой сигнал "текст вдвое короче" молча сгниёт при рефакторинге.
    from docx_editor.edit import _check

    seen = {}

    def fake_llm_chat(messages):
        seen["user"] = messages[1]["content"]
        return {"ok": True, "reason": "ok"}

    import docx_editor.llm as llm_module
    real_chat = llm_module.chat
    llm_module.chat = fake_llm_chat
    try:
        diff = [{"id": "p35", "before": "а" * 456, "after": "б" * 332}]
        _check("правка", diff)
    finally:
        llm_module.chat = real_chat

    assert "456" in seen["user"] and "332" in seen["user"], seen["user"]
    print("edit_demo: diff, отдаваемый Проверяющему, несёт длины было/стало")


def test_rule_no_op_yields_already():
    # Ф8 item A: правка типа rule, ничего не изменившая (документ уже
    # нормализован — как правка 3 "слипшиеся предложения" после того, как
    # правка 2 того же запроса уже прогнала typography по всему документу),
    # обязана дать verdict="already", а не "failed": требуемое состояние УЖЕ
    # в документе, и "не удалось" вводит в заблуждение не меньше ложного
    # "сделано" (инвариант 5).
    doc = Document()
    doc.add_paragraph("Уже нормальный абзац: без лишних пробелов; и без слипшихся предложений.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "global", "rule": "typography", "ids": [], "anchors": []}

    result, doc, idx = run_edit(doc, idx, "раздели слипшиеся предложения", navigator=fake_navigator)

    assert result["verdict"] == "already", result
    assert result["reason"], "у already обязана быть непустая причина"
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "already не должен менять документ"
    print(f"edit_demo: rule без изменений даёт already: {result['reason']!r}")


def test_non_rule_empty_diff_stays_failed():
    # Тот же симптом (ops применились, diff пуст), но НЕ через rule — здесь
    # "already" не полагается (см. правило item A): нет оснований заключить,
    # что желаемое состояние достигнуто, только что Редактор ничего не
    # поменял по факту. set_style на тот же стиль, что уже есть, — валидная
    # операция (validate не запрещает старый==новый для стиля), которая
    # применяется, но не меняет ТЕКСТ, который сравнивает _diff.
    doc = Document()
    doc.add_paragraph("Абзац без изменений текста.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "set_style", "id": "p0", "style": "Normal"}]}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — diff пуст раньше")

    result, doc, idx = run_edit(
        doc, idx, "измени абзац",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "failed", result
    assert result["reason"] == "операции применились, но текст не изменился", result
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before
    print("edit_demo: non-rule путь с пустым diff остаётся failed (already — только через rule/fallthrough, не отсюда)")


def test_quoted_phrase_always_merges_with_navigator_hit():
    # Ф8 item C: правки 4/8/14 отказывались с «во фрагменте только один
    # случай», потому что дословная цитата «ёлочками» из текста правки
    # искалась ТОЛЬКО как запасной путь — когда Навигатор не назвал ни
    # одного id/anchor'а. Навигатор чаще называет ОДИН случай (тот, что
    # увидел первым), и запасной путь никогда не срабатывал, хотя цитата
    # реально встречается в нескольких местах документа (правка 8 — три
    # разных написания года, каждое в «ёлочках», в трёх разных блоках).
    # Теперь цитата всегда доливает совпадения, а не только когда всё
    # остальное пусто.
    doc = Document()
    doc.add_paragraph("Первое: термин встречается тут.")
    doc.add_paragraph("Второе предложение без искомого слова.")
    doc.add_paragraph("Третье: термин встречается снова тут.")
    doc.add_paragraph("Четвёртое: термин встречается и здесь тоже.")
    idx = index(doc)

    seen_fragments = []

    def fake_navigator(outline_text, request):
        # Навигатор нашёл только ОДИН случай (p0) — как измерено на живом Навигаторе.
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        seen_fragments.append(fragment_text)
        return {"ops": []}

    request = "Замени везде «термин встречается» на другую формулировку."
    run_edit(doc, idx, request, navigator=fake_navigator, editor=fake_editor)

    assert seen_fragments, "Редактор обязан быть вызван"
    fragment_text = seen_fragments[0]
    assert "p0" in fragment_text, fragment_text
    assert "p2" in fragment_text, fragment_text
    assert "p3" in fragment_text, fragment_text
    print("edit_demo: дословная цитата в «ёлочках» слита с id Навигатора, все 3 случая дошли до фрагмента Редактора")


def test_checker_exception_restores_document_and_cleans_snapshot():
    # Ф8 item D: если Проверяющий падает (оборвался транспорт), патч уже
    # применён к живому doc, снимок ещё жив. Исключение обязано долететь
    # наружу ГРОМКО (инвариант 6), но doc обязан вернуться к состоянию ДО
    # правки — ТОЙ ЖЕ ссылкой, потому что вызывающий не получает новую пару
    # (result, doc, idx) через return (управление уходит через raise), а
    # снимок обязан быть удалён, а не течь во временные файлы.
    doc = Document()
    doc.add_paragraph("Оригинальный текст абзаца до правки.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "Оригинальный", "new": "Изменённый"}]}

    def fake_checker(request, diff):
        raise RuntimeError("обрыв транспорта на проверке (тест)")

    captured = {}
    real_snapshot = edit_mod._snapshot

    def spy_snapshot(d):
        path = real_snapshot(d)
        captured["path"] = path
        return path

    edit_mod._snapshot = spy_snapshot
    try:
        raised = False
        try:
            run_edit(
                doc, idx, "замени 'Оригинальный' на 'Изменённый'",
                navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
            )
        except RuntimeError:
            raised = True
    finally:
        edit_mod._snapshot = real_snapshot

    assert raised, "исключение Проверяющего обязано долететь наружу, а не проглотиться"
    assert "path" in captured, "снимок обязан был создаваться (ops успели примениться)"
    assert not os.path.exists(captured["path"]), "снимок обязан быть удалён даже при исключении, а не течь"

    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "doc (та же ссылка) обязан вернуться к состоянию ДО правки"
    print("edit_demo: исключение Проверяющего долетает наружу, doc восстановлен на месте, снимок не течёт")


def test_editor_exception_mid_batch_restores_earlier_cluster():
    # В11 (замер w18, colbert#4/math#4): обрыв транспорта у Редактора ВТОРОГО
    # кластера — уже ПОСЛЕ того, как первый кластер успешно применил set_style
    # к живому doc, — раньше улетал наружу, минуя восстановление: doc
    # оставался наполовину мутированным, а вызывающий (bench/server.py),
    # ретраящий ТУ ЖЕ правку на ТЕХ ЖЕ doc/idx, получал уже испорченный
    # документ и не знал об этом. Стиль первого кластера обязан откатиться
    # ТАК ЖЕ, как checker-исключение откатывает текст (см. тест выше).
    doc = Document()
    for i in range(12):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)
    before_styles = [b["style"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p8"], "anchors": []}

    calls = []

    def fake_editor(fragment_text, request, feedback=None):
        calls.append(fragment_text)
        if len(calls) == 1:
            return {"ops": [{"op": "set_style", "id": "p0", "style": "Heading 1"}]}
        raise ConnectionError("обрыв транспорта на втором кластере (тест)")

    def no_checker(request, diff):
        raise AssertionError("до Проверяющего дойти не должно — исключение случается раньше")

    raised = False
    try:
        run_edit(
            doc, idx, "сделай оба абзаца заголовками",
            navigator=fake_navigator, editor=fake_editor, checker=no_checker,
        )
    except ConnectionError:
        raised = True

    assert raised, "исключение Редактора обязано долететь наружу, а не проглотиться"
    assert len(calls) == 2, "второй кластер обязан был получить свой вызов Редактора"
    after_styles = [b["style"] for b in doc_map(doc, idx)]
    assert after_styles == before_styles, (
        "стиль p0, применённый первым кластером, обязан откатиться при исключении второго: "
        f"{after_styles}"
    )
    print("edit_demo: исключение Редактора между кластерами откатывает уже применённый первый кластер")


def test_set_style_diff_reaches_checker():
    # Change A: _diff раньше сравнивал только текст — set_style была ей
    # невидима, Проверяющий получал пустой diff, и run_edit откатывал правку
    # с "текст не изменился", хотя стиль реально сменился (это и произошло на
    # живой правке 4 корпуса Математика). Теперь diff обязан нести стиль, и
    # такая правка обязана доходить до Проверяющего и закрываться done.
    doc = Document()
    doc.add_paragraph("Заголовок раздела")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "set_style", "id": "p0", "style": "Heading 1"}]}

    seen_diff = []

    def fake_checker(request, diff):
        seen_diff.append(diff)
        return {"ok": True, "reason": "стиль сменился на заголовок"}

    result, doc, idx = run_edit(
        doc, idx, "сделай заголовком",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert seen_diff and seen_diff[0], "diff обязан быть непустым — Проверяющий должен увидеть изменение"
    assert "Heading 1" in seen_diff[0][0].get("note", ""), seen_diff[0]
    assert result["verdict"] == "done", result
    print(f"edit_demo: set_style дошёл до Проверяющего с note={seen_diff[0][0]['note']!r}, verdict=done")


def test_diff_move_and_insert():
    # Change A p.2: move_after обязан дать РОВНО одну запись diff'а на
    # переместившийся блок; insert_after не должен порождать НИКАКИХ записей
    # с пометкой перемещения — множество id там другое (появился новый), а
    # не просто переставленное старое.
    from docx_editor.edit import _diff

    doc = Document()
    for t in ("Первый", "Второй", "Третий", "Четвёртый"):
        doc.add_paragraph(t)
    idx = index(doc)
    before = doc_map(doc, idx)

    patch.apply(doc, idx, {"op": "move_after", "id": "p0", "after": "p2"})
    move_diff = _diff(before, doc_map(doc, idx))
    move_entries = [d for d in move_diff if "перемещ" in d.get("note", "")]
    assert len(move_entries) == 1, move_diff
    assert move_entries[0]["id"] == "p0", move_diff

    doc2 = Document()
    for t in ("Первый", "Второй", "Третий", "Четвёртый"):
        doc2.add_paragraph(t)
    idx2 = index(doc2)
    before2 = doc_map(doc2, idx2)
    patch.apply(doc2, idx2, {"op": "insert_after", "id": "p1", "text": "Вставленный", "style": "Normal"})
    insert_diff = _diff(before2, doc_map(doc2, idx2))
    spurious = [d for d in insert_diff if "перемещ" in d.get("note", "")]
    assert not spurious, insert_diff
    print("edit_demo: move_after даёт ровно одну запись в diff, insert_after не даёт ложных move-записей")


def test_move_that_cancels_itself_is_not_done():
    # Н68 (живой прогон): «перенеси пункт про окно на первое место» дала два
    # move_after на ТОТ ЖЕ блок (сначала после C, потом обратно после B) —
    # итоговая позиция блока совпала с исходной, перенос отменил сам себя, а
    # вердикт был done. patch.apply не бросает исключение просто потому, что
    # итог бессмыслен — судить обязан код по ИТОГОВОЙ позиции, не по факту
    # применения. Батч несёт ЕЩЁ и настоящую правку текста (typo), чтобы diff
    # был непустым и правка не перехватывалась более ранним гвардом «текст не
    # изменился» — проверяем именно новый, более точный, гвард переноса.
    doc = Document()
    for t in ("Пункт А.", "Пункт Б.", "Пункт про окно.", "Пункт В."):
        doc.add_paragraph(t)
    doc.add_paragraph("Тут опечтака в отдельном абзаце.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p2", "p4"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [
            {"op": "move_after", "id": "p2", "after": "p3"},  # сначала уводит "окно" в конец списка
            {"op": "move_after", "id": "p2", "after": "p1"},  # и тут же возвращает его на исходное место
            {"op": "replace_text", "id": "p4", "old": "опечтака", "new": "опечатка"},
        ]}

    def no_checker(request, diff):
        raise AssertionError("гвард переноса обязан отказать ДО вызова Проверяющего")

    result, doc, idx = run_edit(
        doc, idx, "Перенеси пункт про окно на первое место и заодно поправь опечатку.",
        navigator=fake_navigator, editor=fake_editor, checker=no_checker,
    )

    assert result["verdict"] == "failed", result
    assert "отменил сам себя" in result["reason"], result["reason"]
    texts = [b["text"] for b in doc_map(doc, idx)]
    assert texts[2] == "Пункт про окно." and "опечтака" in texts[4], (
        "документ обязан остаться нетронутым — самоотменившийся перенос откачен целиком: " + str(texts)
    )
    print(f"edit_demo: перенос, отменивший сам себя, не считается перемещением: {result['reason']!r}")


def test_set_list_level_diff_reaches_checker():
    # Ф15: set_list_level не трогает текст абзаца вовсе, только w:ilvl —
    # без сравнения b.get("list") != a.get("list") в _struct_note (уже было
    # у Ф11 для style/level/list) правка была бы невидима Проверяющему, как
    # раньше были невидимы set_style/move_after. Это и проверяем: без вызова
    # модели, прямо по _diff.
    from docx_editor.edit import _diff

    doc = Document()
    p = doc.add_paragraph("Пункт списка")
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), "0")
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), "5")
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)
    idx = index(doc)
    before = doc_map(doc, idx)

    patch.apply(doc, idx, {"op": "set_list_level", "id": "p0", "ilvl": 1})
    diff = _diff(before, doc_map(doc, idx))
    assert diff, "set_list_level обязан дать непустой diff — иначе Проверяющий его не увидит"
    assert "список" in diff[0].get("note", ""), diff
    print(f"edit_demo: set_list_level дошёл до diff'а с note={diff[0]['note']!r}")


def test_rule_fallthrough_to_editor_done():
    # Change C: rule, не найдя типографских дефектов, раньше был терминальным
    # "already", даже когда правка вообще не про типографику (Математика
    # edit 1 — двойной пробел и висящий пробел остались в документе, а rule
    # искал не тот класс дефектов и смолчал). Пустой diff у rule теперь не
    # терминал: код восстанавливает документ и продолжает ТЕМ ЖЕ запросом
    # обычным локальным путём. Здесь rule ожидаемо ничего не находит
    # (никаких типографских дефектов), а фейковый Редактор честно предлагает
    # рабочую операцию — итог обязан быть done, а не already.
    doc = Document()
    doc.add_paragraph("Обычный абзац без опечаток.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": "typography", "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "Обычный", "new": "Другой"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'Обычный' на 'Другой'",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    after_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p0")
    assert after_text == "Другой абзац без опечаток.", after_text
    print("edit_demo: rule без изменений упал в локальный путь, Редактор довёл правку до done")


def test_rule_fallthrough_editor_declines_stays_already():
    # Тот же fallthrough, но Редактор честно отказывается (пустые ops) — оба
    # пути (rule и Редактор) согласны, что менять нечего, already здесь
    # легитимен (в отличие от отказа validate/Проверяющего — item C.4).
    doc = Document()
    doc.add_paragraph("Обычный абзац без опечаток.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": "typography", "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": []}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — Редактор честно отказался")

    result, doc, idx = run_edit(
        doc, idx, "замени 'Обычный' на 'Другой'",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "already", result
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "already не должен менять документ"
    print(f"edit_demo: fallthrough с честным отказом Редактора даёт already: {result['reason']!r}")


def test_edit_12_on_real_doc_fixture():
    # Раньше это был единственный живой вызов модели во всём демо — упирался в
    # реальный API при каждом прогоне run_all.py (см. CLAUDE.md: демо,
    # закрываемое фикстурой, закрываем фикстурой). Живая приёмка ровно этого
    # же корпуса и этой же правки №12 остаётся в examples/colbert_run.py (не
    # *_demo.py, run_all.py его не подхватывает) — там весь прогон из 20 правок
    # проверяет тот же инвариант "вердикт правдив" по-настоящему, живой
    # моделью. Здесь — фейковые Навигатор/Редактор/Проверяющий в форме их
    # реального JSON-ответа (как в остальных тестах файла), но реальный
    # документ, реальный текст правки №12 (через split реального файла правок)
    # и реальный p76: инвариант — verdict=done обязан означать, что p76
    # реально поменялся, а остальные блоки документа не тронуты.
    doc = Document(REAL_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)
    p76_before = next(b["text"] for b in blocks if b["id"] == "p76")
    other_before = [b["text"] for b in blocks if b["id"] != "p76"]

    request = split(open(EDITS_FILE, encoding="utf-8").read())[11]
    old = "на порядок (в 6–10 раз)"
    new = "более чем в 6–10 раз"
    assert old in p76_before, p76_before

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p76"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        # Правка цитирует ещё «на два порядка (в ~170 раз)», и _resolve честно
        # доливает по этой цитате приблизительное совпадение где-то ещё в
        # документе (locate) — с кластеризацией это отдельный, СВОЙ вызов
        # Редактора на СВОЁМ маленьком фрагменте. Фейк обязан быть таким же
        # избирательным, как настоящий: отвечать только когда p76 реально
        # виден в ЭТОМ фрагменте, иначе честно "нечего делать здесь".
        if "p76 " not in fragment_text:
            return {"ops": []}
        return {"ops": [{"op": "replace_text", "id": "p76", "old": old, "new": new}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "кратность и формулировка теперь согласованы"}

    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    p76_after = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p76")
    assert p76_after != p76_before and old not in p76_after and new in p76_after, p76_after
    other_after = [b["text"] for b in doc_map(doc, idx) if b["id"] != "p76"]
    assert other_after == other_before, "правка №12 не должна трогать остальные блоки документа"
    print(f"edit_demo: правка №12 на реальном документе (фикстура): p76 было {p76_before!r}, стало {p76_after!r}")


def test_scattered_targets_three_calls_one_verdict():
    # Спецификация кластеризации: цели в трёх разных частях документа (индексы
    # 0/5/10 из 12 абзацев — при around=1 окна p0/p5/p10 не пересекаются)
    # обязаны дать ТРИ маленьких вызова Редактора (каждый видит ровно одну
    # цель плюс соседей), но ОДИН снимок/diff/вердикт на всю правку — цикл
    # дробит шаги, не обязательство (см. _run_clusters/_finish).
    doc = Document()
    for i in range(12):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)
    target_ids = ["p0", "p5", "p10"]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": target_ids, "anchors": []}

    calls = []

    def fake_editor(fragment_text, request, feedback=None):
        calls.append(fragment_text)
        # считаем ТОЛЬКО строки блоков: перед фрагментом код может положить
        # свои приписки (инвентарь вариантов, «уже сделано в этой правке»),
        # и цитаты в них не являются целями кластера
        block_lines = [ln for ln in fragment_text.splitlines() if re.match(r"[pt]\d+ \[", ln)]
        present = [t for t in target_ids if any(ln.startswith(f"{t} ") for ln in block_lines)]
        assert len(present) == 1, f"кластер обязан нести ровно одну цель: {fragment_text!r}"
        return {"ops": [{"op": "replace_text", "id": present[0], "old": "обычный", "new": "изменённый"}]}

    checker_calls = []

    def fake_checker(request, diff):
        checker_calls.append(diff)
        assert len(diff) == 3, diff
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'обычный' на 'изменённый' в трёх местах",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert len(calls) == 3, f"три разнесённых блока обязаны дать три вызова Редактора, получили {len(calls)}"
    assert len(checker_calls) == 1, "Проверяющий обязан вызываться РОВНО один раз на всю правку"
    assert result["verdict"] == "done", result
    assert result["iter"] == 3, result
    print("edit_demo: три разнесённые цели дали три маленьких вызова Редактора и один общий вердикт")


def test_adjacent_targets_one_call():
    # Клаcтеризация: два соседних абзаца (p0/p1, окна пересекаются) — ОДИН
    # кластер и, значит, ОДИН вызов Редактора, а не два.
    doc = Document()
    doc.add_paragraph("Абзац A: обычный.")
    doc.add_paragraph("Абзац B: обычный.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p1"], "anchors": []}

    calls = []

    def fake_editor(fragment_text, request, feedback=None):
        calls.append(fragment_text)
        return {"ops": [
            {"op": "replace_text", "id": "p0", "old": "обычный", "new": "другой"},
            {"op": "replace_text", "id": "p1", "old": "обычный", "new": "другой"},
        ]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'обычный' на 'другой' в обоих абзацах",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert len(calls) == 1, f"два соседних абзаца обязаны дать ОДИН вызов Редактора, получили {len(calls)}"
    assert result["verdict"] == "done", result
    print("edit_demo: два соседних резолвленных id слились в один кластер — один вызов Редактора")


def test_empty_ops_cluster_skipped_others_apply():
    # Кластер, где Редактор честно вернул пустые ops ("нечего делать здесь"),
    # не должен проваливать правку — другие кластеры всё ещё могут дать работу,
    # и итоговый вердикт обязан быть done.
    doc = Document()
    for i in range(8):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)

    def fake_editor(fragment_text, request, feedback=None):
        if "p0 " in fragment_text:
            return {"ops": []}
        return {"ops": [{"op": "replace_text", "id": "p7", "old": "обычный", "new": "другой"}]}

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p7"], "anchors": []}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'обычный' на 'другой' там, где уместно",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    p7_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p7")
    assert "другой" in p7_text, p7_text
    print("edit_demo: пустой ops одного кластера пропущен, другой кластер применился, вердикт done")


def test_invalid_cluster_aborts_whole_edit():
    # Кластер, чьи ops не проходят валидацию даже после ретрая внутри
    # _apply_ops, обязан прервать ВСЮ правку — включая откат уже применённого
    # предыдущего кластера. Атомарность остаётся на уровне всей правки, а не
    # кластера (см. спецификацию задачи: "шаги мельче, обязательство то же").
    doc = Document()
    doc.add_paragraph("Абзац A: обычный текст.")
    for i in range(1, 5):
        doc.add_paragraph(f"Абзац {i}: наполнитель.")
    doc.add_paragraph("Абзац B: обычный текст.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p5"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        if "p0 " in fragment_text:
            return {"ops": [{"op": "replace_text", "id": "p0", "old": "обычный", "new": "другой"}]}
        # p5-кластер: "old" которого в блоке нет — невалидно что при первой
        # попытке, что при ретрае (feedback игнорируется нарочно).
        return {"ops": [{"op": "replace_text", "id": "p5", "old": "текста, которого тут нет", "new": "х"}]}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — правка обязана прерваться раньше")

    result, doc, idx = run_edit(
        doc, idx, "замени 'обычный' на 'другой'",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "failed", result
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "невалидный кластер обязан откатить ВСЮ правку, включая уже применённый первый кластер"
    print("edit_demo: невалидный кластер после ретрая откатил всю правку, документ побайтово как был")


def test_many_scattered_targets_no_cap():
    # Владелец явно снял верхнюю границу на число целей: сколько бы
    # резолвленных id ни было, каждый непересекающийся кластер даёт свой
    # вызов Редактора — без кэпа и без отката на "один большой фрагмент".
    doc = Document()
    n_targets, total_paras = 15, 45
    for i in range(total_paras):
        doc.add_paragraph(f"Абзац {i}: наполнитель.")
    idx = index(doc)
    target_ids = [f"p{i}" for i in range(0, total_paras, 3)]
    assert len(target_ids) == n_targets

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": target_ids, "anchors": []}

    calls = []

    def fake_editor(fragment_text, request, feedback=None):
        calls.append(fragment_text)
        # цели считаем по строкам блоков: приписки кода перед фрагментом
        # (инвентарь вариантов, «уже сделано») целями кластера не являются
        block_lines = [ln for ln in fragment_text.splitlines() if re.match(r"[pt]\d+ \[", ln)]
        present = [t for t in target_ids if any(ln.startswith(f"{t} ") for ln in block_lines)]
        assert len(present) == 1, fragment_text
        return {"ops": [{"op": "replace_text", "id": present[0], "old": "наполнитель", "new": "правка"}]}

    def fake_checker(request, diff):
        assert len(diff) == n_targets, diff
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'наполнитель' на 'правка' везде, где нужно",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert len(calls) == n_targets, f"{n_targets} разнесённых целей без кэпа обязаны дать {n_targets} вызовов, получили {len(calls)}"
    assert result["verdict"] == "done", result
    assert result["iter"] == n_targets, result
    print(f"edit_demo: {n_targets} разнесённых целей — {len(calls)} вызовов Редактора, кэп на 12 отсутствует")


def test_out_of_lane_op_blocked_document_unchanged():
    # Находка: _apply_ops звал patch.validate с картой ВСЕГО документа, поэтому
    # операция на любой существующий id проходила, даже если Редактор видел
    # только свой фрагмент. Редактор, которому показали окрестность p0, но
    # который отвечает операцией на далёкий p9, обязан быть отклонён кодом, а
    # не применён — документ должен остаться побайтово как был, verdict != done.
    doc = Document()
    for i in range(10):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        assert "p9 " not in fragment_text, "p9 не должен попадать во фрагмент вокруг p0"
        return {"ops": [{"op": "replace_text", "id": "p9", "old": "обычный", "new": "чужой"}]}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — операция вне фрагмента обязана быть отклонена раньше")

    result, doc, idx = run_edit(
        doc, idx, "замени 'обычный' на 'чужой' в p0",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] != "done", result
    assert "вне фрагмента" in result["reason"], result
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "документ обязан остаться побайтово как был — блок вне фрагмента не должен меняться"
    print(f"edit_demo: операция на блок вне фрагмента отклонена, документ не тронут: {result['reason']!r}")


def test_out_of_lane_retry_recovers_within_fragment():
    # Тот же случай, но Редактор, получив ошибку об "вне фрагмента" как
    # feedback, на ретрае отвечает верно — в пределах СВОЕГО фрагмента.
    # Единственный существующий ретрай внутри _apply_ops обязан это спасти.
    doc = Document()
    for i in range(10):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        if feedback is None:
            return {"ops": [{"op": "replace_text", "id": "p9", "old": "обычный", "new": "чужой"}]}
        assert "вне фрагмента" in feedback, feedback
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "обычный", "new": "исправленный"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'обычный' на 'исправленный' в p0",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    p9_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p9")
    assert "чужой" not in p9_text, p9_text
    p0_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p0")
    assert "исправленный" in p0_text, p0_text
    print("edit_demo: после ретрая с feedback Редактор попал в свой фрагмент, verdict=done")


def test_replace_all_not_blocked_by_lane_guard():
    # replace_all документ-широкая ПО КОНТРАКТУ и id не несёт — гвард обязан
    # её пропускать, хотя формально она меняет блоки вне показанного фрагмента.
    doc = Document()
    for i in range(10):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "replace_all", "old": "обычный", "new": "новый"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'обычный' на 'новый' везде",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    p9_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p9")
    assert "новый" in p9_text, p9_text
    print("edit_demo: replace_all не заблокирован гвардом фрагмента, применился по всему документу")


def test_op_on_block_created_in_same_batch_applies():
    # insert_after создаёт свежий id — операция на него в ТОМ ЖЕ батче обязана
    # быть разрешена, хотя этого id не было во фрагменте, который видел
    # Редактор изначально (см. пополнение fragment_ids в _apply_ops).
    doc = Document()
    doc.add_paragraph("Первый абзац.")
    doc.add_paragraph("Второй абзац.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [
            {"op": "insert_after", "id": "p0", "text": "Новый абзац.", "style": "Normal"},
            {"op": "set_style", "id": "p2", "style": "Heading 1"},
        ]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "вставь абзац после первого и сделай его заголовком",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    new_block = next(b for b in doc_map(doc, idx) if b["id"] == "p2")
    assert new_block["text"] == "Новый абзац.", new_block
    assert new_block["style"] == "Heading 1", new_block
    print("edit_demo: операция на блок, созданный этим же батчем, применилась, verdict=done")


def test_two_insert_after_same_anchor_keep_order():
    # Н67 (замер, живой документ): «добавь раздел «6. Пересмотр» с текстом» —
    # Редактор дал ДВЕ операции insert_after с ОДНИМ и тем же якорем p16
    # (заголовок раздела, затем текст раздела). Раньше вторая вставка ложилась
    # ВПЛОТНУЮ к якорю и вытесняла первую вниз — в файле получался текст
    # раздела, а заголовок ПОСЛЕ него (перевёрнутый порядок). Второй
    # insert_after на тот же якорь обязан цепляться за ПЕРВУЮ вставку, а не
    # снова за якорь.
    doc = Document()
    doc.add_paragraph("5. Некоторый предыдущий раздел.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": [], "trace": "heading+text"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [
            {"op": "insert_after", "id": "p0", "text": "6. Пересмотр", "style": "Heading 1"},
            {"op": "insert_after", "id": "p0", "text": "Регламент пересматривается ежегодно.", "style": "Normal"},
        ]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "раздел добавлен"}

    result, doc, idx = run_edit(
        doc, idx, "Добавь раздел «6. Пересмотр» с текстом: Регламент пересматривается ежегодно.",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    texts = [b["text"] for b in doc_map(doc, idx)]
    assert texts == [
        "5. Некоторый предыдущий раздел.",
        "6. Пересмотр",
        "Регламент пересматривается ежегодно.",
    ], texts
    print("edit_demo: два insert_after с одним якорем сохраняют порядок операций (заголовок, затем текст)")


def test_render_full_text_outline_stays_short():
    # Ф12: 50 из 116 абзацев ColBERT-документа длиннее 300 знаков — render()
    # обязан отдавать их ЦЕЛИКОМ (иначе Редактор режет цитату на границе
    # обрыва и рвёт слово пополам), а find.outline() по-прежнему режет
    # коротким префиксом — это другой потребитель (покрывает ВЕСЬ документ
    # для Навигатора) с другим бюджетом, его не трогаем.
    doc = Document(REAL_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)
    long_blocks = [b for b in blocks if b["kind"] == "p" and len(b["text"]) > 300 and b["level"] is None]
    assert len(long_blocks) >= 30, f"ожидали много длинных абзацев без level, получили {len(long_blocks)}"
    b = long_blocks[0]

    rendered = render([b])
    assert b["text"] in rendered, "render() обязан отдавать длинный абзац целиком, без обрыва"

    outline_line = find.outline([b])
    assert b["text"][:20] in outline_line and "…" in outline_line, outline_line
    assert len(outline_line) < len(rendered), "outline() обязан оставаться короче полного render()"
    print(f"edit_demo: render() отдал абзац {b['id']} ({len(b['text'])} зн.) целиком, "
          f"outline() по-прежнему режет коротким префиксом ({len(outline_line)} зн.)")


def test_defect1_mid_word_cut_rejected():
    # Ф12: даже без 300-знакового обрыва в render() модель МОЖЕТ обрубить
    # цитату сама (см. отчёт строителя) — гвард в patch.validate обязан
    # ловить это независимо от источника обрыва, через единственный ретрай
    # _apply_ops, и не дать словам склеиться.
    doc = Document()
    doc.add_paragraph("Для больших чисел используют специальные методы, "
                       "которые требуют серьёзных вычислительных ресурсов.")
    idx = index(doc)
    before = doc_map(doc, idx)[0]["text"]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        if feedback is None:
            # "old" обрублен ровно посреди слова «серьёзных» — так же, как
            # раньше обрубал старый _truncate на границе в 300 знаков
            return {"ops": [{"op": "replace_text", "id": "p0",
                              "old": "которые требуют с", "new": "которые требуют больших мощностей, х"}]}
        assert "слова" in feedback, feedback
        return {"ops": [], "note": "не могу процитировать целым словом"}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — обрубленная цитата обязана быть отбита раньше")

    result, doc, idx = run_edit(
        doc, idx, "замени методы на более мощные",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] != "done", result
    after = doc_map(doc, idx)[0]["text"]
    assert after == before, after
    assert "херьёзных" not in after, after
    print("edit_demo: обрубленная посреди слова цитата отбита, слово не склеилось")


def test_defect2_batch_collision_colbert11():
    # Реальный случай ColBERT edit 11 (см. отчёт строителя), один батч:
    # 1) «Кросс-энкодер» → «Кросс-энкодер (Cross-Encoder, англ.)»
    # 2) «Cross-Encoder» → «Кросс-энкодер»  ← матчит ВНУТРИ вставки шага 1
    # без гварда даёт «Кросс-энкодер, англ.)» → склейку вида «...энкодерs».
    doc = Document()
    doc.add_paragraph("Кросс-энкодер — важная архитектура поиска.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        if feedback is None:
            return {"ops": [
                {"op": "replace_text", "id": "p0", "old": "Кросс-энкодер",
                 "new": "Кросс-энкодер (или Cross-Encoder, англ.)"},
                {"op": "replace_text", "id": "p0", "old": "Cross-Encoder", "new": "Кросс-энкодер"},
            ]}
        assert "батча" in feedback, feedback
        return {"ops": [], "note": "op2 бьёт по собственному выводу op1 — оставляю как есть"}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — батч обязан провалиться раньше")

    result, doc, idx = run_edit(
        doc, idx, "добавь англоязычный термин рядом с русским",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    text = doc_map(doc, idx)[0]["text"]
    assert "Кросс-энкодерs" not in text and "энкодерыs" not in text, text
    print(f"edit_demo: op2 не смог откусить от вывода op1 в том же батче, verdict={result['verdict']!r}")


def test_defect2_batch_collision_replace_all():
    # Реальный случай ColBERT edit 7: op1 этого же батча вписывает «Bi-encoders»,
    # а op2 (replace_all «Bi-encoder»→«Bi-encoders») находит «Bi-encoder» как
    # префикс СВЕЖЕГО вывода op1 и без гварда даёт «Bi-encoderss».
    doc = Document()
    doc.add_paragraph("Bi-encoder — базовая архитектура.")
    doc.add_paragraph("Второй абзац без термина.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        if feedback is None:
            return {"ops": [
                {"op": "replace_text", "id": "p1", "old": "Второй абзац",
                 "new": "Второй абзац, известный как Bi-encoders"},
                {"op": "replace_all", "old": "Bi-encoder", "new": "Bi-encoders"},
            ]}
        # Ф16, item 4: validate теперь считается ДО коллизии батча (targets/
        # collision требуют полей, которые validate и проверяет) — здесь
        # replace_all матчит «Bi-encoder» как префикс уже вписанного op1
        # «Bi-encoders», и это одновременно и самоколлизия батча, и обрыв
        # посреди слова; validate.py ловит обрыв первым, тем же исходом
        # (батч проваливается, документ не портится).
        assert "посреди слова" in feedback, feedback
        return {"ops": [], "note": "replace_all бьёт по только что вписанному Bi-encoders"}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — батч обязан провалиться раньше")

    result, doc, idx = run_edit(
        doc, idx, "приведи термин к множественному числу",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    texts = " ".join(b["text"] for b in doc_map(doc, idx))
    assert "Bi-encoderss" not in texts, texts
    print(f"edit_demo: replace_all не тронул Bi-encoders, вписанный op1 этого же батча, verdict={result['verdict']!r}")


def test_batch_identity_op_skipped_not_retried():
    # ПАРТИЯ 1 (BUILD_PLAN): воспроизводит форму ColBERT 7 (bench/runs/w8,
    # n=7) — два реальных replace_all и третий, тождественный (old==new),
    # который Редактор добавляет, "унифицируя" правку к варианту, уже
    # стоящему в документе. До фикса: op3 бьёт "old и new совпадают",
    # единственный ретрай тратится впустую, батч проваливается — ОБЕ верные
    # замены откатываются вместе с ним. После фикса: op3 молча пропускается,
    # обе замены остаются в силе.
    doc = Document()
    doc.add_paragraph("Используется Single-Vector Bi-encoders. Также известен как Bi-Encoders для сравнения.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [
            {"op": "replace_all", "old": "Single-Vector Bi-encoders", "new": "Single-Vector Bi-encoder"},
            {"op": "replace_all", "old": "Bi-Encoders", "new": "Bi-encoder"},
            {"op": "replace_all", "old": "Bi-encoder", "new": "Bi-encoder"},
        ]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "Bi-encoder пишется по-разному, выбери один вариант и поставь везде",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    text = doc_map(doc, idx)[0]["text"]
    assert "Single-Vector Bi-encoder." in text, text
    assert "Bi-encoder для сравнения" in text, text
    assert "Bi-Encoders" not in text and "Single-Vector Bi-encoders" not in text, text
    print(f"edit_demo: тождественная операция (old==new) в батче пропущена без ретрая, verdict={result['verdict']!r}")


def test_batch_consumed_old_skipped_not_retried():
    # ПАРТИЯ 1 (BUILD_PLAN): воспроизводит Математика 14 (bench/runs/w8,
    # n=14) — op1 (replace_all «читатель»→«вы») съедает подстроку, которую
    # op2 того же батча искал целой фразой. До фикса: op2 бьёт "не найден",
    # единственный ретрай не спасает, батч (и уже верная работа op1)
    # откатывается. После фикса: op2 распознаётся как объяснимый собственной
    # перепиской батча ("не найден", потому что уже переписан, а не потому
    # что документ не такой) и пропускается, работа op1 остаётся в силе.
    doc = Document()
    doc.add_paragraph("Я хочу, чтобы читатель знал важные детали.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [
            {"op": "replace_all", "old": "читатель", "new": "вы"},
            {"op": "replace_all", "old": "Я хочу, чтобы читатель знал", "new": "Я хочу, чтобы вы знали"},
        ]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "приведи обращение к читателю к одному лицу по всей главе",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    text = doc_map(doc, idx)[0]["text"]
    assert "вы знал" in text, text
    assert "читатель" not in text, text
    assert any("пропущ" in a for a in result["applied"]), result["applied"]
    print(f"edit_demo: операция, чей old съеден более ранней записью батча, честно помечена пропущенной, verdict={result['verdict']!r}")


def test_new_containing_old_still_applies():
    # "new", включающий "old" целиком (добавили пояснение в скобках), — сам
    # по себе легитимная правка: под подозрение попадает только ПОЗДНЕЙШАЯ
    # операция, которая матчит этот вывод (см. тесты выше), а не эта.
    doc = Document()
    doc.add_paragraph("Кросс-энкодер — важная архитектура поиска.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "replace_text", "id": "p0",
                          "old": "Кросс-энкодер", "new": "Кросс-энкодер (Cross-Encoder)"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "добавь англоязычный термин в скобках",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )
    assert result["verdict"] == "done", result
    print("edit_demo: new, включающий old целиком, применился без ложного срабатывания гварда")


def test_variant_inventory_appears_for_multiple_targets():
    # Change 1 (Ф13): при дроблении на кластеры Редактор кластера видит
    # только тот вариант, что уже в нём, и "унифицирует" правку к нему же
    # (ColBERT 7/15). Инвентарь вариантов с частотами обязан попасть в
    # fragment_text КАЖДОГО кластера, когда в документе больше одного
    # варианта с реальными попаданиями.
    doc = Document()
    doc.add_paragraph("Bi-encoder — первая архитектура.")
    doc.add_paragraph("Абзац без термина.")
    doc.add_paragraph("Bi-encoder встречается снова здесь.")
    doc.add_paragraph("Абзац без термина.")
    doc.add_paragraph("А тут пишут Bi-Encoders с большой буквы.")
    idx = index(doc)

    seen = []

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        seen.append(fragment_text)
        return {"ops": []}

    request = "Термин пишется то как «Bi-encoder», то как «Bi-Encoders» — выбери один вариант и поставь везде."
    run_edit(doc, idx, request, navigator=fake_navigator, editor=fake_editor)

    assert seen, "Редактор обязан быть вызван"
    assert "Найдено в документе:" in seen[0], seen[0]
    assert "«Bi-encoder» — 2 блока" in seen[0], seen[0]
    assert "«Bi-Encoders» — 1 блок" in seen[0], seen[0]
    print("edit_demo: инвентарь вариантов с частотами дошёл до фрагмента Редактора")


def test_variant_inventory_absent_for_single_target():
    # Тот же механизм, но с ОДНИМ вариантом-целью — секция обязана
    # отсутствовать, иначе это шум в обычной точечной правке.
    doc = Document()
    doc.add_paragraph("Bi-encoder — единственное упоминание.")
    idx = index(doc)

    seen = []

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        seen.append(fragment_text)
        return {"ops": []}

    request = "Замени «Bi-encoder» на более подробное объяснение."
    run_edit(doc, idx, request, navigator=fake_navigator, editor=fake_editor)

    assert seen, "Редактор обязан быть вызван"
    assert "Найдено в документе:" not in seen[0], seen[0]
    print("edit_demo: инвентарь не появляется для одноцелевой правки")


def test_collision_feedback_lets_retry_succeed():
    # Change 2: гвард самоколлизии (Ф12, дефект 2) обязан не просто отказывать,
    # а подсказывать выход — текст ошибки называет, ЧТО написала более ранняя
    # операция батча. Редактор, который прочитал совет и на ретрае процитировал
    # ДРУГОЙ, нетронутый кусок блока вместо повторного поиска только что
    # написанного текста, обязан пройти единственный ретрай и получить done.
    doc = Document()
    doc.add_paragraph("Кросс-энкодер — важная архитектура поиска. Также известен как CE.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        if feedback is None:
            return {"ops": [
                {"op": "replace_text", "id": "p0", "old": "Кросс-энкодер",
                 "new": "Кросс-энкодер (или Cross-Encoder, англ.)"},
                {"op": "replace_text", "id": "p0", "old": "Cross-Encoder", "new": "CE"},
            ]}
        assert "более ранняя операция" in feedback and "батча" in feedback, feedback
        # Совет учтён: цитируем ДРУГОЙ, нетронутый кусок блока, а не текст,
        # который только что записала op1.
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "как CE.", "new": "как Cross-Encoder."}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "добавь англоязычный термин и распиши сокращение",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    text = doc_map(doc, idx)[0]["text"]
    assert text == ("Кросс-энкодер (или Cross-Encoder, англ.) — важная архитектура поиска. "
                     "Также известен как Cross-Encoder."), text
    print(f"edit_demo: после совета в тексте ошибки самоколлизии ретрай прошёл, verdict={result['verdict']!r}")


def test_set_text_over_earlier_write_in_same_batch():
    # Регресс m11 замера w19 (в w18 та же правка была done): Редактор в одном
    # батче сначала правит термин в p12, потом переписывает ВЕСЬ абзац через
    # set_text. Цель set_text — весь блок, значит она пересекается с любой
    # более ранней записью батча, и гвард самоколлизии отказывал ей всегда,
    # печатая «текст «None»» (old у set_text нет). Переписывание блока целиком
    # не читает собственный вывод — оно его выбрасывает, батч обязан пройти.
    doc = Document()
    doc.add_paragraph("Тегмарк называет это «математическим универсализмом» и на этом строит вывод.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        assert feedback is None, f"ретрая быть не должно: {feedback}"
        return {"ops": [
            {"op": "replace_text", "id": "p0", "old": "«математическим универсализмом»",
             "new": "гипотезой математической Вселенной"},
            {"op": "set_text", "id": "p0",
             "text": "Тегмарк выдвигает гипотезу математической Вселенной: физическая реальность и есть математическая структура."},
        ]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "поправь название гипотезы Тегмарка",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    text = doc_map(doc, idx)[0]["text"]
    assert text.startswith("Тегмарк выдвигает гипотезу математической Вселенной"), text
    assert "универсализм" not in text, text
    print(f"edit_demo: set_text поверх более ранней записи батча прошёл, verdict={result['verdict']!r}")


def test_partial_cluster_processing_not_done():
    # Change 3 (Ф13): набор кластеров зафиксирован ДО мутации (p0 и p7 —
    # два непересекающихся кластера). Кластер p0 честно применяет операцию,
    # но она не даёт diff (set_style на тот же стиль) — цель не обработана
    # по факту, хотя кластер p7 реально поменял текст. Итог не должен быть
    # done, и документ обязан остаться побайтово как был.
    doc = Document()
    for i in range(10):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p7"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        if "p0 " in fragment_text:
            return {"ops": [{"op": "set_style", "id": "p0", "style": "Normal"}]}
        return {"ops": [{"op": "replace_text", "id": "p7", "old": "обычный", "new": "другой"}]}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — частичная обработка решается кодом")

    result, doc, idx = run_edit(
        doc, idx, "поправь два места",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] != "done", result
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "документ обязан остаться побайтово как был при частичной обработке кластеров"
    print(f"edit_demo: частично обработанный набор кластеров не даёт done, verdict={result['verdict']!r}")


def test_retry_recomputes_fragment_after_batch_mutation():
    # Ф13-бис (ColBERT 19, живое воспроизведение): батч удаляет p2, p3, p4,
    # затем insert_after падает на p3, которого сам же батч и удалил. Ретрай
    # обязан увидеть ТЕКУЩЕЕ состояние документа (p2/p3/p4 уже нет), а не
    # снимок, снятый до батча, — иначе Редактор честно видит их живыми и
    # предлагает удалить снова, и это бьёт "блок не найден в документе".
    doc = Document()
    for i in range(7):
        doc.add_paragraph(f"Абзац {i}: наполнитель.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p1", "p2", "p3", "p4", "p5"], "anchors": []}

    retry_fragments = []

    def fake_editor(fragment_text, request, feedback=None):
        if feedback is None:
            return {"ops": [
                {"op": "delete", "id": "p2"},
                {"op": "delete", "id": "p3"},
                {"op": "delete", "id": "p4"},
                {"op": "insert_after", "id": "p3", "text": "Новый абзац.", "style": "Normal"},
            ]}
        retry_fragments.append(fragment_text)
        # Улика бага: если бы фрагмент был снят ДО батча, p2/p3/p4 всё ещё
        # выглядели бы живыми здесь.
        assert "p2 [" not in fragment_text and "p3 [" not in fragment_text and "p4 [" not in fragment_text, fragment_text
        assert "p1 [" in fragment_text and "p5 [" in fragment_text, fragment_text
        return {"ops": [{"op": "insert_after", "id": "p5", "text": "Новый абзац.", "style": "Normal"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "удали три абзаца из середины и добавь новый после последнего оставшегося",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert len(retry_fragments) == 1, "ретрай обязан случиться ровно один раз"
    assert result["verdict"] == "done", result
    texts = [b["text"] for b in doc_map(doc, idx)]
    assert not any("Абзац 2" in t or "Абзац 3" in t or "Абзац 4" in t for t in texts), texts
    assert any("Новый абзац" in t for t in texts), texts
    print("edit_demo: ретрай после мутации батча получил ТЕКУЩЕЕ состояние документа, а не снимок до батча")


def test_failed_records_ops_applied_before_batch_failure():
    # Ф13-бис (ColBERT 19): даже когда правка в итоге failed (ретрай тоже не
    # спас — Редактор игнорирует feedback и повторяет невалидный ops), три
    # реальных delete, применённые ДО сбоя, обязаны остаться в "applied", а
    # не пропадать бесследно — документ при этом всё равно откатывается.
    doc = Document()
    for i in range(7):
        doc.add_paragraph(f"Абзац {i}: наполнитель.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p1", "p2", "p3", "p4", "p5"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [
            {"op": "delete", "id": "p2"},
            {"op": "delete", "id": "p3"},
            {"op": "delete", "id": "p4"},
            # невалидна что при первой попытке, что при ретрае (feedback игнорируется нарочно)
            {"op": "replace_text", "id": "p1", "old": "текста, которого тут нет", "new": "х"},
        ]}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться — правка обязана прерваться раньше")

    result, doc, idx = run_edit(
        doc, idx, "удали три абзаца из середины",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "failed", result
    assert len(result["applied"]) == 3, result["applied"]
    assert all("удал" in a for a in result["applied"]), result["applied"]
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "документ обязан откатиться, несмотря на непустой applied"
    print(f"edit_demo: failed сохранил {len(result['applied'])} операций, реально применённых до сбоя, документ откачен")


def test_id_fields_covers_all_ops_except_document_wide():
    # Ф16: _ID_FIELDS — контракт, на который опирается гвард полосы фрагмента
    # (_out_of_lane). Операция, забытая в этом словаре, проходит гвард молча
    # и может задеть любой блок документа, а не только показанный Редактору —
    # так и было с "footnote" (добавлена в patch.py, забыта здесь). Полнота
    # проверялась только против _HANDLERS (patch_demo.py), не против гварда.
    #
    # В2-бис: set_header_footer — третье законное исключение. Она не называет
    # ни один id ТЕЛА документа (колонтитул не часть doc_map()) — тот же класс,
    # что normalize/replace_all, документ-широкая операция без адресации блоком.
    missing = patch._OPS - {"normalize", "replace_all", "set_header_footer"} - set(edit_mod._ID_FIELDS)
    assert not missing, f"_ID_FIELDS не покрывает операции {missing} — гвард полосы их не проверяет"
    print("edit_demo: _ID_FIELDS покрывает все id-несущие операции patch._OPS")


def test_edit_prompt_mentions_every_op():
    # Ф17: тот же класс дефекта, что и test_id_fields_covers_all_ops_except_
    # document_wide, но для другого потребителя контракта — _EDIT_PROMPT.
    # normalize идёт путём rule и Редактору вообще не показывается (см.
    # run_edit), поэтому единственное законное исключение. Операция, которую
    # добавили в patch._OPS, но забыли упомянуть в _EDIT_PROMPT, реализована,
    # но Редактор о ней никогда не узнает и не сможет её предложить — именно
    # это уже случилось с "footnote".
    missing = [op for op in patch._OPS - {"normalize"} if op not in edit_mod._EDIT_PROMPT]
    assert not missing, f"_EDIT_PROMPT не упоминает операции {missing} — Редактор их никогда не увидит"
    print("edit_demo: _EDIT_PROMPT упоминает каждую операцию patch._OPS (кроме normalize)")


def test_edit_prompt_frames_insert_col_header_row():
    # В11 (находка Н70, замер w18): «добавь колонку „Телефон“» дала колонку
    # из пяти пустых ячеек, включая шапку, — операция cells принимает, модель
    # их не заполнила. Рамка промпта (не механизм): первая строка таблицы —
    # обычно шапка, cells[0] обязана нести осмысленный заголовок колонки.
    insert_col_line = next(ln for ln in edit_mod._EDIT_PROMPT.splitlines() if ln.startswith('{"op":"insert_col"'))
    assert "шапка" in insert_col_line and "не оставляй" in insert_col_line, insert_col_line
    print("edit_demo: _EDIT_PROMPT указывает Редактору заполнять заголовок колонки (insert_col), не оставлять пустым")


def test_replace_all_deduped_across_clusters():
    # Item 1 (di-base math#1, живой прогон: verdict rolled_back, iter=3,
    # applied нёс одну и ту же пару дважды). Каждый кластер получает ПОЛНЫЙ
    # текст правки — один и тот же replace_all может прийти от нескольких
    # кластеров подряд, а written внутри _apply_ops заводится заново на
    # каждый вызов и не видит соседний кластер. Повтор не меняет текст
    # (уже применён первым кластером), diff второго кластера пуст →
    # partial=True → откатывается ВЕСЬ верный результат первого кластера.
    doc = Document()
    doc.add_paragraph("Термин обладает свойством адекватности. Точка.")
    for i in range(1, 9):
        doc.add_paragraph(f"Абзац {i}: наполнитель.")
    doc.add_paragraph("Последний абзац без термина.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p9"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        # Оба кластера (p0 и p9 — непересекающиеся окна) честно видят один и
        # тот же полный текст правки и оба предлагают одну и ту же замену.
        return {"ops": [{"op": "replace_all", "old": "адекватности. ", "new": "адекватности."}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "убери лишний пробел после «адекватности.»",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    p0_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p0")
    assert p0_text == "Термин обладает свойством адекватности.Точка.", p0_text
    print(f"edit_demo: повторный replace_all от второго кластера дедуплицирован, verdict={result['verdict']!r}")


def test_retry_widens_lane_to_match_rerendered_fragment():
    # Item 2 (дефект в сегодняшнем ретрай-фиксе, edit.py): фрагмент на ретрае
    # перерисовывается С ОКРУЖЕНИЕМ (around=1) от УЖЕ показанных id, поэтому
    # реально видимый фрагмент шире fragment_ids. Живой случай: показаны
    # p52,p53,p54, ретрай перерисовал p51..p55, а правка на p51 отбита как
    # "вне фрагмента, показанного в этом вызове" — хотя p51 только что был
    # на экране у Редактора. Здесь: единственная цель p3, окно {p2,p3,p4},
    # первая попытка невалидна (текста нет) → ретрай перерисовывает с
    # around=1 и заодно показывает p1 — Редактор отвечает на p1, и это не
    # должно быть отбито гвардом.
    doc = Document()
    for i in range(10):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p3"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        if feedback is None:
            return {"ops": [{"op": "replace_text", "id": "p3", "old": "текста, которого тут нет", "new": "х"}]}
        assert "p1 [" in fragment_text, fragment_text
        return {"ops": [{"op": "replace_text", "id": "p1", "old": "обычный", "new": "другой"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "правка",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    p1_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p1")
    assert "другой" in p1_text, p1_text
    print("edit_demo: ретрай на блок, появившийся только в перерисованном (around=1) фрагменте, не отбит гвардом")


def test_cluster_with_removed_targets_skipped_not_dead_end():
    # Item 3 (находка контрольного анализа потока, без живого воспроизведения):
    # clusters вычисляются один раз до всех мутаций. Если к моменту обработки
    # кластера его цели уже пропали из документа, find.fragment() вернёт [] —
    # пустой фрагмент даёт ПУСТОЕ множество fragment_ids (не None!), гвард
    # полосы активен и отклонит ЛЮБОЙ ответ Редактора; ретрай перерисует тот
    # же пустой фрагмент, вторая неудача фатальна и откатывает ВСЮ правку,
    # включая уже верную работу более раннего кластера. Здесь find.fragment
    # подменена, чтобы для p6 детерминированно вернуть [] (имитация "цель
    # уже пропала"), не полагаясь на конкретный способ, которым это могло
    # случиться в реальном документе.
    doc = Document()
    for i in range(8):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)
    before_p6 = doc_map(doc, idx)[6]["text"]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p6"], "anchors": []}

    real_fragment = find.fragment

    def fake_fragment(blocks, ids, around=1):
        if set(ids) == {"p6"}:
            return []
        return real_fragment(blocks, ids, around=around)

    calls = []

    def fake_editor(fragment_text, request, feedback=None):
        calls.append(fragment_text)
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "обычный", "new": "другой"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    find.fragment = fake_fragment
    try:
        result, doc, idx = run_edit(
            doc, idx, "правка",
            navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
        )
    finally:
        find.fragment = real_fragment

    assert len(calls) == 1, f"кластер с пропавшими целями обязан быть пропущен без вызова редактора, вызовов: {len(calls)}"
    assert result["verdict"] == "done", result
    p0_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p0")
    assert "другой" in p0_text, p0_text
    p6_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p6")
    assert p6_text == before_p6, p6_text
    print("edit_demo: кластер с find.fragment()==[] пропущен, остальная правка дошла до done")


def test_targets_computed_after_validate_not_before():
    # Item 4: _op_targets вычислялся ДО patch.validate и падал TypeError на
    # операциях, которые validate и создан отклонять (например old=None у
    # replace_text/replace_all — воспроизведено на {'op':'replace_all',
    # 'new':'x'} и {'op':'replace_text','id':'p10','new':'x'}). Одна такая
    # операция от модели убивала весь батч необработанным исключением.
    # patch.validate подменён (детерминированный отказ), чтобы проверка не
    # зависела от того, есть ли уже в patch.py (чужой файл в этой сессии)
    # собственный гвард на old=None — здесь проверяется порядок вызовов
    # внутри edit.py: реальный _op_targets НЕ должен вызываться раньше
    # validate, иначе он упадёт тем же TypeError на этом же old=None.
    doc = Document()
    doc.add_paragraph("Абзац с текстом.")
    idx = index(doc)

    real_validate = patch.validate

    def fake_validate(blocks, op, d):
        return "симулированный отказ валидатора"

    patch.validate = fake_validate
    try:
        malformed = {"op": "replace_text", "id": "p0", "new": "x"}  # без "old"
        applied, err, tries = edit_mod._apply_ops(doc, idx, [malformed], "", "правка", None)
    finally:
        patch.validate = real_validate

    assert applied == [] and err == "симулированный отказ валидатора", (applied, err)
    print("edit_demo: _op_targets не вызывается раньше patch.validate — malformed op не роняет TypeError")


def test_numeric_literals_are_anchors():
    """ПАРТИЯ 2: числа с дробной частью — такие же якоря, как цитаты.

    ColBERT 4 просит «Сделай через запятую: 4.2%, 6.7%, 36.2%, 0.4448,
    0.4436, 8.8 млн», Навигатор вернул пустые ids и anchors, кавычек в
    правке нет — искать было нечем, хотя все шесть строк есть в документе.
    """
    task = ("Проценты и числа записаны через точку, как в английском. "
            "Сделай через запятую: 4.2%, 6.7%, 36.2%, 0.4448, 0.4436, 8.8 млн.")
    lits = edit_mod._literals(task)
    assert lits == ["4.2%", "6.7%", "36.2%", "0.4448", "0.4436", "8.8"], lits

    # шум без дробной части не берём, цитаты продолжают работать
    assert edit_mod._literals("«Bi-encoder» и 20 правок") == ["Bi-encoder"]

    doc = Document(REAL_DOC)
    blocks = doc_map(doc, index(doc))
    for lit in lits:
        assert find.by_text(blocks, lit), f"{lit} не найден в документе"
    print("edit_demo: числа из текста правки работают якорями наравне с цитатами")


def test_editor_already_gives_already_not_failed():
    """ПАРТИЯ 4: «здесь уже всё так» — правильный исход, а не отказ.

    ColBERT 15 просит переставить уровни списка, которые в документе УЖЕ
    расставлены верно, и Редактор честно это сообщает. Раньше вердикт «уже
    так» жил только на пути rule, и верный ответ модели записывался как
    failed. Мерило проекта — совпадение вердикта с файлом, а не число
    выполненных правок, поэтому такая метка была ложью в свою сторону.
    """
    doc = Document()
    for i in range(10):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p7"], "anchors": []}

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не зовётся: менять нечего, diff пуст")

    def editor_already(fragment_text, request, feedback=None):
        return {"ops": [], "already": True, "note": "уже в требуемом виде"}

    result, doc, idx = run_edit(
        doc, idx, "приведи к нужному виду",
        navigator=fake_navigator, editor=editor_already, checker=fake_checker,
    )
    assert result["verdict"] == "already", result
    assert [b["text"] for b in doc_map(doc, idx)] == before

    # «не умею» остаётся отказом — путать эти два ответа нельзя
    def editor_cannot(fragment_text, request, feedback=None):
        return {"ops": [], "already": False, "note": "нет подходящей операции"}

    result, doc, idx = run_edit(
        doc, idx, "сделай сноску",
        navigator=fake_navigator, editor=editor_cannot, checker=fake_checker,
    )
    assert result["verdict"] == "failed", result
    print("edit_demo: already на локальном пути отличается от «не умею»")


def test_later_clusters_see_what_earlier_ones_did():
    """ПАРТИЯ 5: кластер видит, что уже сделали предыдущие.

    Замерено: ColBERT 14 расшифровала MRR@10 дважды в двух слепых друг к
    другу кластерах (Проверяющий завернул всю правку), ColBERT 4 — третий
    кластер предложил вернуть точку вместо запятой в числе, которое первый
    уже исправил. Текст правки каждый кластер получает ЦЕЛИКОМ, поэтому без
    приписки он не может знать, что часть работы сделана.
    """
    doc = Document()
    for i in range(12):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)
    seen = []

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0", "p6"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        seen.append(fragment_text)
        target = "p0" if any(ln.startswith("p0 ") for ln in fragment_text.splitlines()) else "p6"
        return {"ops": [{"op": "replace_text", "id": target, "old": "обычный", "new": "изменённый"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "замени 'обычный' на 'изменённый' в двух местах",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )
    assert result["verdict"] == "done", result
    assert len(seen) == 2, seen
    assert "Уже сделано" not in seen[0], "первому кластеру приписывать нечего"
    assert "Уже сделано в этой правке" in seen[1], seen[1]
    assert "В p0 заменено" in seen[1], seen[1]
    print("edit_demo: второй кластер получает список уже сделанного в этой правке")


def test_unify_replaces_every_variant_and_reports_done():
    # Ф18: Навигатор маршрутизирует правку "приведи к одному" в unify —
    # модель только выбирает пары ("старое"→канон.), код заменяет ПО ВСЕМУ
    # документу и сам считает вердикт (нулевой остаток → done). Раньше это
    # шло кластерами и на 4 написаниях Bi-encoder дало ЛОЖНЫЙ done (ColBERT 7).
    doc = Document()
    doc.add_paragraph("Модель COLBERT показывает высокое качество.")
    doc.add_paragraph("ColBERT — базовая архитектура.")
    doc.add_paragraph("Далее Colbert используется как основа.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": []}

    def fake_unifier(inventory_text, request):
        assert "COLBERT" in inventory_text and "Colbert" in inventory_text, inventory_text
        return {"pairs": [["COLBERT", "ColBERT"], ["Colbert", "ColBERT"]], "note": ""}

    def no_editor(*a, **kw):
        raise AssertionError("Редактор не должен звонить на пути unify")

    def no_checker(*a, **kw):
        raise AssertionError("Проверяющий-LLM не должен звонить на пути unify — вердикт считает код")

    request = 'Термин записан как «COLBERT», «ColBERT», «Colbert» — приведи к одному варианту по всему документу.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=no_editor, checker=no_checker, unifier=fake_unifier,
    )
    assert result["verdict"] == "done", result
    assert result["reason"] == "заменено 2 вхождений 2 вариантов", result["reason"]
    text = " ".join(b["text"] for b in doc_map(doc, idx))
    assert "COLBERT" not in text and "Colbert" not in text, text
    assert text.count("ColBERT") == 3, text
    print("edit_demo: unify заменяет все варианты и даёт done по нулевому остатку")


def test_unify_leftover_variant_rolls_back_to_failed():
    # Абзац несёт ОБА написания: "Bi-encoder" внутри "Bi-encoders" (не по
    # границе слова) первым в тексте и отдельно стоящее "Bi-encoder" вторым.
    # validate() смотрит только ПЕРВОЕ совпадение _flex_span в блоке — оно
    # приклеено к "s" и не считается (word=True), поэтому блок помечается как
    # "not found" и пара не применяется вовсе, хотя отдельно стоящее написание
    # реально есть. Итоговый словограничный скан (find.by_regex) это ловит:
    # правка обязана честно откатиться, а не соврать done с недоделанным полем.
    doc = Document()
    doc.add_paragraph("Bi-encoders and Bi-encoder both appear here.")
    doc.add_paragraph("Здесь пишут Bi-Encoder с большой буквы.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": []}

    def fake_unifier(inventory_text, request):
        return {"pairs": [["Bi-encoder", "Bi-Encoder"]], "note": ""}

    def no_editor(*a, **kw):
        raise AssertionError("Редактор не должен звонить на пути unify")

    def no_checker(*a, **kw):
        raise AssertionError("Проверяющий-LLM не должен звонить на пути unify")

    request = 'Термин пишется то «Bi-encoder», то «Bi-Encoder» — приведи к одному варианту по всему документу.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=no_editor, checker=no_checker, unifier=fake_unifier,
    )
    assert result["verdict"] == "failed", result
    assert "Bi-encoder" in result["reason"], result["reason"]
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "документ обязан остаться прежним после отката"
    print("edit_demo: unify откатывается и отказывает честно, если остался вариант")


def test_unify_word_flag_protects_plural_form():
    # Тот же класс, что и test_unify_leftover..., но абзацы разные: единичное
    # "Bi-encoder" стоит ОТДЕЛЬНО от абзаца с "Bi-encoders" — validate находит
    # его по границе слова, замена проходит, а множественное число ("Bi-
    # encoders" из другого абзаца) не задето вовсе — это и есть тот самый
    # случай ColBERT 7 из BUILD_PLAN Ф18.
    doc = Document()
    doc.add_paragraph("Bi-encoder — архитектура для быстрого поиска.")
    doc.add_paragraph("Bi-encoders используются широко в индустрии.")
    doc.add_paragraph("Здесь пишут Bi-Encoder с большой буквы.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": []}

    def fake_unifier(inventory_text, request):
        return {"pairs": [["Bi-encoder", "Bi-Encoder"]], "note": ""}

    def no_editor(*a, **kw):
        raise AssertionError("Редактор не должен звонить на пути unify")

    def no_checker(*a, **kw):
        raise AssertionError("Проверяющий-LLM не должен звонить на пути unify")

    request = 'Термин пишется то «Bi-encoder», то «Bi-Encoder» — приведи к одному варианту по всему документу.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=no_editor, checker=no_checker, unifier=fake_unifier,
    )
    assert result["verdict"] == "done", result
    texts = [b["text"] for b in doc_map(doc, idx)]
    assert texts[0] == "Bi-Encoder — архитектура для быстрого поиска.", texts
    assert texts[1] == "Bi-encoders используются широко в индустрии.", "множественное число задето быть не должно: " + texts[1]
    print("edit_demo: флаг word защищает «Bi-encoders» от замены внутри неё «Bi-encoder»")


def test_unify_empty_pairs_is_honest_failed():
    # Модель честно отказалась выбрать канонический вариант — pairs пуст.
    # Это не провал системы, а легитимный ответ (см. _UNIFY_PROMPT), но
    # результат правки для документа обязан остаться "failed", а документ —
    # нетронутым.
    doc = Document()
    doc.add_paragraph("Модель COLBERT показывает высокое качество.")
    doc.add_paragraph("ColBERT — базовая архитектура.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": []}

    def fake_unifier(inventory_text, request):
        return {"pairs": [], "note": "неясно, какой вариант канонический"}

    def no_editor(*a, **kw):
        raise AssertionError("Редактор не должен звонить на пути unify")

    def no_checker(*a, **kw):
        raise AssertionError("Проверяющий-LLM не должен звонить на пути unify")

    request = 'Термин записан как «COLBERT», «ColBERT» — приведи к одному варианту по всему документу.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=no_editor, checker=no_checker, unifier=fake_unifier,
    )
    assert result["verdict"] == "failed", result
    assert result["reason"] == "неясно, какой вариант канонический", result["reason"]
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "документ не должен измениться при пустых pairs"
    print("edit_demo: unify с пустыми pairs — честный отказ, документ не тронут")


def test_unify_falls_through_to_local_when_inventory_thin():
    # Навигатор промахнулся маршрутом: в документе только ОДНО реальное
    # написание термина — unify тут не при чём (см. _variant_inventory,
    # порог < 2 вариантов). Правка обязана уйти обычным локальным путём
    # (через editor/checker), а unifier — не звониться вовсе.
    doc = Document()
    doc.add_paragraph("Bi-encoder — единственное упоминание в документе.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_unifier(inventory_text, request):
        raise AssertionError("unifier не должен звониться при инвентаре < 2 вариантов")

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "Bi-encoder", "new": "BiEncoder"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    request = 'Термин «Bi-encoder» пишется неверно, исправь.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker, unifier=fake_unifier,
    )
    assert result["verdict"] == "done", result
    assert doc_map(doc, idx)[0]["text"] == "BiEncoder — единственное упоминание в документе."
    print("edit_demo: unify с тонким инвентарём падает обратно на локальный путь")


def test_unify_verdict_checks_own_inventory_not_model_reply():
    # Defect w13 (ColBERT 7): вердикт done сверялся только со списком пар,
    # который назвала модель — вариант, который она забыла назвать, молча
    # оставался в документе, а код всё равно говорил "заменено". Три реальных
    # написания термина, модель называет пары только для ДВУХ из них.
    doc = Document()
    doc.add_paragraph("В тексте упомянут Bi-encoder несколько раз.")
    doc.add_paragraph("Здесь тоже есть Bi-Encoder с большой буквы.")
    doc.add_paragraph("А тут вариант BiEncoder слитно.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": []}

    def fake_unifier(inventory_text, request):
        assert "Bi-encoder" in inventory_text and "BiEncoder" in inventory_text, inventory_text
        # "BiEncoder" забыт — ровно дефект ColBERT 7.
        return {"pairs": [["Bi-encoder", "Bi-Encoder"]], "note": ""}

    def no_editor(*a, **kw):
        raise AssertionError("Редактор не должен звонить на пути unify")

    def no_checker(*a, **kw):
        raise AssertionError("Проверяющий-LLM не должен звонить на пути unify")

    request = 'Термин записан как «Bi-encoder», «Bi-Encoder», «BiEncoder» — приведи к одному варианту по всему документу.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=no_editor, checker=no_checker, unifier=fake_unifier,
    )
    assert result["verdict"] == "failed", result
    assert "BiEncoder" in result["reason"], result["reason"]
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "не названный моделью вариант обязан провалить вердикт и откатить документ"
    print("edit_demo: unify сверяет done со своим инвентарём, а не со списком пар модели")


def test_unify_without_single_canonical_falls_through_to_local():
    # Замер w14: пары модели разошлись по нескольким канонам (ColBERT 4 —
    # шесть разных чисел, ColBERT 12 — две формулировки) — это не унификация
    # одного написания. Отказ здесь был чистой потерей: в w12 те же правки
    # делал локальный путь. Маршрут обязан уступить, а не перехватить.
    doc = Document()
    doc.add_paragraph("Доля 4.2% и метрика 0.4448 записаны через точку.")
    doc.add_paragraph("Ещё раз 4.2% в другом абзаце.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": []}

    def fake_unifier(inventory_text, request):
        return {"pairs": [["4.2%", "4,2%"], ["0.4448", "0,4448"]], "note": ""}

    called = []

    def fake_editor(fragment_text, request):
        called.append(fragment_text)
        return {"ops": [{"op": "replace_all", "old": "4.2%", "new": "4,2%"}]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "числа приведены к запятой"}

    request = 'Числа «4.2%» и «0.4448» записаны через точку — приведи к запятой.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker, unifier=fake_unifier,
    )
    assert called, "правка обязана уйти на локальный путь, а не отказаться на unify"
    assert result["verdict"] == "done", result
    print("edit_demo: unify без единого канона уступает локальному пути, а не отказывает")


def test_unify_old_outside_inventory_never_applied():
    # Defect w13 (ColBERT 7, "Single-vector"→"Bi-encoder"): модель предложила
    # переписать прилагательное "Single-vector" в термин "Bi-Encoder" — этого
    # old код не считал (не входит в инвентарь), поэтому применять его нельзя,
    # даже если find.by_text где-то его находит. Реальные варианты термина при
    # этом всё равно унифицируются.
    doc = Document()
    doc.add_paragraph("Single-vector representations differ from Bi-encoder ones.")
    doc.add_paragraph("Bi-Encoders are also used widely.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": []}

    def fake_unifier(inventory_text, request):
        return {"pairs": [["Bi-encoder", "Bi-Encoder"], ["Bi-Encoders", "Bi-Encoder"],
                           ["Single-vector", "Bi-Encoder"]], "note": ""}

    def no_editor(*a, **kw):
        raise AssertionError("Редактор не должен звонить на пути unify")

    def no_checker(*a, **kw):
        raise AssertionError("Проверяющий-LLM не должен звонить на пути unify")

    request = 'Термин записан как «Bi-encoder», «Bi-Encoders» — приведи к одному варианту по всему документу.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=no_editor, checker=no_checker, unifier=fake_unifier,
    )
    assert result["verdict"] == "done", result
    text = " ".join(b["text"] for b in doc_map(doc, idx))
    assert "Single-vector representations" in text, "old вне инвентаря обязан быть пропущен, не применён"
    assert text.count("Bi-Encoder") == 2 and "Bi-encoder" not in text and "Bi-Encoders" not in text, text
    print("edit_demo: unify пропускает пару, чей old не входит в посчитанный инвентарь")


def test_unify_canonical_containing_old_is_done_not_leftover():
    # Математика 5: канон "«Bi-encoder»" СОДЕРЖИТ старый вариант "Bi-encoder"
    # как подстроку (модель добавила кавычки) — старый (case-sensitive)
    # вариант неизбежно "находится" внутри канона, и наивная проверка
    # остатка ложно откатывала верную правку. Подстрока канона обязана быть
    # исключена из проверки остатка.
    doc = Document()
    doc.add_paragraph("В тексте упомянут Bi-encoder несколько раз.")
    doc.add_paragraph("Здесь тоже есть bi-encoder с маленькой буквы.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": []}

    def fake_unifier(inventory_text, request):
        return {"pairs": [["Bi-encoder", "«Bi-encoder»"], ["bi-encoder", "«Bi-encoder»"]], "note": ""}

    def no_editor(*a, **kw):
        raise AssertionError("Редактор не должен звонить на пути unify")

    def no_checker(*a, **kw):
        raise AssertionError("Проверяющий-LLM не должен звонить на пути unify")

    request = 'Термин пишется то «Bi-encoder», то «bi-encoder» — возьми в кавычки одинаково по всему документу.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=no_editor, checker=no_checker, unifier=fake_unifier,
    )
    assert result["verdict"] == "done", result
    text = " ".join(b["text"] for b in doc_map(doc, idx))
    assert text.count("«Bi-encoder»") == 2, text
    print("edit_demo: канон, содержащий старый вариант как подстроку, не считается ложным остатком")


def test_rule_wins_when_unify_also_set():
    # Замер w14: обратный порядок (unify выше rule) сломал ColBERT 6 (кавычки)
    # и Математику 1 (пробелы) — обе делались нормализацией, обе перехватывал
    # unify и валил. Навигатор иногда ставит оба поля; выигрывает rule, а от
    # НЕВЕРНОГО rule защищает не приоритет, а fallthrough после отката
    # Проверяющего (см. test_rule_rolled_back_falls_through_to_local_done).
    doc = Document()
    doc.add_paragraph('Здесь стоят "прямые кавычки" в тексте.')
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": "quotes", "ids": [], "anchors": []}

    def no_unifier(*a, **kw):
        raise AssertionError("unifier не должен звониться, когда Навигатор назвал rule")

    def no_editor(*a, **kw):
        raise AssertionError("Редактор не должен звониться на пути rule")

    def ok_checker(request, diff):
        return {"ok": True, "reason": "кавычки заменены"}

    request = 'Прямые кавычки замени на ёлочки по всему документу.'
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=no_editor, checker=ok_checker, unifier=no_unifier,
    )
    assert result["verdict"] == "done", result
    text = doc_map(doc, idx)[0]["text"]
    assert '"' not in text and "«прямые кавычки»" in text, text
    print("edit_demo: rule побеждает, когда Навигатор проставил unify вместе с ним")


def test_rule_rolled_back_falls_through_to_local_done():
    # Defect w13 (Математика 2/6): rule нашёл дефект, но не тот, каким была
    # правка — normalize поправил посторонний двойной пробел, Проверяющий
    # справедливо сказал "это не та правка". Раньше rolled_back у rule был
    # терминальным, и решаемая правка (замена термина) гибла вместе с ним.
    # Теперь rolled_back продолжает ТЕМ ЖЕ запросом обычный локальный путь.
    doc = Document()
    doc.add_paragraph("Тут  двойной пробел, а ниже старый термин, который нужно заменить.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": "typography", "ids": [], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "replace_text", "id": "p0", "old": "старый термин", "new": "новый термин"}]}

    def fake_checker(request, diff):
        text = " ".join(d.get("after", "") for d in diff)
        if "новый термин" in text:
            return {"ok": True, "reason": "термин заменён"}
        return {"ok": False, "reason": "это не та правка — исправился только пробел"}

    result, doc, idx = run_edit(
        doc, idx, "Замени «старый термин» на «новый термин».",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    after_text = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p0")
    assert "Тут  двойной пробел" in after_text, "нормализация обязана быть полностью откатана перед локальным путём"
    assert "новый термин" in after_text and "старый термин" not in after_text, after_text
    print("edit_demo: rolled_back у rule продолжает локальным путём и доводит правку до done")


def test_footnote_only_edit_reaches_done_not_failed():
    # Ф19-бис: footnote применяется, но не трогает ни текст, ни style/level/
    # list абзаца — раньше _diff видел пустой набор изменений и цикл откатывал
    # уже сделанную работу с "операции применились, но текст не изменился".
    # Единственная операция кластера — footnote; verdict обязан быть done.
    doc = Document()
    doc.add_paragraph("Термин имеет спорное определение в литературе.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "footnote", "id": "p0", "old": "спорное определение",
                          "text": "См. обсуждение в разделе 2."}]}

    def fake_checker(request, diff):
        assert any("сносок" in d.get("note", "") for d in diff), diff
        return {"ok": True, "reason": "сноска добавлена"}

    result, doc, idx = run_edit(
        doc, idx, "Добавь сноску после «спорное определение».",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    refs = idx["p0"].findall(".//" + qn("w:footnoteReference"))
    assert len(refs) == 1, "сноска обязана была остаться в документе (не откачена)"
    print("edit_demo: правка из одной сноски даёт done, а не 'текст не изменился'")


def test_compose_single_editor_call_over_scattered_targets():
    # Ф20: kind="compose" — правка строит одно целое из НЕСКОЛЬКИХ мест
    # документа, поэтому кластеризация по соседству (обычный local-путь,
    # см. test_scattered_targets_three_calls_one_verdict — те же три
    # разнесённых id дают там ТРИ вызова) здесь обязана быть выключена: ОДИН
    # вызов Редактора должен увидеть ВСЕ резолвленные цели сразу.
    doc = Document()
    for i in range(12):
        doc.add_paragraph(f"Абзац {i}: обычный текст.")
    idx = index(doc)
    target_ids = ["p0", "p5", "p10"]

    def fake_navigator(outline_text, request):
        return {"kind": "compose", "rule": None, "ids": target_ids, "anchors": []}

    calls = []

    def fake_editor(fragment_text, request, feedback=None):
        calls.append(fragment_text)
        block_lines = [ln for ln in fragment_text.splitlines() if re.match(r"[pt]\d+ \[", ln)]
        present = [t for t in target_ids if any(ln.startswith(f"{t} ") for ln in block_lines)]
        assert present == target_ids, f"compose обязан показать ВСЕ цели в одном фрагменте: {fragment_text!r}"
        return {"ops": [{"op": "replace_text", "id": t, "old": "обычный", "new": "составной"} for t in present]}

    checker_calls = []

    def fake_checker(request, diff):
        checker_calls.append(diff)
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "построй таблицу из этих трёх абзацев",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert len(calls) == 1, f"compose обязан дать ОДИН вызов Редактора, получили {len(calls)}"
    assert len(checker_calls) == 1, "Проверяющий остаётся один на правку, как и для local"
    assert result["verdict"] == "done", result
    print("edit_demo: compose дал один вызов Редактора над всеми тремя разнесёнными целями")


def test_compose_oversized_fragment_refused_document_unchanged():
    # Ф20: гвард объёма. compose собирает ОДИН фрагмент из всех целей — именно
    # здесь он способен вырасти за окно модели (ContextWindowExceededError уже
    # случался, см. Ф13-бис). Порог — код, не модель: Редактор и Проверяющий
    # не должны вызываться вовсе, документ не должен быть тронут.
    doc = Document()
    doc.add_paragraph("Слово " * 3000)  # ~18 тыс. знаков — заведомо выше порога 15000
    doc.add_paragraph("Второй обычный абзац.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "compose", "rule": None, "ids": ["p0", "p1"], "anchors": []}

    def fake_editor(fragment_text, request, feedback=None):
        raise AssertionError("Редактор не должен вызываться — гвард объёма обязан отклонить фрагмент раньше")

    def fake_checker(request, diff):
        raise AssertionError("Проверяющий не должен вызываться")

    result, doc, idx = run_edit(
        doc, idx, "сравни оба абзаца и сделай вывод",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "failed", result
    assert "15000" in result["reason"] and "знаков" in result["reason"], result["reason"]
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "документ не должен измениться при отказе по объёму"
    print(f"edit_demo: compose-фрагмент сверх порога честно отклонён: {result['reason']!r}")


def test_position_end_resolves_to_last_block():
    # Ф19-бис (battery 20, тот же корень у ColBERT 20): «добавь раздел в конец
    # документа» не адресуется ни id, ни якорем, ни цитатой — Навигатор
    # обязан выставить position="end", а _resolve обязан подставить последний
    # блок документа как адрес, когда обычный резолв ничего не дал.
    doc = Document()
    for t in ("Первый абзац.", "Второй абзац.", "Третий и последний абзац."):
        doc.add_paragraph(t)
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": [], "anchors": [], "position": "end"}

    def fake_editor(fragment_text, request, feedback=None):
        assert "Третий и последний абзац" in fragment_text, fragment_text
        return {"ops": [{"op": "insert_after", "id": "p2", "text": "Новый раздел в конце.", "style": "Normal"}]}

    def fake_checker(request, diff):
        return {"ok": True, "reason": "ok"}

    result, doc, idx = run_edit(
        doc, idx, "добавь раздел в конец документа",
        navigator=fake_navigator, editor=fake_editor, checker=fake_checker,
    )

    assert result["verdict"] == "done", result
    texts = [b["text"] for b in doc_map(doc, idx)]
    assert texts[-1] == "Новый раздел в конце.", texts
    print("edit_demo: position=end резолвится в последний блок документа, вставка в конец прошла")


def test_trace_table_catches_paragraph_faking_a_row():
    # В1, ложь №1 живого прогона: «Добавь строку в таблицу» — Редактор вписал
    # абзац "r4c0:Служба безопасности | r4c1:... | r4c2:нет" рядом с таблицей,
    # которая осталась 4x3. Текстовый diff видит появившийся текст с нужными
    # словами и раньше говорил done. trace="table" обязан поймать это ДО
    # Проверяющего — checker вообще не должен вызываться.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)
    before = doc_map(doc, idx)
    table_before = next(b for b in before if b["kind"] == "t")

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["t0"], "anchors": [], "trace": "table"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "insert_after", "id": "t0",
                          "text": "r4c0:Служба безопасности | r4c1:Согласование внепланового выпуска | r4c2:нет",
                          "style": "Normal"}]}

    def no_checker(request, diff):
        raise AssertionError("trace=table обязан отклонить подделку ДО вызова Проверяющего")

    result, doc, idx = run_edit(
        doc, idx, "Добавь строку в таблицу: Служба безопасности | Согласование внепланового выпуска | нет",
        navigator=fake_navigator, editor=fake_editor, checker=no_checker,
    )

    assert result["verdict"] == "failed", result
    after = doc_map(doc, idx)
    table_after = next(b for b in after if b["kind"] == "t")
    assert table_after["rows"] == table_before["rows"], "таблица не должна была измениться"
    assert not any("r4c0" in b["text"] for b in after if b["kind"] == "p"), "поддельный абзац обязан быть откачен"
    print(f"edit_demo: trace=table поймал абзац, подделывающий строку таблицы: {result['reason']!r}")


def test_trace_heading_text_catches_heading_without_body():
    # В1, ложь №2 живого прогона (второй путь): \\n-гвард (В2) уже блокирует
    # "6. Пересмотр\\nРегламент...", но модель может просто ОПУСТИТЬ текст
    # раздела и вставить один Heading 1 без абзаца текста после него —
    # trace="heading+text" требует ДВУХ новых абзацев, не одного.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)
    before = [(b["id"], b.get("text"), b.get("rows")) for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p16"], "anchors": [], "trace": "heading+text"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "insert_after", "id": "p16", "text": "6. Пересмотр", "style": "Heading 1"}]}

    def no_checker(request, diff):
        raise AssertionError("trace=heading+text обязан отклонить одинокий заголовок ДО Проверяющего")

    result, doc, idx = run_edit(
        doc, idx, "Добавь раздел «6. Пересмотр» с текстом: Регламент пересматривается ежегодно приказом директора.",
        navigator=fake_navigator, editor=fake_editor, checker=no_checker,
    )

    assert result["verdict"] == "failed", result
    after = [(b["id"], b.get("text"), b.get("rows")) for b in doc_map(doc, idx)]
    assert after == before, "документ обязан остаться нетронутым — заголовок без текста раздела откачен"
    print(f"edit_demo: trace=heading+text поймал заголовок без текста раздела: {result['reason']!r}")


def test_trace_heading_text_passes_for_real_section():
    # Регресс-контроль: та же правка, сделанная ПРАВИЛЬНО (insert_paragraphs,
    # заголовок и текст — два новых абзаца), обязана дойти до Проверяющего и
    # получить done — trace не должен рубить верные правки этого класса.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p16"], "anchors": [], "trace": "heading+text"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "insert_paragraphs", "id": "p16", "items": [
            {"text": "6. Пересмотр", "style": "Heading 1"},
            {"text": "Регламент пересматривается ежегодно приказом директора.", "style": "Normal"},
        ]}]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "раздел добавлен"}

    result, doc, idx = run_edit(
        doc, idx, "Добавь раздел «6. Пересмотр» с текстом: Регламент пересматривается ежегодно приказом директора.",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    texts = [b["text"] for b in doc_map(doc, idx) if b["kind"] == "p"]
    assert "6. Пересмотр" in texts and "Регламент пересматривается ежегодно приказом директора." in texts
    print("edit_demo: trace=heading+text пропускает настоящий раздел (заголовок + текст) до done")


def test_trace_table_passes_for_real_row_insert():
    # Регресс-контроль: настоящий insert_row обязан пройти trace=table и
    # дойти до Проверяющего, как и раньше.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["t0"], "anchors": [], "trace": "table"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "insert_row", "id": "t0", "at": 1,
                          "cells": ["Служба безопасности", "Согласование внепланового выпуска", "нет"]}]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "строка добавлена"}

    result, doc, idx = run_edit(
        doc, idx, "Добавь строку в таблицу: Служба безопасности | Согласование внепланового выпуска | нет",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    table = next(b for b in doc_map(doc, idx) if b["kind"] == "t")
    assert len(table["rows"]) == 5 and table["rows"][1][0] == "Служба безопасности"
    print("edit_demo: trace=table пропускает настоящую вставку строки до done")


def test_trace_heading_passes_when_paragraph_already_had_outline_level():
    # В11 (находка Н66, «В1 инвертирован», замер w18 colbert#9/math#4): реальный
    # набор документов несёт w:outlineLvl ПРЯМО на абзаце независимо от именного
    # стиля («жирный абзац» может уже структурно значиться как заголовок в
    # навигации Word, даже стилем Normal) — старая проверка требовала ПЕРЕХОДА
    # "не был заголовком → стал", и такой абзац не проходил её никогда: он и
    # ДО set_style уже читался как заголовок кодом (_is_heading по level), хотя
    # для человека это была просто "жирная строка". Правка выполнена ПРАВИЛЬНОЙ
    # операцией (set_style на Heading 1) — она обязана дойти до done, ярлык
    # trace тут ни при чём: реальный applied-op доказывает сделанное.
    doc = Document()
    doc.add_paragraph("Заголовок документа")
    doc.add_paragraph("Обычный абзац ниже.")
    idx = index(doc)
    p0 = idx["p0"]
    pPr = p0.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "0")
    pPr.append(outline)
    assert doc_map(doc, idx)[0]["level"] == 0, "фикстура обязана уже нести level ДО правки"

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": [], "trace": "heading"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "set_style", "id": "p0", "style": "Heading 1"}]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "заголовок оформлен стилем"}

    result, doc, idx = run_edit(
        doc, idx, "Заголовок документа сейчас просто жирный абзац. Сделай его нормальным заголовком через стиль Word.",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    assert doc_map(doc, idx)[0]["style"] == "Heading 1"
    print("edit_demo: trace=heading пропускает set_style, даже если абзац уже нёс outlineLvl до правки")


def test_trace_table_passes_when_navigator_mislabels_unrelated_edit():
    # В11 (находка Н66, замер w18 colbert#16): Навигатор наугад проставил
    # trace=table правке про перепутанные обозначения в формулах — таблиц в
    # деле вообще не было, только replace_text по абзацам. Раньше это рубилось
    # («ни одна таблица не изменилась»), хотя правка не имела к таблицам
    # никакого отношения. Среди применённого нет НИ операции, способной
    # оставить table-след, НИ операции, способной его подделать текстом
    # (replace_text — точечная правка существующего текста, не тот класс) —
    # ярлыку Навигатора в этом случае просто нечего было подделывать.
    doc = Document()
    doc.add_paragraph("В формуле сумма идёт по m, максимум по n.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": [], "trace": "table"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "replace_text", "id": "p0",
                          "old": "сумма идёт по m, максимум по n",
                          "new": "сумма идёт по n, максимум по m"}]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "обозначения исправлены"}

    result, doc, idx = run_edit(
        doc, idx, "В формуле перепутаны обозначения — сумма должна идти по n, максимум по m.",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    assert doc_map(doc, idx)[0]["text"] == "В формуле сумма идёт по n, максимум по m."
    print("edit_demo: ошибочный trace=table на правке без единой таблицы не рубит верную правку")


def test_trace_footnote_regression_real_footnote_not_blocked():
    # Регресс-контроль из ТЗ: настоящая сноска Word не оставляет текстового
    # следа в теле документа вовсе (её текст — в word/footnotes.xml,
    # _struct_note уже считает w:footnoteReference) — trace=footnote не
    # должен рубить её так же, как раньше не рубил её _diff (Ф19-бис).
    doc = Document()
    doc.add_paragraph("Термин имеет спорное определение в литературе.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": [], "trace": "footnote"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "footnote", "id": "p0", "old": "спорное определение",
                          "text": "См. обсуждение в разделе 2."}]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "сноска добавлена"}

    result, doc, idx = run_edit(
        doc, idx, "Добавь сноску после «спорное определение».",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    refs = idx["p0"].findall(".//" + qn("w:footnoteReference"))
    assert len(refs) == 1
    print("edit_demo: trace=footnote не блокирует настоящую сноску (регресс-контроль)")


def test_trace_field_catches_paragraph_faking_a_page_reference():
    # В2-бис: «проставь номер страницы» — Редактор мог бы подделать это
    # обычным текстом («см. стр. 5») вместо настоящего поля Word. trace=field
    # обязан поймать это ДО Проверяющего — реальных w:fldSimple/w:fldChar
    # в документе не появилось.
    doc = Document()
    doc.add_paragraph("Раздел документа продолжается на следующей странице.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": [], "trace": "field"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "insert_after", "id": "p0", "text": "см. стр. 5", "style": "Normal"}]}

    def no_checker(request, diff):
        raise AssertionError("trace=field обязан отклонить подделку текстом ДО вызова Проверяющего")

    result, doc, idx = run_edit(
        doc, idx, "Добавь поле с номером страницы в конце абзаца.",
        navigator=fake_navigator, editor=fake_editor, checker=no_checker,
    )

    assert result["verdict"] == "failed", result
    after = doc_map(doc, idx)
    assert not any("см. стр" in b["text"] for b in after if b["kind"] == "p"), "поддельный абзац обязан быть откачен"
    print(f"edit_demo: trace=field поймал абзац, подделывающий номер страницы текстом: {result['reason']!r}")


def test_trace_field_passes_for_real_field_insert():
    # Регресс-контроль: настоящий field обязан пройти trace=field и дойти до
    # Проверяющего, как и раньше.
    doc = Document()
    doc.add_paragraph("Раздел документа продолжается на следующей странице.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": ["p0"], "anchors": [], "trace": "field"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "field", "id": "p0", "instr": "PAGE"}]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "поле добавлено"}

    result, doc, idx = run_edit(
        doc, idx, "Добавь поле с номером страницы в конце абзаца.",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    assert idx["p0"].find(qn("w:fldSimple")) is not None
    print("edit_demo: trace=field пропускает настоящую вставку поля до done")


def test_trace_header_footer_catches_paragraph_faking_a_footer():
    # В2-бис: «добавь номер страницы в нижний колонтитул» — Редактор мог бы
    # подделать это обычным абзацем в теле документа вместо настоящего
    # колонтитула. trace=header_footer обязан поймать это ДО Проверяющего —
    # ни один колонтитул не стал отдельной частью документа.
    doc = Document()
    doc.add_paragraph("Тело документа.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": [], "anchors": [], "position": "end", "trace": "header_footer"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "insert_after", "id": "p0", "text": "Страница 1", "style": "Normal"}]}

    def no_checker(request, diff):
        raise AssertionError("trace=header_footer обязан отклонить подделку абзацем ДО вызова Проверяющего")

    result, doc, idx = run_edit(
        doc, idx, "Добавь номер страницы в нижний колонтитул.",
        navigator=fake_navigator, editor=fake_editor, checker=no_checker,
    )

    assert result["verdict"] == "failed", result
    after = doc_map(doc, idx)
    assert not any("Страница 1" in b["text"] for b in after if b["kind"] == "p"), "поддельный абзац обязан быть откачен"
    print(f"edit_demo: trace=header_footer поймал абзац, подделывающий колонтитул: {result['reason']!r}")


def test_trace_header_footer_passes_for_real_header_footer_set():
    # Регресс-контроль: настоящий set_header_footer обязан пройти
    # trace=header_footer и дойти до Проверяющего, как и раньше.
    doc = Document()
    doc.add_paragraph("Тело документа.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": [], "anchors": [], "position": "end", "trace": "header_footer"}

    def fake_editor(fragment_text, request, feedback=None):
        return {"ops": [{"op": "set_header_footer", "which": "footer", "text": "Страница ", "field": "PAGE"}]}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "колонтитул изменён"}

    result, doc, idx = run_edit(
        doc, idx, "Добавь номер страницы в нижний колонтитул.",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    assert doc.sections[-1].footer.is_linked_to_previous is False
    print("edit_demo: trace=header_footer пропускает настоящее изменение колонтитула до done")


def test_pattern_addresses_dates_without_quotable_anchor():
    # В5: «даты приведи к формату ДД.ММ.ГГГГ» не адресуется ни id, ни цитатой
    # — формат нельзя процитировать. Навигатор возвращает pattern (regex),
    # код сам находит ВСЕ совпадения через find.by_regex.
    doc = Document()
    doc.add_paragraph("Совещание состоится 5 января 2024 года, отчёт сдать позже.")
    doc.add_paragraph("Вторая встреча запланирована на 12 марта 2025 года.")
    doc.add_paragraph("Повестка дня согласуется отдельно с каждым участником.")
    doc.add_paragraph("Протокол ведёт секретарь комиссии.")
    doc.add_paragraph("Итоги рассылаются всем участникам совещания.")
    idx = index(doc)

    date_pattern = r"\d{1,2}\s+[а-яё]+\s+\d{4}\s+года"

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": [], "anchors": [], "pattern": date_pattern}

    def fake_editor(fragment_text, request, feedback=None):
        ops = []
        if "5 января 2024 года" in fragment_text:
            ops.append({"op": "replace_text", "id": "p0", "old": "5 января 2024 года", "new": "05.01.2024"})
        if "12 марта 2025 года" in fragment_text:
            ops.append({"op": "replace_text", "id": "p1", "old": "12 марта 2025 года", "new": "12.03.2025"})
        return {"ops": ops}

    def ok_checker(request, diff):
        return {"ok": True, "reason": "формат дат унифицирован"}

    result, doc, idx = run_edit(
        doc, idx, "Даты приведи к единому формату ДД.ММ.ГГГГ.",
        navigator=fake_navigator, editor=fake_editor, checker=ok_checker,
    )

    assert result["verdict"] == "done", result
    texts = [b["text"] for b in doc_map(doc, idx)]
    assert "05.01.2024" in texts[0] and "12.03.2025" in texts[1], texts
    assert "января" not in texts[0] and "марта" not in texts[1]
    print("edit_demo: pattern без цитаты и id нашёл обе даты, обе приведены к ДД.ММ.ГГГГ")


def test_pattern_invalid_regex_is_honest_refusal_not_exception():
    # Гвард В5: модель-автор регулярки может прислать невалидный regex —
    # честный отказ ("не нашёл адреса"), а не исключение наружу.
    doc, idx = _fake_doc()

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": [], "anchors": [], "pattern": "[неверная("}

    def fake_editor(*a, **kw):
        raise AssertionError("Редактор не должен вызываться — адреса нет вовсе")

    result, doc, idx = run_edit(
        doc, idx, "Приведи формат к единому виду.",
        navigator=fake_navigator, editor=fake_editor,
    )

    assert result["verdict"] == "failed", result
    print("edit_demo: невалидный regex от Навигатора — честный отказ, не исключение")


def test_pattern_matching_implausible_share_treated_as_no_address():
    # Гвард В5: жадная/сломанная регулярка, нашедшая неправдоподобную долю
    # блоков документа, не должна адресовать правку на ВЕСЬ документ —
    # это тот же класс риска, что и document-wide replace_all без разбора.
    doc = Document()
    doc.add_paragraph("Первый абзац: значение 1.")
    doc.add_paragraph("Второй абзац: значение 5.")
    doc.add_paragraph("Третий абзац: значение 10 и 20.")
    doc.add_paragraph("Четвёртый абзац: значение 42.")
    idx = index(doc)
    before = [b["text"] for b in doc_map(doc, idx)]

    def fake_navigator(outline_text, request):
        return {"kind": "local", "rule": None, "ids": [], "anchors": [], "pattern": r"\d+"}

    def fake_editor(*a, **kw):
        raise AssertionError("Редактор не должен вызываться — доля совпадений неправдоподобна")

    result, doc, idx = run_edit(
        doc, idx, "Приведи числа к единому формату.",
        navigator=fake_navigator, editor=fake_editor,
    )

    assert result["verdict"] == "failed", result
    after = [b["text"] for b in doc_map(doc, idx)]
    assert after == before, "документ не должен был измениться"
    print("edit_demo: regex, нашедший неправдоподобную долю блоков (4 из 4), отклонён как адрес")


def test_unify_inflected_russian_word_routes_to_editor_not_replace_all():
    # В6, первый дефект (измерено дважды на battery): unify заменяет буквально
    # — канон «литий-ионный» (им.п., ед.ч.) вставал во ВСЕ позиции, включая
    # те, где текст требовал другого падежа/числа («литий-ионный
    # аккумуляторы», «литий-ионный состоит»). Русская словоформа обязана
    # уйти к Редактору (локальный путь), а не через document-wide replace_all.
    doc = Document()
    doc.add_paragraph("Новые литий-ионные аккумуляторы поставляются заводом.")
    doc.add_paragraph("Устройство состоит из литий-ионного накопителя энергии.")
    idx = index(doc)

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": ["литий-ионные", "литий-ионного"]}

    def fake_unifier(inventory_text, request):
        return {"pairs": [["литий-ионные", "литий-ионный"], ["литий-ионного", "литий-ионный"]], "note": ""}

    editor_calls = []

    def fake_editor(fragment_text, request, feedback=None):
        editor_calls.append(fragment_text)
        return {"ops": [], "already": False,
                "note": "склонение зависит от контекста — унификация написания без контроля падежа испортит согласование"}

    def no_checker(request, diff):
        raise AssertionError("до Проверяющего не должно дойти — Редактор ничего не применил")

    request = "Термин «литий-ионный» пишется в разных формах — приведи к одному написанию по всему документу."
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=fake_editor, checker=no_checker, unifier=fake_unifier,
    )

    assert editor_calls, "русская словоформа обязана была направить правку Редактору, а не в document-wide replace_all"
    assert result["verdict"] == "failed", result
    text = " ".join(b["text"] for b in doc_map(doc, idx))
    assert "литий-ионные" in text and "литий-ионного" in text, (
        "unify не должен был стереть словоформы буквальной заменой на именительный падеж: " + text
    )
    print("edit_demo: русская словоформа не унифицируется document-wide replace_all — правка ушла к Редактору")


def test_unify_word_not_term_kandidat_release_sentence():
    # В6, второй дефект (живой прогон, реальный документ Регламент.docx):
    # unify спутал СЛОВО и ТЕРМИН. p16 документа обсуждает единообразие
    # термина «релиз-кандидат» («кандидат на релиз», «релиз кандидат» —
    # варианты записи ОДНОГО термина), а p5 использует обычное слово
    # «кандидата» в другом предложении («Сборка кандидата собирается из ветки
    # release») — не вариант термина, а самостоятельное слово. Замена его на
    # «релиз-кандидата» дала ложный done в живом прогоне. Гвард (_inflects,
    # тот же, что и для склонения) обязан направить это к Редактору.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)
    p5_before = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p5")
    assert "Сборка кандидата собирается из ветки release." in p5_before

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": [], "anchors": ["релиз-кандидат", "кандидата"]}

    def fake_unifier(inventory_text, request):
        return {"pairs": [["кандидата", "релиз-кандидата"]], "note": ""}

    editor_calls = []

    def fake_editor(fragment_text, request, feedback=None):
        editor_calls.append(fragment_text)
        return {"ops": [], "already": False, "note": "'кандидата' здесь не вариант термина, а обычное слово"}

    def no_checker(request, diff):
        raise AssertionError("до Проверяющего не должно дойти")

    request = "Термин «релиз-кандидат» записан по-разному — приведи к одному написанию по всему документу."
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=fake_editor, checker=no_checker, unifier=fake_unifier,
    )

    assert editor_calls, "слово 'кандидата' — кириллица не в верхнем регистре, гвард обязан направить правку Редактору"
    p5_after = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p5")
    assert "Сборка кандидата собирается из ветки release." in p5_after, (
        "unify не должен был подменить обычное слово 'кандидата' в несвязанном предложении: " + p5_after
    )
    print("edit_demo: 'Сборка кандидата...' из Регламент.docx больше не искажается unify — слово ушло к Редактору")


def test_off_term_op_blocked_on_local_path_after_unify_fallthrough():
    # В11 (находка Н69): предыдущий тест проверял ГВАРД `_inflects` (unify
    # уступает Редактору), но Редактор там — заглушка, которая сама честно
    # отказывается. В живом прогоне РЕАЛЬНЫЙ Редактор попытку не отклонил —
    # он предложил ту же порчу (replace_text «кандидата» → «релиз-кандидата»)
    # уже НА ЛОКАЛЬНОМ пути, потому что защита В6/Н62 стояла только на
    # маршруте unify. Навигатор здесь называет p5 явным id (эту же ошибку
    # ловит В6 на маршруте unify, но здесь p5 попадает в резолв НЕ через
    # инвентарь) — а инвентарь вариантов термина (anchors — РЕАЛЬНЫЕ варианты
    # из p16: «релиз-кандидат», «кандидат на релиз», «релиз кандидат») его не
    # содержит вовсе. Код обязан отклонить операцию, чей old не входит в уже
    # посчитанный инвентарь, ДО Проверяющего.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)
    p5_before = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p5")
    assert "Сборка кандидата собирается из ветки release." in p5_before

    def fake_navigator(outline_text, request):
        return {"kind": "unify", "rule": None, "ids": ["p5"],
                "anchors": ["релиз-кандидат", "кандидат на релиз", "релиз кандидат"]}

    def fake_unifier(inventory_text, request):
        return {"pairs": [["кандидат на релиз", "релиз-кандидат"], ["релиз кандидат", "релиз-кандидат"]], "note": ""}

    def misbehaving_editor(fragment_text, request, feedback=None):
        # Настоящий Редактор в живом прогоне не отказался — он честно (но
        # ошибочно) предложил заменить обычное слово термином. Повторяет то
        # же самое и на ретрае — как модель, не понимающая, в чём проблема.
        if not any(ln.startswith("p5 ") for ln in fragment_text.splitlines()):
            return {"ops": [], "already": False, "note": "нет цели в этом кластере"}
        return {"ops": [{"op": "replace_text", "id": "p5", "old": "кандидата", "new": "релиз-кандидата"}]}

    def no_checker(request, diff):
        raise AssertionError("гвард обязан отклонить операцию ДО вызова Проверяющего")

    request = "Термин «релиз-кандидат» записан по-разному — приведи к одному написанию по всему документу."
    result, doc, idx = run_edit(
        doc, idx, request,
        navigator=fake_navigator, editor=misbehaving_editor, checker=no_checker, unifier=fake_unifier,
    )

    assert result["verdict"] == "failed", result
    assert "не один из вариантов термина" in result["reason"], result["reason"]
    p5_after = next(b["text"] for b in doc_map(doc, idx) if b["id"] == "p5")
    assert "Сборка кандидата собирается из ветки release." in p5_after, (
        "документ обязан остаться нетронутым: " + p5_after
    )
    print("edit_demo: Редактор на локальном пути (после unify-фолбэка) не смог тронуть чужое слово 'кандидата'")


if __name__ == "__main__":
    test_split_real_file()
    test_fallback_search_then_honest_refusal()
    test_rollback_on_checker_reject()
    test_resolve_normalizes_integer_id()
    test_ellipsis_truncation_retries_to_full_text()
    test_check_diff_carries_lengths()
    test_rule_no_op_yields_already()
    test_non_rule_empty_diff_stays_failed()
    test_quoted_phrase_always_merges_with_navigator_hit()
    test_checker_exception_restores_document_and_cleans_snapshot()
    test_editor_exception_mid_batch_restores_earlier_cluster()
    test_set_style_diff_reaches_checker()
    test_diff_move_and_insert()
    test_move_that_cancels_itself_is_not_done()
    test_set_list_level_diff_reaches_checker()
    test_rule_fallthrough_to_editor_done()
    test_rule_fallthrough_editor_declines_stays_already()
    test_edit_12_on_real_doc_fixture()
    test_scattered_targets_three_calls_one_verdict()
    test_adjacent_targets_one_call()
    test_empty_ops_cluster_skipped_others_apply()
    test_invalid_cluster_aborts_whole_edit()
    test_many_scattered_targets_no_cap()
    test_out_of_lane_op_blocked_document_unchanged()
    test_out_of_lane_retry_recovers_within_fragment()
    test_replace_all_not_blocked_by_lane_guard()
    test_op_on_block_created_in_same_batch_applies()
    test_two_insert_after_same_anchor_keep_order()
    test_render_full_text_outline_stays_short()
    test_defect1_mid_word_cut_rejected()
    test_defect2_batch_collision_colbert11()
    test_defect2_batch_collision_replace_all()
    test_batch_identity_op_skipped_not_retried()
    test_batch_consumed_old_skipped_not_retried()
    test_new_containing_old_still_applies()
    test_variant_inventory_appears_for_multiple_targets()
    test_variant_inventory_absent_for_single_target()
    test_collision_feedback_lets_retry_succeed()
    test_set_text_over_earlier_write_in_same_batch()
    test_partial_cluster_processing_not_done()
    test_retry_recomputes_fragment_after_batch_mutation()
    test_failed_records_ops_applied_before_batch_failure()
    test_id_fields_covers_all_ops_except_document_wide()
    test_edit_prompt_mentions_every_op()
    test_edit_prompt_frames_insert_col_header_row()
    test_replace_all_deduped_across_clusters()
    test_retry_widens_lane_to_match_rerendered_fragment()
    test_cluster_with_removed_targets_skipped_not_dead_end()
    test_targets_computed_after_validate_not_before()
    test_numeric_literals_are_anchors()
    test_editor_already_gives_already_not_failed()
    test_later_clusters_see_what_earlier_ones_did()
    test_unify_replaces_every_variant_and_reports_done()
    test_unify_leftover_variant_rolls_back_to_failed()
    test_unify_word_flag_protects_plural_form()
    test_unify_empty_pairs_is_honest_failed()
    test_unify_falls_through_to_local_when_inventory_thin()
    test_unify_verdict_checks_own_inventory_not_model_reply()
    test_unify_without_single_canonical_falls_through_to_local()
    test_unify_old_outside_inventory_never_applied()
    test_unify_canonical_containing_old_is_done_not_leftover()
    test_rule_wins_when_unify_also_set()
    test_rule_rolled_back_falls_through_to_local_done()
    test_footnote_only_edit_reaches_done_not_failed()
    test_compose_single_editor_call_over_scattered_targets()
    test_compose_oversized_fragment_refused_document_unchanged()
    test_position_end_resolves_to_last_block()
    test_trace_table_catches_paragraph_faking_a_row()
    test_trace_heading_text_catches_heading_without_body()
    test_trace_heading_text_passes_for_real_section()
    test_trace_table_passes_for_real_row_insert()
    test_trace_heading_passes_when_paragraph_already_had_outline_level()
    test_trace_table_passes_when_navigator_mislabels_unrelated_edit()
    test_trace_footnote_regression_real_footnote_not_blocked()
    test_trace_field_catches_paragraph_faking_a_page_reference()
    test_trace_field_passes_for_real_field_insert()
    test_trace_header_footer_catches_paragraph_faking_a_footer()
    test_trace_header_footer_passes_for_real_header_footer_set()
    test_pattern_addresses_dates_without_quotable_anchor()
    test_pattern_invalid_regex_is_honest_refusal_not_exception()
    test_pattern_matching_implausible_share_treated_as_no_address()
    test_unify_inflected_russian_word_routes_to_editor_not_replace_all()
    test_unify_word_not_term_kandidat_release_sentence()
    test_off_term_op_blocked_on_local_path_after_unify_fallthrough()
