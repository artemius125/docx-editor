"""Приёмка rules.py: typography/quotes на реальном документе, normalize в patch.py.

Числа — приёмочные значения из BUILD_PLAN.md (Ф3), перепроверенные на реальном
файле перед написанием кода.
"""

import re

from docx import Document

from docx_editor.parse import doc_map, index
from docx_editor.patch import apply, validate
from docx_editor import rules
from docx_editor.rules import typography

REAL_DOC = "/home/artem/Загрузки/Архитектура_ColBERT.docx"
MATH_DOC = "/home/artem/Загрузки/Математика как основа.docx"
_WS_BEFORE_PUNCT = re.compile(r"[^\S\n]+(?=[,.:;!?)\]}»])")
_GLUED = re.compile(r"([A-Za-zА-Яа-яЁё])\.([A-ZА-ЯЁ])")
_DECIMALS = ["4.2%", "0.4448", "0.4436", "8.8", "36.2%", "6.7%"]


def _text(doc, idx):
    return "\n".join(b["text"] for b in doc_map(doc, idx) if b["kind"] == "p")


def test_typography_on_real_doc():
    doc = Document(REAL_DOC)
    idx = index(doc)

    desc = apply(doc, idx, {"op": "normalize", "rule": "typography"})
    assert "7" in desc and "3" in desc, desc

    whole = _text(doc, idx)
    assert not _WS_BEFORE_PUNCT.search(whole), "остался пробел перед пунктуацией"
    assert not _GLUED.search(whole), "остались слипшиеся предложения"
    for num in _DECIMALS:
        assert num in whole, f"десятичное число {num} задето нормализацией"

    print("rules_demo: typography — пробелы и слипшиеся предложения убраны, числа целы")


def test_quotes_on_real_doc_and_bad_rule():
    doc = Document(REAL_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)

    err = validate(blocks, {"op": "normalize", "rule": "casing"}, doc)
    assert isinstance(err, str) and "typography" in err and "quotes" in err, err
    assert validate(blocks, {"op": "normalize", "rule": "quotes"}, doc) is None

    desc = apply(doc, idx, {"op": "normalize", "rule": "quotes"})
    assert "22" in desc, desc

    whole = _text(doc, idx)
    assert whole.count('"') == 0, "остались прямые кавычки"
    # 13 исходных пар + 11 новых (22 прямых / 2) = 24 каждой
    assert whole.count("«") == 24 and whole.count("»") == 24

    print("rules_demo: quotes заменены, исходные ёлочки целы, неизвестное правило отбито")


def test_invisible_chars_synthetic():
    # В реальном документе невидимых символов нет (см. BUILD_PLAN, Ф3) — правило
    # проверяется на синтетике.
    doc = Document()
    doc.add_paragraph("сло⁠во с​ невидимыми﻿ символами­ тут")

    desc = typography(doc)
    assert "4" in desc, desc

    whole = _text(doc, index(doc))
    assert whole == "слово с невидимыми символами тут", whole
    print("rules_demo: невидимые символы вырезаны на синтетике")


def test_whitespace_rules_on_real_doc():
    """ПАРТИЯ 3: двойной и висячий пробелы по ВСЕМУ документу.

    Математика 1 просит убрать два названных места «и заодно проверить весь
    текст на такое же». В замере w9 это уходило Редактору, он чинил ровно
    процитированное и не доходил до блока 50 — вердикт «сделано» при
    невыполненной правке. Смысл здесь не нужен, поэтому чинит код.
    """
    doc = Document(MATH_DOC)
    blocks = doc_map(doc, index(doc))
    bad = [b["id"] for b in blocks
           if b["kind"] == "p" and ("  " in b["text"] or b["text"].endswith(" "))]
    assert len(bad) == 3, bad  # 1 двойной + 2 висячих, замерено на файле

    note = rules.typography(doc)
    assert "схлопнуто двойных пробелов — 1" in note, note
    assert "убрано висячих пробелов в конце абзаца — 2" in note, note

    blocks = doc_map(doc, index(doc))
    left = [b["id"] for b in blocks
            if b["kind"] == "p" and ("  " in b["text"] or b["text"].endswith(" "))]
    assert not left, left
    print("rules_demo: двойные и висячие пробелы убраны по всему документу")


if __name__ == "__main__":
    test_typography_on_real_doc()
    test_quotes_on_real_doc_and_bad_rule()
    test_invisible_chars_synthetic()
    test_whitespace_rules_on_real_doc()
