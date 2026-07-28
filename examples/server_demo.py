"""Приёмка server.py: upload -> /edit по настоящему циклу (Навигатор/Редактор/
Проверяющий подменены фикстурой llm.chat — без сети и без денег) -> download.

edit.py зовёт модель через docx_editor.llm.chat как атрибут модуля, поэтому
подмена llm.chat отражается и в вызовах edit.run_edit изнутри /edit — без
передачи фейковых navigator/editor/checker через HTTP-контракт (там их нет).
"""

import asyncio
import io
import json
import tempfile
import time

import httpx
import uvicorn
from docx import Document
from fastapi.testclient import TestClient

import docx_editor.llm as llm
from docx_editor.server import app

# Настоящая chat() — остальные тесты в файле подменяют llm.chat фикстурой
# _fake_chat, поэтому ссылку на реальную функцию берём ДО этой подмены.
_real_chat = llm.chat

DOC = "/home/artem/Загрузки/Архитектура_ColBERT.docx"

# Правка 1 — существующий в документе текст (p7), Навигатор находит его сразу,
# Редактор предлагает реальную замену, Проверяющий подтверждает: verdict=done.
OLD_TEXT = "высочайшая скорость на этапе инференса"
NEW_TEXT = "высокую скорость на этапе инференса"

# Правка 2 — Навигатор промахивается (несуществующий id, ненаходимый anchor), а
# в тексте самой правки нет ни одной «ёлочной» цитаты, которая реально есть в
# документе — запасной поиск по цитатам тоже пуст: честный verdict=failed.
PROMPT = (
    f"1. Замени «{OLD_TEXT}» на «{NEW_TEXT}».\n"
    "2. Впиши требование «совершенно вымышленная и точно отсутствующая в документе фраза»."
)

# Правка 2' — намеренный сбой Навигатора (не осознанный отказ, а падение) на
# ВТОРОЙ правке второй, отдельной сессии: проверяет, что (а) исключение не
# глотается где-то в потоке /edit, долетает до вызывающего, и (б) файл на
# диске уже несёт правку 1, применённую и сохранённую ДО падения — сохранение
# внутри цикла не должно "отставать" от событий, которые уже ушли в поток.
CRASH_MARKER = "ломает конвейер"
PROMPT_CRASH = (
    f"1. Замени «{OLD_TEXT}» на «{NEW_TEXT}».\n"
    f"2. Правка, которая {CRASH_MARKER} (тест громких ошибок)."
)

# Правка 3 (отдельная сессия) — Редактор сначала предлагает "old", которого
# в блоке p7 нет (проваливает validate), затем, получив фидбэк, отвечает
# правильно — проверяет, что журнал честно пишет iter=2 (а не декоративную 1).
RETRY_MARKER = "запуск с ретраем редактора"

# Правка 4 (отдельная сессия, отдельный чистый документ) — Ф8 item A: rule на
# документе, уже нормализованном, диффа нет — обязан течь как status=done,
# verdict=already, а не как провал.
ALREADY_MARKER = "уже нормально по типографике"

# Правка 5 (отдельная сессия) — Ф10: Навигатор нарочно спит NAV_SLEEP_SECONDS,
# дольше интервала пинга (5с) в server.py, чтобы гарантированно поймать пинг
# ДО первого op-события — доказательство, что поток не молчит, пока run_edit
# считает в фоновом потоке.
SLOW_MARKER = "медленная правка с пингом"
NAV_SLEEP_SECONDS = 6

_REQUIRED_OP_FIELDS = {
    "type", "status", "text", "task", "task_text", "model", "at", "dt", "blocks", "verdict",
}
_JOURNAL_FIELDS = {"task", "task_text", "model", "iter", "verdict", "reason", "ids", "reply"}


def _fake_chat(messages):
    system, user = messages[0]["content"], messages[1]["content"]
    if "навигатор" in system:
        if SLOW_MARKER in user:
            time.sleep(NAV_SLEEP_SECONDS)
            return {"kind": "local", "rule": None, "ids": ["p7"], "anchors": []}
        if OLD_TEXT in user:
            return {"kind": "local", "rule": None, "ids": ["p7"], "anchors": []}
        if CRASH_MARKER in user:
            raise RuntimeError("сбой навигатора (тест громких ошибок)")
        if ALREADY_MARKER in user:
            return {"kind": "global", "rule": "typography", "ids": [], "anchors": []}
        return {"kind": "local", "rule": None, "ids": ["p999"], "anchors": ["текст, которого точно нет"]}
    if "редактор" in system:
        if "не прошла проверку" in user:
            # ретрай после фидбэка о невалидной операции — исправленный, валидный ответ
            return {"ops": [{"op": "replace_text", "id": "p7", "old": OLD_TEXT, "new": NEW_TEXT}]}
        if RETRY_MARKER in user:
            # первый ответ Редактора нарочно невалиден: такого "old" в p7 нет
            return {"ops": [{"op": "replace_text", "id": "p7", "old": "текст, которого нет в блоке p7", "new": NEW_TEXT}]}
        return {"ops": [{"op": "replace_text", "id": "p7", "old": OLD_TEXT, "new": NEW_TEXT}]}
    if "проверяющий" in system:
        return {"ok": True, "reason": "текст заменён корректно"}
    raise AssertionError(f"неожиданный system prompt: {system[:40]!r}")


def main():
    llm.chat = _fake_chat
    client = TestClient(app)

    assert client.get("/").status_code == 200

    with open(DOC, "rb") as f:
        data = f.read()
    files = {"file": ("doc.docx", io.BytesIO(data), "application/octet-stream")}
    up = client.post("/upload", files=files)
    assert up.status_code == 200, up.text
    payload = up.json()
    session = payload["session"]
    assert len(payload["blocks"]) == 116, payload["blocks"][:1]

    prev = client.get(f"/preview/{session}")
    assert prev.status_code == 200
    assert len(prev.json()) == 116

    edit = client.post("/edit", data={"prompt": PROMPT, "session": session})
    assert edit.status_code == 200
    lines = [line for line in edit.text.splitlines() if line.strip()]
    events = [json.loads(line) for line in lines]

    assert events[0]["type"] == "planning"
    assert events[-1]["type"] == "result"
    op_events = [e for e in events[1:-1] if e["type"] == "op"]
    assert len(op_events) == 2, op_events

    done_evt, failed_evt = op_events
    for evt in op_events:
        assert _REQUIRED_OP_FIELDS <= evt.keys(), evt.keys()

    assert done_evt["verdict"] == "done", done_evt
    assert done_evt["status"] == "done"
    assert done_evt["task"] == 1
    assert done_evt["task_text"].startswith("Замени «")
    p7 = next(b for b in done_evt["blocks"] if b["id"] == "p7")
    assert NEW_TEXT in p7["text"], p7["text"]
    assert OLD_TEXT not in p7["text"], p7["text"]

    assert failed_evt["verdict"] == "failed", failed_evt
    assert failed_evt["status"] == "failed"
    assert failed_evt["task"] == 2
    assert failed_evt["text"], "у отказа должен быть непустой текст причины"

    result = events[-1]
    assert result["done"] == [done_evt["text"]]
    assert result["failed"] == [failed_evt["text"]]

    # Журнал (docx_editor/log.py): одна запись на правку, тем же порядком и
    # с тем же вердиктом, что уже проверен по событиям потока.
    records = client.get(f"/logs?session={session}").json()
    assert len(records) == 2, records
    assert [r["task"] for r in records] == [1, 2], records
    assert records[0]["verdict"] == done_evt["verdict"], records[0]
    assert records[1]["verdict"] == failed_evt["verdict"], records[1]
    assert records[1]["reason"], "у отказа в журнале должна быть непустая причина"
    for r in records:
        assert _JOURNAL_FIELDS <= r.keys(), r.keys()

    # reply — разобранный JSON-ответ ролей, сериализованный в строку: у done
    # виден реальный ответ Редактора (что он предложил), у failed Навигатор
    # промахнулся и Редактора вообще не звали ("editor": null, а не {}).
    assert NEW_TEXT in records[0]["reply"], records[0]["reply"]
    assert '"editor": null' in records[1]["reply"], records[1]["reply"]

    dl = client.get(f"/download/{session}")
    assert dl.status_code == 200
    with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
        tmp.write(dl.content)
        tmp.flush()
        doc = Document(tmp.name)
        assert len(doc.paragraphs) == 116, len(doc.paragraphs)
        assert NEW_TEXT in doc.paragraphs[7].text, "правка не дошла до файла на диске"

    print("server_demo: ok")
    return session


def test_crash_does_not_lose_saved_progress(main_session=None):
    """Правка 2 роняет Навигатора. Инвариант 6 — падение обязано долететь
    наружу, а не превратиться в тихий result. Инвариант 5 — правка 1 к этому
    моменту уже сохранена на диск (save внутри цикла, а не после него), так
    что /download отдаёт файл С этой правкой, а не старую версию.

    main_session (если передан) — сессия из main(), уже несущая 2 записи
    журнала: используется, чтобы доказать, что session-фильтр в /logs
    реально фильтрует, а не просто игнорирует параметр."""
    llm.chat = _fake_chat
    client = TestClient(app)
    with open(DOC, "rb") as f:
        data = f.read()
    files = {"file": ("doc.docx", io.BytesIO(data), "application/octet-stream")}
    session = client.post("/upload", files=files).json()["session"]

    try:
        client.post("/edit", data={"prompt": PROMPT_CRASH, "session": session})
        raised = False
    except RuntimeError:
        raised = True
    assert raised, "падение Навигатора обязано долететь до вызывающего, а не потеряться в потоке"

    dl = client.get(f"/download/{session}")
    assert dl.status_code == 200
    with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
        tmp.write(dl.content)
        tmp.flush()
        doc = Document(tmp.name)
        assert NEW_TEXT in doc.paragraphs[7].text, "правка 1 обязана была сохраниться на диск до падения правки 2"

    # Крэш-сессия получила ровно одну запись журнала (правка 1, done) — правка
    # 2 упала внутри run_edit до того, как /edit успел её залогировать.
    crash_logs = client.get(f"/logs?session={session}").json()
    assert len(crash_logs) == 1, crash_logs
    assert crash_logs[0]["verdict"] == "done", crash_logs[0]
    assert all(r["session"] == session for r in crash_logs), crash_logs

    if main_session is not None:
        # Фильтрация по session реальна: запись крэш-сессии не утекла в хвост
        # первой сессии (main()), и наоборот — обе сессии сосуществуют в одном
        # общем out/log.jsonl к этому моменту.
        main_logs = client.get(f"/logs?session={main_session}").json()
        assert len(main_logs) == 2, main_logs
        assert all(r["session"] == main_session for r in main_logs), main_logs
        assert not any(CRASH_MARKER in r["task_text"] for r in main_logs), main_logs

        # Ф10-находка: GET /logs без session отдавал журнал ВСЕХ сессий —
        # session теперь обязателен, ручка обязана отказать (422), а не
        # молча слить обе сессии разом.
        no_session = client.get("/logs")
        assert no_session.status_code == 422, no_session.text

    print("server_demo: падение внутри /edit долетает наружу, правка 1 уже на диске, журнал фильтрует по сессии")


def test_journal_iter_retry():
    """Редактор ошибается с первого раза ("old" не найден в p7), _apply_ops
    делает один ретрай с фидбэком, второй ответ уже валиден: правка проходит
    (verdict=done), а журнал обязан честно отразить это как iter=2 — иначе
    поле было бы декоративной константой."""
    llm.chat = _fake_chat
    client = TestClient(app)
    with open(DOC, "rb") as f:
        data = f.read()
    files = {"file": ("doc.docx", io.BytesIO(data), "application/octet-stream")}
    session = client.post("/upload", files=files).json()["session"]

    prompt = f"1. Замени «{OLD_TEXT}» на «{NEW_TEXT}» ({RETRY_MARKER})."
    edit = client.post("/edit", data={"prompt": prompt, "session": session})
    assert edit.status_code == 200
    events = [json.loads(line) for line in edit.text.splitlines() if line.strip()]
    op_evt = next(e for e in events if e["type"] == "op")
    assert op_evt["verdict"] == "done", op_evt

    records = client.get(f"/logs?session={session}").json()
    assert len(records) == 1, records
    assert _JOURNAL_FIELDS <= records[0].keys(), records[0].keys()
    assert records[0]["verdict"] == "done", records[0]
    assert records[0]["iter"] == 2, records[0]

    print("server_demo: ретрай Редактора отражён в журнале честно — iter=2")


def test_already_streams_as_done():
    """Ф8 item A: rule на уже нормализованном документе (свежий чистый
    документ, отдельный от DOC, чтобы гарантированно не иметь опечаток
    типографики) даёт verdict=already с пустым diff'ом — событие потока
    обязано нести status="done" (это честный успех, не провал) и попасть
    в done итогового result, при этом verdict в событии остаётся "already",
    отличая его от обычного done для отладки/UI-чипа."""
    llm.chat = _fake_chat
    client = TestClient(app)

    clean_doc = Document()
    clean_doc.add_paragraph("Чистый абзац без опечаток и лишних пробелов.")
    buf = io.BytesIO()
    clean_doc.save(buf)
    buf.seek(0)
    files = {"file": ("clean.docx", buf, "application/octet-stream")}
    session = client.post("/upload", files=files).json()["session"]

    edit = client.post("/edit", data={"prompt": f"1. {ALREADY_MARKER}.", "session": session})
    assert edit.status_code == 200
    events = [json.loads(line) for line in edit.text.splitlines() if line.strip()]
    op_evt = next(e for e in events if e["type"] == "op")

    assert op_evt["verdict"] == "already", op_evt
    assert op_evt["status"] == "done", op_evt
    assert op_evt["text"], "у already обязан быть непустой текст (причина)"

    result_evt = events[-1]
    assert result_evt["type"] == "result"
    assert result_evt["done"] == [op_evt["text"]], result_evt
    assert result_evt["failed"] == [], result_evt

    records = client.get(f"/logs?session={session}").json()
    assert len(records) == 1 and records[0]["verdict"] == "already", records

    print("server_demo: already-правка стримится как status=done (verdict=already) и попадает в done итога")


async def _run_ping_scenario():
    """Ф10: TestClient (используемый везде выше) буферизует всё тело ответа
    целиком и не даёт увидеть, КОГДА пришла каждая строка — для этого нужен
    настоящий сокет. httpx.ASGITransport тоже не годится: он дожидается
    завершения всего ASGI-приложения ДО того, как вернуть Response, и все
    строки NDJSON приходят клиенту одним пакетом с одинаковой временной
    меткой (проверено отдельно — send() в ASGITransport ничего не отдаёт
    наружу, пока стрим не закрыт). Поэтому здесь настоящий uvicorn на
    127.0.0.1 со случайным портом и настоящий httpx.AsyncClient поверх TCP —
    только так видно реальный интервал между байтами, ту самую паузу,
    которую рвёт облачный туннель."""
    config = uvicorn.Config(app, host="127.0.0.1", port=0, log_level="warning")
    server = uvicorn.Server(config)
    server_task = asyncio.create_task(server.serve())
    try:
        while not server.started:
            await asyncio.sleep(0.01)
        port = server.servers[0].sockets[0].getsockname()[1]

        # таймаут httpx по умолчанию — 5с на чтение, ровно интервал пинга: без
        # явного увеличения клиент сам оборвётся на ожидании следующего байта,
        # той самой паузой, которую в проде рвёт облачный туннель.
        async with httpx.AsyncClient(base_url=f"http://127.0.0.1:{port}", timeout=30.0) as client:
            with open(DOC, "rb") as f:
                data = f.read()
            files = {"file": ("doc.docx", io.BytesIO(data), "application/octet-stream")}
            session = (await client.post("/upload", files=files)).json()["session"]

            timeline = []
            t0 = time.perf_counter()
            prompt = f"1. {SLOW_MARKER}."
            async with client.stream("POST", "/edit", data={"prompt": prompt, "session": session}) as resp:
                assert resp.status_code == 200
                async for raw in resp.aiter_lines():
                    if raw.strip():
                        timeline.append((time.perf_counter() - t0, json.loads(raw)))
    finally:
        server.should_exit = True
        await server_task

    types = [evt["type"] for _, evt in timeline]
    assert "op" in types, types
    first_op = types.index("op")
    assert "ping" in types[:first_op], f"пинга не было до первого op-события: {types}"

    op_evt = next(evt for _, evt in timeline if evt["type"] == "op")
    assert _REQUIRED_OP_FIELDS <= op_evt.keys(), op_evt.keys()
    assert op_evt["verdict"] == "done", op_evt
    assert op_evt["status"] == "done", op_evt

    result_evt = timeline[-1][1]
    assert result_evt["type"] == "result", result_evt
    assert result_evt["done"] == [op_evt["text"]], result_evt
    assert result_evt["failed"] == [], result_evt

    gaps = [b - a for (a, _), (b, _) in zip(timeline, timeline[1:])]
    max_gap = max(gaps)
    assert max_gap < 6.5, f"пауза между строками потока превысила лимит пинга: {gaps}"

    print(
        f"server_demo: во время {NAV_SLEEP_SECONDS}с сна Навигатора пришло "
        f"{types.count('ping')} пинг(а/ов) до op, макс. пауза между строками {max_gap:.2f}с"
    )


def test_ping_keeps_stream_alive_during_slow_edit():
    asyncio.run(_run_ping_scenario())


def _fake_response(status_code, *, json_body=None, text_body=None):
    request = httpx.Request("POST", "https://fake.test/chat/completions")
    content = json.dumps(json_body).encode() if json_body is not None else text_body.encode()
    return httpx.Response(status_code, content=content, request=request)


def test_llm_chat_http_error_surfaces_body_and_size():
    """Находка: raise_for_status() выбрасывает тело ответа, а на живом сбое
    (400 от litellm) там лежала ЕДИНСТВЕННАЯ причина. llm.chat теперь
    собирает httpx.HTTPStatusError сам, с телом и размером запроса в тексте."""
    calls = []
    server_message = "DISTINCTIVE_CONTEXT_WINDOW_EXCEEDED_MARKER"

    def fake_post(url, **kwargs):
        calls.append(kwargs)
        return _fake_response(400, text_body=json.dumps({"error": {"message": server_message}}))

    llm.httpx.post = fake_post
    try:
        messages = [{"role": "user", "content": "x" * 123}]
        try:
            _real_chat(messages)
            raised = None
        except httpx.HTTPStatusError as e:
            raised = e
        assert raised is not None, "не-2xx ответ обязан бросать httpx.HTTPStatusError"
        assert server_message in str(raised), str(raised)
        assert "123" in str(raised), f"размер запроса (123 знака) не попал в сообщение: {raised}"
        assert len(calls) == 1, "HTTP-ошибка не должна ретраиться (ретрай — только на обрыве транспорта)"
    finally:
        llm.httpx.post = httpx.post
    print("server_demo: llm.chat на 400 бросает HTTPStatusError с телом ответа и размером запроса")


def test_llm_chat_200_still_returns_parsed_json():
    def fake_post(url, **kwargs):
        body = {"choices": [{"message": {"content": json.dumps({"ok": True, "n": 1})}}]}
        return _fake_response(200, json_body=body)

    llm.httpx.post = fake_post
    try:
        result = _real_chat([{"role": "user", "content": "hi"}])
        assert result == {"ok": True, "n": 1}, result
    finally:
        llm.httpx.post = httpx.post
    print("server_demo: llm.chat на 200 по-прежнему отдаёт разобранный JSON")


def test_llm_chat_transport_retry_unchanged():
    """Ретрай остаётся ровно один и только на httpx.TransportError — HTTP-статус
    (проверено в test_llm_chat_http_error_surfaces_body_and_size) не ретраится."""
    attempts = []

    def fake_post(url, **kwargs):
        attempts.append(1)
        if len(attempts) == 1:
            raise httpx.ConnectError("boom", request=httpx.Request("POST", url))
        body = {"choices": [{"message": {"content": json.dumps({"ok": True})}}]}
        return _fake_response(200, json_body=body)

    llm.httpx.post = fake_post
    try:
        result = _real_chat([{"role": "user", "content": "hi"}])
        assert result == {"ok": True}, result
        assert len(attempts) == 2, "обрыв транспорта обязан дать ровно один ретрай"
    finally:
        llm.httpx.post = httpx.post
    print("server_demo: ретрай на httpx.TransportError не тронут (1 обрыв -> 1 ретрай -> успех)")


if __name__ == "__main__":
    session_from_main = main()
    test_crash_does_not_lose_saved_progress(session_from_main)
    test_already_streams_as_done()
    test_journal_iter_retry()
    test_ping_keeps_stream_alive_during_slow_edit()
    test_llm_chat_http_error_surfaces_body_and_size()
    test_llm_chat_200_still_returns_parsed_json()
    test_llm_chat_transport_retry_unchanged()
