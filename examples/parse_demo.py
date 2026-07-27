"""Проверка parse.py на реальном документе: 116 блоков, 0 таблиц, все Normal.

Ф9: документ размечен ПРЯМЫМ форматированием, а не именованными стилями —
`style` остаётся "Normal" везде, но `level`/`list`, читаемые из pPr напрямую,
обязаны вскрыть реальную структуру (24 заголовка, 31 элемент списка, из них
7 вложенных)."""

from collections import Counter

from docx import Document

from docx_editor.parse import doc_map, index, render

DOC = "/home/artem/Загрузки/Архитектура_ColBERT.docx"


def main():
    doc = Document(DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)

    assert len(blocks) == 116, f"ожидали 116 блоков, получили {len(blocks)}"
    tables = [b for b in blocks if b["kind"] == "t"]
    assert len(tables) == 0, f"ожидали 0 таблиц, получили {len(tables)}"
    paragraphs = [b for b in blocks if b["kind"] == "p"]
    styles = {b["style"] for b in paragraphs}
    assert styles == {"Normal"}, f"ожидали только стиль Normal, получили {styles}"

    with_level = [b for b in paragraphs if b["level"] is not None]
    assert len(with_level) == 24, f"ожидали 24 блока с level, получили {len(with_level)}"
    by_level = Counter(b["level"] for b in with_level)
    assert by_level == {0: 1, 1: 9, 2: 14}, f"неожиданное распределение по level: {by_level}"

    with_list = [b for b in paragraphs if b["list"] is not None]
    assert len(with_list) == 31, f"ожидали 31 блок с list, получили {len(with_list)}"
    nested = [b for b in with_list if b["list"]["ilvl"] == 1]
    assert len(nested) == 7, f"ожидали 7 вложенных (ilvl=1), получили {len(nested)}"

    text = render(blocks)
    # p0 — сам приёмочный случай Ф10: заголовок документа размечен как
    # обычный жирный абзац (H0 из outlineLvl, оформление — не стиль).
    assert text.startswith("p0 [Normal, H0, весь жирный]")

    print(f"parse_demo: ok, {len(blocks)} блоков, стили {styles}, "
          f"level: {dict(by_level)}, list: {len(with_list)} (вложенных {len(nested)})")


def test_render_format_tag():
    """Находка Ф10: render() был обязан выбросить оформление из тега, чтобы
    Редактор мог видеть «заголовок — просто жирный абзац». Три случая,
    которые падают, если это сломается: 1) абзац без level/list/оформления
    рендерится байт-в-байт как раньше; 2) частичный жирный — текст после ]
    остаётся дословным, span уходит в тег; 3) сплошной жирный сворачивается
    в «весь жирный», а не в список фрагментов."""
    doc = Document()
    doc.add_paragraph("Обычный абзац.")
    p2 = doc.add_paragraph()
    p2.add_run("Обычный текст, ")
    run = p2.add_run("жирный кусок")
    run.bold = True
    p3 = doc.add_paragraph()
    run3 = p3.add_run("Весь абзац жирный")
    run3.bold = True

    idx = index(doc)
    lines = render(doc_map(doc, idx)).splitlines()

    assert lines[0] == "p0 [Normal] Обычный абзац.", lines[0]
    assert lines[1] == 'p1 [Normal, жирным: «жирный кусок»] Обычный текст, жирный кусок', lines[1]
    assert lines[2] == "p2 [Normal, весь жирный] Весь абзац жирный", lines[2]

    print("parse_demo: render() — без оформления тег как раньше, частичный жирный "
          "уходит в span с дословным текстом после ], сплошной — в «весь жирный»")


if __name__ == "__main__":
    main()
    test_render_format_tag()
