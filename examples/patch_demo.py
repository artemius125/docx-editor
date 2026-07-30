"""Приёмка patch.py: девять операций на маленьком документе + валидатор невалидных патчей."""

import io
import zipfile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.table import Table

from docx_editor.parse import doc_map, index
from docx_editor.patch import _HANDLERS, _OPS, apply, validate

REAL_DOC = "/home/artem/Загрузки/Архитектура_ColBERT.docx"
REGLAMENT_DOC = "/home/artem/Документы/artemius125/docx-editor/bench/fixtures/Регламент.docx"


def _build():
    doc = Document()
    doc.add_paragraph("Первый абзац.")
    doc.add_paragraph("Второй абзац для удаления.")
    doc.add_paragraph("Третий абзац.")
    p = doc.add_paragraph()
    p.add_run("Срок ")
    p.add_run("3").bold = True
    p.add_run("0 дн").italic = True
    r = p.add_run("ей")
    r.bold, r.underline = True, True
    p.add_run(" истекает.")
    return doc


def test_ops():
    # idx создаётся один раз и переиспользуется: apply дописывает в него новые
    # элементы (insert_after/create_table), повторный index(doc) звал бы
    # позиционную перенумерацию и разрушил бы стабильность id.
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    op = {"op": "replace_text", "id": "p0", "old": "Первый", "new": "Изменённый"}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    assert blocks[0]["text"] == "Изменённый абзац."

    # фраза лежит на стыке трёх ранов разного начертания — соседи не тронуты
    apply(doc, idx, {"op": "replace_text", "id": "p3", "old": "30 дней", "new": "60 дней"})
    blocks = doc_map(doc, idx)
    assert blocks[3]["text"] == "Срок 60 дней истекает."
    runs = doc.paragraphs[3].runs
    assert runs[0].text == "Срок " and not runs[0].bold
    assert runs[-1].text == " истекает." and not runs[-1].bold

    apply(doc, idx, {"op": "set_text", "id": "p2", "text": "Совсем другой текст."})
    blocks = doc_map(doc, idx)
    assert blocks[2]["text"] == "Совсем другой текст."

    apply(doc, idx, {"op": "insert_after", "id": "p2", "text": "Новый абзац.", "style": "Normal"})
    blocks = doc_map(doc, idx)
    assert blocks[3]["text"] == "Новый абзац."

    apply(doc, idx, {"op": "set_style", "id": "p0", "style": "Heading 1"})
    blocks = doc_map(doc, idx)
    assert blocks[0]["style"] == "Heading 1"

    before = len(blocks)
    apply(doc, idx, {"op": "delete", "id": "p1"})
    blocks = doc_map(doc, idx)
    assert len(blocks) == before - 1 and all(b["id"] != "p1" for b in blocks)

    last_id = blocks[-1]["id"]
    apply(doc, idx, {"op": "move_after", "id": "p0", "after": last_id})
    blocks = doc_map(doc, idx)
    assert blocks[-1]["id"] == "p0"

    apply(doc, idx, {"op": "create_table", "after": "p0", "rows": [["a", "b"], ["c", "d"]], "header": True})
    blocks = doc_map(doc, idx)
    table = [b for b in blocks if b["kind"] == "t"][0]
    assert table["rows"] == [["a", "b"], ["c", "d"]]

    apply(doc, idx, {"op": "set_cell", "id": table["id"], "row": 1, "col": 1, "text": "абзац в ячейке"})
    blocks = doc_map(doc, idx)
    table = [b for b in blocks if b["kind"] == "t"][0]
    assert table["rows"][1][1] == "абзац в ячейке"

    apply(doc, idx, {"op": "replace_all", "old": "абзац", "new": "блок"})
    blocks = doc_map(doc, idx)
    table = [b for b in blocks if b["kind"] == "t"][0]
    assert table["rows"][1][1] == "блок в ячейке"
    assert not any("абзац" in b["text"] for b in blocks if b["kind"] == "p")

    print("patch_demo: девять операций применены и проверены")


def test_invalid():
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    assert isinstance(validate(blocks, {"op": "frobnicate"}, doc), str)
    assert isinstance(validate(blocks, {"op": "replace_text", "id": "p99", "old": "a", "new": "b"}, doc), str)
    assert isinstance(validate(blocks, {"op": "move_after", "id": "p0", "after": "p99"}, doc), str)
    assert isinstance(validate(blocks, {"op": "replace_text", "id": "p0", "old": "Первый", "new": "Первый"}, doc), str)
    assert isinstance(validate(blocks, {"op": "replace_all", "old": "абзац", "new": "абзац"}, doc), str)
    assert isinstance(validate(blocks, {"op": "replace_text", "id": "p0", "old": "нет такого текста", "new": "x"}, doc), str)

    err = validate(blocks, {"op": "set_style", "id": "p0", "style": "НесуществующийСтиль"}, doc)
    assert isinstance(err, str) and "Normal" in err

    # insert_after с несуществующим стилем — тот же класс ошибки, что и set_style
    err_ins = validate(blocks, {"op": "insert_after", "id": "p0", "text": "x", "style": "НесуществующийСтиль"}, doc)
    assert isinstance(err_ins, str) and "Normal" in err_ins

    apply(doc, idx, {"op": "create_table", "after": "p0", "rows": [["a", "b"]]})
    blocks = doc_map(doc, idx)
    table_id = [b["id"] for b in blocks if b["kind"] == "t"][0]
    assert isinstance(validate(blocks, {"op": "set_cell", "id": table_id, "row": 5, "col": 0, "text": "x"}, doc), str)

    assert validate(blocks, {"op": "replace_text", "id": "p0", "old": "Первый", "new": "Другой"}, doc) is None

    print("patch_demo: невалидные патчи отбиты валидатором")


def test_hyperlink():
    # текст гиперссылки лежит в w:hyperlink/w:r — не в r_lst абзаца верхнего
    # уровня. p.text (то, что видит validate) его включает, поэтому apply
    # обязан ходить по тем же ранам, иначе offset'ы разъезжаются и
    # _replace_span молча портит абзац (см. находку в BUILD_PLAN).
    doc = Document()
    p = doc.add_paragraph("см. ")
    hl = p._p.makeelement(qn("w:hyperlink"), {})
    r = p._p.makeelement(qn("w:r"), {})
    t = p._p.makeelement(qn("w:t"), {})
    t.text = "документацию"
    r.append(t)
    hl.append(r)
    p._p.append(hl)

    idx = index(doc)
    blocks = doc_map(doc, idx)
    assert blocks[0]["text"] == "см. документацию"

    op = {"op": "replace_text", "id": "p0", "old": "документацию", "new": "инструкцию"}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    assert blocks[0]["text"] == "см. инструкцию", blocks[0]["text"]

    print("patch_demo: replace_text через w:hyperlink не портит абзац")


def test_create_table_has_borders():
    doc = _build()
    idx = index(doc)
    apply(doc, idx, {"op": "create_table", "after": "p0", "rows": [["a", "b"], ["c", "d"]]})
    tbl = next(el for tid, el in idx.items() if tid.startswith("t"))
    assert tbl.tblPr.find(qn("w:tblBorders")) is not None, "нет w:tblBorders — таблица без рамок"
    print("patch_demo: create_table ставит рамки прямо в XML")


def test_real_doc_style_error():
    doc = Document(REAL_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)
    err = validate(blocks, {"op": "set_style", "id": "p0", "style": "List Bullet"}, doc)
    assert isinstance(err, str), "ожидали отказ: List Bullet отсутствует в документе"
    available = err.split("доступны стили абзацев:")[1]
    assert "List Bullet" not in available, err
    assert "List Paragraph" in available, err
    print("patch_demo: set_style(List Bullet) на реальном документе честно отказан")


def test_replace_text_tolerates_nbsp():
    # документ реально содержит U+00A0 посреди предложений (находка Ф4: 61 из
    # 116 абзацев, 147 вхождений), а цитата в патче от модели набрана обычным
    # пробелом. old написан обычным пробелом, в абзаце — U+00A0.
    doc = Document()
    doc.add_paragraph("слово\xa0слово синее небо.")
    idx = index(doc)
    blocks = doc_map(doc, idx)

    op = {"op": "replace_text", "id": "p0", "old": "слово слово", "new": "два слова"}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    # результат целиком, не только начало: неверный end() обрежет или откусит
    # лишнее от хвоста абзаца — падать должно громко, а не тихо на префиксе
    assert blocks[0]["text"] == "два слова синее небо.", blocks[0]["text"]

    # точное совпадение (быстрый путь) не сломано
    op2 = {"op": "replace_text", "id": "p0", "old": "синее небо", "new": "ясное небо"}
    assert validate(blocks, op2, doc) is None
    apply(doc, idx, op2)
    blocks = doc_map(doc, idx)
    assert blocks[0]["text"] == "два слова ясное небо.", blocks[0]["text"]

    print("patch_demo: replace_text находит и корректно режет спан сквозь U+00A0")


def test_set_text_rejects_ellipsis_truncation():
    # Находка Ф8: живой Редактор ответил set_text, укоротив абзац 456→332
    # знака и оборвав его буквальным «…» — смысловая правка была на месте,
    # поэтому Проверяющий сказал ok, а машинная проверка "текст изменился"
    # такое по устройству не ловит. validate обязан отбить это ДО применения.
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    bad = {"op": "set_text", "id": "p0", "text": "Первый…"}
    err = validate(blocks, bad, doc)
    assert isinstance(err, str) and err, "обрезанный set_text с многоточием обязан быть отбит"
    assert "p0" in err, err

    # многоточие тремя точками — тот же случай
    bad_dots = {"op": "set_text", "id": "p0", "text": "Первый..."}
    assert isinstance(validate(blocks, bad_dots, doc), str)

    # replace_text — та же защита, new короче old и обрывается многоточием
    bad_rt = {"op": "replace_text", "id": "p3", "old": "Срок 30 дней истекает.", "new": "Срок 30…"}
    assert isinstance(validate(blocks, bad_rt, doc), str)

    # легитимный случай: исходный текст УЖЕ оканчивается многоточием —
    # запрет не должен срабатывать на нём (в приёмочном документе таких
    # абзацев 0, но правило не должно бить по ним, если появятся)
    doc2 = Document()
    doc2.add_paragraph("Абзац с оборванной мыслью…")
    idx2 = index(doc2)
    blocks2 = doc_map(doc2, idx2)
    legit = {"op": "set_text", "id": "p0", "text": "Абзац с другой оборванной мыслью…"}
    assert validate(blocks2, legit, doc2) is None, validate(blocks2, legit, doc2)

    print("patch_demo: обрезание абзаца с многоточием в конце отбито, легитимное многоточие пропущено")


def test_replace_text_rejects_mid_word_cut():
    # Ф12: render() резал абзац ровно на границе в 300 знаков — "old" от
    # модели заканчивался посреди слова, _flex_span его честно находил (он
    # РЕАЛЬНО есть в документе), и хвост слова оставался приклеен к новому
    # тексту («...новых мощностей, херьёзных...»). validate обязан отбивать
    # такой "old" ДО применения: правая граница совпадения должна быть
    # границей слова, а не серединой.
    doc = Document()
    doc.add_paragraph("которые требуют серьёзных вычислительных ресурсов")
    idx = index(doc)
    blocks = doc_map(doc, idx)

    cut = {"op": "replace_text", "id": "p0", "old": "которые требуют с", "new": "которые требуют новых"}
    err = validate(blocks, cut, doc)
    assert isinstance(err, str) and "слова" in err, err

    # replace_all — та же защита (Ф12: тот же класс дефекта у replace_all)
    err_all = validate(blocks, {"op": "replace_all", "old": "которые требуют с", "new": "которые требуют новых"}, doc)
    assert isinstance(err_all, str) and "слова" in err_all, err_all

    # легитимный случай: правка суффикса внутри более длинного слова —
    # ЛЕВАЯ граница совпадения приходится на середину слова (как и в обрыве
    # выше), а ПРАВАЯ — на настоящий конец слова; гвард смотрит только на
    # правую и не должен сработать
    doc2 = Document()
    doc2.add_paragraph("Современные кросс-энкодеры работают быстро.")
    idx2 = index(doc2)
    blocks2 = doc_map(doc2, idx2)
    legit = {"op": "replace_text", "id": "p0", "old": "энкодеры", "new": "энкодера"}
    assert validate(blocks2, legit, doc2) is None, validate(blocks2, legit, doc2)

    # set_format — тот же класс защиты, что и replace_text/replace_all (Ф15)
    doc3 = Document()
    doc3.add_paragraph("которые требуют серьёзных вычислительных ресурсов")
    idx3 = index(doc3)
    blocks3 = doc_map(doc3, idx3)
    err_fmt = validate(blocks3, {"op": "set_format", "id": "p0", "old": "которые требуют с", "i": True}, doc3)
    assert isinstance(err_fmt, str) and "слова" in err_fmt, err_fmt

    print("patch_demo: обрыв «old» посреди слова отбит (replace_text, replace_all, set_format), легитимный суффикс пропущен")


def _add_numpr(p, ilvl, num_id):
    """Помечает абзац как элемент списка: w:pPr/w:numPr/{w:ilvl,w:numId} —
    ровно то, что parse._list читает из XML (Ф15: у _build() списков нет,
    их приходится собирать вручную, как test_hyperlink собирает гиперссылку)."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def test_set_format():
    # p3: "Срок 30 дней истекает." — "30 дней" лежит ровно на границах трёх
    # ранов разного начертания ("3" bold, "0 дн" italic, "ей" bold+underline);
    # соседи ("Срок " и " истекает.") без форматирования вовсе.
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)
    original_text = blocks[3]["text"]

    op = {"op": "set_format", "id": "p3", "old": "30 дней", "i": True}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    assert blocks[3]["text"] == original_text, "set_format не должен менять текст"

    runs = doc.paragraphs[3].runs
    assert runs[0].text == "Срок " and not runs[0].italic, "сосед слева получил курсив по ошибке"
    assert runs[-1].text == " истекает." and not runs[-1].italic, "сосед справа получил курсив по ошибке"
    assert all(r.italic for r in runs[1:-1]), [(r.text, r.italic) for r in runs]
    assert "".join(r.text for r in runs[1:-1]) == "30 дней"
    assert runs[1].bold and runs[3].bold and runs[3].underline, "старое форматирование внутри спана не должно было слететь"

    # false снимает то, что было true: "3" и "ей" были bold
    op2 = {"op": "set_format", "id": "p3", "old": "30 дней", "b": False}
    assert validate(blocks, op2, doc) is None
    apply(doc, idx, op2)
    runs = doc.paragraphs[3].runs
    assert not any(r.bold for r in runs[1:-1]), [(r.text, r.bold) for r in runs]
    assert all(r.italic for r in runs[1:-1]), "i не был указан в этой операции — не должен был тронуться"
    assert runs[3].underline, "u не был указан — не должен был слететь"

    print("patch_demo: set_format ставит и снимает начертание ровно на спане, соседи и текст не тронуты")


def test_set_format_splits_run_boundary():
    # p0: "Первый абзац." — один ран целиком. old="рвый абзац" начинается и
    # заканчивается ВНУТРИ этого рана — set_format обязан физически разрезать
    # ран (не просто переписать текст, как делает _replace_span), иначе
    # префикс/суффикс получат чужое форматирование вместе со спаном.
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    op = {"op": "set_format", "id": "p0", "old": "рвый абзац", "b": True}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    assert blocks[0]["text"] == "Первый абзац.", blocks[0]["text"]

    runs = doc.paragraphs[0].runs
    assert "".join(r.text for r in runs) == "Первый абзац."
    assert not runs[0].bold and runs[0].text == "Пе"
    assert not runs[-1].bold and runs[-1].text == "."
    middle = runs[1:-1]
    assert "".join(r.text for r in middle) == "рвый абзац"
    assert all(r.bold for r in middle)

    print("patch_demo: set_format физически режет ран по границам спана")


def test_set_format_rejects_no_flags():
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)
    err = validate(blocks, {"op": "set_format", "id": "p0", "old": "Первый"}, doc)
    assert isinstance(err, str) and err, "без b/i/u операция не меняет ничего — обязан быть отказ"
    print("patch_demo: set_format без b/i/u отбит валидатором")


def test_set_list_level():
    doc = _build()
    doc.add_paragraph("Пункт списка")
    idx = index(doc)
    blocks = doc_map(doc, idx)
    list_id = blocks[-1]["id"]
    _add_numpr(doc.paragraphs[-1], ilvl=0, num_id=5)
    blocks = doc_map(doc, idx)
    assert blocks[-1]["list"] == {"ilvl": 0, "numId": 5}, blocks[-1]["list"]

    op = {"op": "set_list_level", "id": list_id, "ilvl": 1}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    assert blocks[-1]["list"] == {"ilvl": 1, "numId": 5}, "numId обязан уцелеть, меняется только ilvl"

    # p0 в списке не состоит вовсе — сменить уровень нельзя
    err = validate(blocks, {"op": "set_list_level", "id": "p0", "ilvl": 1}, doc)
    assert isinstance(err, str) and "списк" in err, err

    print("patch_demo: set_list_level меняет ilvl и сохраняет numId, отказан на абзаце вне списка")


def test_set_list_turns_plain_paragraph_into_list_item():
    # Ф19-бис (battery 10): set_list_level умеет менять уровень уже существующего
    # списка, но не умеет превратить ОБЫЧНЫЙ абзац в элемент списка — set_list
    # закрывает именно это, ссылаясь на нумерацию, УЖЕ определённую в документе
    # (word/numbering.xml), а не выдумывая маркер-символ.
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)
    assert blocks[0]["list"] is None, "p0 не должен быть списком до операции"

    op = {"op": "set_list", "id": "p0"}  # ilvl не указан — по умолчанию 0
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    assert blocks[0]["list"]["ilvl"] == 0, blocks[0]["list"]
    real_num_ids = {int(n.get(qn("w:numId"))) for n in doc.part.part_related_by(RT.NUMBERING).element.findall(qn("w:num"))}
    assert blocks[0]["list"]["numId"] in real_num_ids, "numId обязан ссылаться на РЕАЛЬНОЕ определение нумерации документа"

    # уже список — set_list не годится, только set_list_level
    err = validate(blocks, {"op": "set_list", "id": "p0"}, doc)
    assert isinstance(err, str) and "set_list_level" in err, err
    print("patch_demo: set_list сделал обычный абзац элементом списка с реальным numId, отказан на уже-списке")


def test_set_list_rejects_without_numbering_definitions():
    # Документ без единого определения нумерации (как Математика как основа.docx,
    # где word/numbering.xml отсутствует вовсе) — сделать абзац элементом списка
    # нечем, честный отказ, а не выдуманный numId.
    doc = _build()
    part = doc.part.part_related_by(RT.NUMBERING)
    for n in part.element.findall(qn("w:num")):
        part.element.remove(n)
    idx = index(doc)
    blocks = doc_map(doc, idx)

    err = validate(blocks, {"op": "set_list", "id": "p0"}, doc)
    assert isinstance(err, str) and "нумерации" in err, err
    print(f"patch_demo: set_list честно отказан без единого определения нумерации: {err!r}")


def test_replace_all_reports_only_real_changes():
    # replace_all отчитывался по числу СОВПАДЕНИЙ, а не изменений: _flex_span
    # нормализует пробельные разрывы при поиске, и повторное применение той
    # же замены снова "находит" совпадение (обычный пробел ~ U+00A0), хотя
    # подстановка пишет то же самое, что уже стоит в абзаце. Отчёт из этого
    # уходит пользователю как есть (server.py рендерит applied дословно) —
    # значит, не должен приписывать правке изменение, которого не было
    # (инвариант 5, «отчёт не врёт»).
    doc = Document()
    doc.add_paragraph("слово слово синее небо.")
    idx = index(doc)
    blocks = doc_map(doc, idx)

    op = {"op": "replace_all", "old": "слово слово", "new": "слово\xa0слово"}
    assert validate(blocks, op, doc) is None
    report1 = apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    assert blocks[0]["text"] == "слово\xa0слово синее небо.", blocks[0]["text"]
    assert "Заменено 1 вхождений" in report1 and "в 1 блоках" in report1, report1

    # повторное применение ТОГО ЖЕ op: _flex_span снова находит совпадение
    # (пробел ~ U+00A0), но подстановка ничего не меняет — блок не изменился
    # и не должен быть засчитан
    assert validate(blocks, op, doc) is None
    report2 = apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    assert blocks[0]["text"] == "слово\xa0слово синее небо.", "текст не должен был измениться повторно"
    assert "Заменено 0 вхождений" in report2 and "в 0 блоках" in report2, report2

    print("patch_demo: replace_all считает и отчитывается только по реально изменённым блокам")


def test_footnote_new_and_existing_part():
    # Математика 19: несколько сносок подряд в документе, где footnotes.xml
    # изначально нет вовсе. Первая сноска создаёт часть, вторая — дописывает
    # в неё же (не пересоздаёт), id не коллизят с зарезервированными -1/0.
    # Проверка идёт через реальный save()/Document() — только чтение живых
    # lxml-элементов в памяти не поймало бы поломанную связь в .rels или
    # отсутствующую запись в [Content_Types].xml.
    doc = Document(REAL_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)

    paras = [b for b in blocks if b["kind"] == "p" and len(b["text"]) > 30]
    p1, p2 = paras[0], paras[1]
    anchor1 = " ".join(p1["text"].split()[:3])
    anchor2 = " ".join(p2["text"].split()[:3])

    op1 = {"op": "footnote", "id": p1["id"], "old": anchor1, "text": "Первая сноска: [уточнить]."}
    assert validate(blocks, op1, doc) is None
    apply(doc, idx, op1)

    rels = [r for r in doc.part.rels.values() if r.reltype == RT.FOOTNOTES]
    assert len(rels) == 1, "footnotes.xml должен создаться ровно один раз"

    blocks = doc_map(doc, idx)
    op2 = {"op": "footnote", "id": p2["id"], "old": anchor2, "text": "Вторая сноска, другая."}
    assert validate(blocks, op2, doc) is None
    apply(doc, idx, op2)

    rels_after = [r for r in doc.part.rels.values() if r.reltype == RT.FOOTNOTES]
    assert len(rels_after) == 1, "вторая сноска не должна была пересоздать часть footnotes.xml"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    idx2 = index(doc2)
    blocks2 = doc_map(doc2, idx2)

    def _real_ids(fpart_element):
        return sorted(int(fn.get(qn("w:id"))) for fn in fpart_element.findall(qn("w:footnote"))
                      if int(fn.get(qn("w:id"))) >= 1)

    def _footnote_text(fpart_element, fid):
        fn = next(fn for fn in fpart_element.findall(qn("w:footnote")) if fn.get(qn("w:id")) == str(fid))
        return "".join(t.text or "" for t in fn.iter(qn("w:t")))

    # doc2.part_related_by(RT.FOOTNOTES) на СВЕЖЕОТКРЫТОМ документе — это
    # обобщённый Part с сырыми байтами (см. _footnotes_part), не XmlPart;
    # разбираем через parse_xml, как это делает сам _footnotes_part.
    fpart2_raw = doc2.part.part_related_by(RT.FOOTNOTES)
    fpart2_el = parse_xml(fpart2_raw.blob)
    assert _real_ids(fpart2_el) == [1, 2], f"ожидали новые id 1 и 2 без коллизий с зарезервированными -1/0: {_real_ids(fpart2_el)}"
    assert "Первая сноска" in _footnote_text(fpart2_el, 1), _footnote_text(fpart2_el, 1)
    assert "Вторая сноска" in _footnote_text(fpart2_el, 2), _footnote_text(fpart2_el, 2)

    refs1 = idx2[p1["id"]].findall(".//" + qn("w:footnoteReference"))
    assert len(refs1) == 1 and refs1[0].get(qn("w:id")) == "1", refs1
    refs2 = idx2[p2["id"]].findall(".//" + qn("w:footnoteReference"))
    assert len(refs2) == 1 and refs2[0].get(qn("w:id")) == "2", refs2

    # третья сноска — уже на документе, у которого footnotes.xml пришёл при
    # открытии как обобщённый Part (случай "часть уже есть"): _footnotes_part
    # обязан обернуть его в XmlPart, дописать, не завести вторую связь.
    p3 = next(b for b in blocks2 if b["kind"] == "p" and b["id"] not in (p1["id"], p2["id"]) and len(b["text"]) > 30)
    anchor3 = " ".join(p3["text"].split()[:3])
    op3 = {"op": "footnote", "id": p3["id"], "old": anchor3, "text": "Третья сноска, после перезагрузки."}
    assert validate(blocks2, op3, doc2) is None
    apply(doc2, idx2, op3)

    rels3 = [r for r in doc2.part.rels.values() if r.reltype == RT.FOOTNOTES]
    assert len(rels3) == 1, "дозапись в существующую footnotes.xml не должна была завести вторую связь"

    buf3 = io.BytesIO()
    doc2.save(buf3)
    buf3.seek(0)
    doc3 = Document(buf3)
    idx3 = index(doc3)
    fpart3_el = parse_xml(doc3.part.part_related_by(RT.FOOTNOTES).blob)
    assert _real_ids(fpart3_el) == [1, 2, 3], _real_ids(fpart3_el)
    assert "Первая сноска" in _footnote_text(fpart3_el, 1)
    assert "Третья сноска" in _footnote_text(fpart3_el, 3)
    refs3 = idx3[p3["id"]].findall(".//" + qn("w:footnoteReference"))
    assert len(refs3) == 1 and refs3[0].get(qn("w:id")) == "3", refs3

    print(
        "patch_demo: footnote создаёт footnotes.xml, дозаписывает в неё же (и в свежеоткрытую тоже) "
        "без коллизий id, все тексты и ссылки видны после save/reopen"
    )


def test_set_text_preserves_run_formatting_via_diff():
    # В3: живой случай архитектора — абзац с жирным «0.5%» set_text переписал
    # правильно, но жирный ран вышел ПУСТЫМ (весь новый текст свалился в
    # первый ран абзаца целиком, каким бы ни было его начертание). Правка:
    # SequenceMatcher находит неизменившийся кусок (общий заголовок с «0.5%»
    # включительно) и не трогает его раны вовсе — начертание остаётся на
    # месте, переписывается только реально изменившийся хвост.
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Порог риска — выше ")
    p.add_run("0.5%").bold = True
    p.add_run(", релиз откатывается. Решение принимается единолично.")
    idx = index(doc)
    blocks = doc_map(doc, idx)

    op = {
        "op": "set_text", "id": "p0",
        "text": "Порог риска — выше 0.5%, релиз откатывается только после согласования с директором.",
    }
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)

    assert blocks[0]["text"] == op["text"], blocks[0]["text"]
    bold_runs = [r for r in doc.paragraphs[0].runs if r.text and r.bold]
    assert any("0.5%" in r.text for r in bold_runs), [
        (r.text, r.bold) for r in doc.paragraphs[0].runs]

    print("patch_demo: set_text (диффом) сохраняет жирное «0.5%» при переписывании хвоста абзаца")


def test_set_text_preserves_structure_loses_run_formatting():
    # Ф19/В3: set_text переписывает абзац целиком (Редактор — автор новой
    # формулировки). Фиксируем, что именно сохраняется (стиль, level из
    # w:outlineLvl, list из w:numPr) даже при полной перезаписи, а что
    # ЛЕГИТИМНО теряется: если новый текст не делит с старым НИ ОДНОГО
    # выровненного по границам ранов куска (здесь — полностью другая фраза),
    # диффу нечего сопоставлять, и начертание отдельных ранов внутри абзаца
    # теряется — это записанный предел В3, а не баг.
    doc = _build()
    idx = index(doc)
    p3 = doc.paragraphs[3]
    p3.style = doc.styles["Heading 2"]
    pPr = p3._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "1")
    pPr.append(outline)
    _add_numpr(p3, ilvl=0, num_id=7)

    blocks = doc_map(doc, idx)
    assert blocks[3]["style"] == "Heading 2"
    assert blocks[3]["level"] == 1
    assert blocks[3]["list"] == {"ilvl": 0, "numId": 7}
    assert any(r.get("b") or r.get("i") or r.get("u") for r in blocks[3]["runs"]), "фикстура обязана нести раны с начертанием"

    op = {"op": "set_text", "id": "p3", "text": "Совсем новая формулировка."}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)

    assert blocks[3]["text"] == "Совсем новая формулировка."
    assert blocks[3]["style"] == "Heading 2", "стиль абзаца обязан был сохраниться"
    assert blocks[3]["level"] == 1, "level (w:outlineLvl) обязан был сохраниться"
    assert blocks[3]["list"] == {"ilvl": 0, "numId": 7}, "list (w:numPr) обязан был сохраниться"
    assert not any(r.get("b") or r.get("i") or r.get("u") for r in blocks[3]["runs"]), (
        "set_text теряет начертание ранов — если этот assert упал, ограничение из BUILD_PLAN устарело"
    )

    print("patch_demo: set_text сохраняет стиль/level/list, теряет начертание ранов внутри абзаца")


def test_insert_delete_row_on_real_table():
    # В2, п.1: реальная таблица 4x3 из Регламент.docx (заголовок + 3 строки).
    # Раньше строку добавить было нечем — модель подделывала её абзацем со
    # своей нотацией "r4c0:… | r4c1:…". insert_row/delete_row обязаны менять
    # РЕАЛЬНУЮ структуру таблицы (число <w:tr>, gridCol не трогается), а не
    # пририсовывать текст рядом.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)
    table = next(b for b in blocks if b["kind"] == "t")
    tbl = idx[table["id"]]
    nrows_before = len(table["rows"])
    ncols = len(table["rows"][0])

    op = {"op": "insert_row", "id": table["id"], "at": 1,
          "cells": ["Служба безопасности", "Согласование внепланового выпуска", "нет"]}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    table = next(b for b in blocks if b["kind"] == "t")
    assert len(table["rows"]) == nrows_before + 1
    assert table["rows"][1] == ["Служба безопасности", "Согласование внепланового выпуска", "нет"]
    assert table["rows"][0][0] == "Роль", "заголовок не должен был сдвинуться"
    assert len(tbl.tblGrid.findall(qn("w:gridCol"))) == ncols, "insert_row не должен трогать колонки"
    assert all(len(tr.findall(qn("w:tc"))) == ncols for tr in tbl.findall(qn("w:tr")))

    op_del = {"op": "delete_row", "id": table["id"], "row": 1}
    assert validate(blocks, op_del, doc) is None
    apply(doc, idx, op_del)
    blocks = doc_map(doc, idx)
    table = next(b for b in blocks if b["kind"] == "t")
    assert len(table["rows"]) == nrows_before
    assert table["rows"][1][0] == "Релиз-менеджер", "после удаления вставленной строки таблица должна вернуться к исходному виду"

    # нельзя удалить последнюю оставшуюся строку
    tiny = Document()
    tiny.add_table(rows=1, cols=2)
    tidx = index(tiny)
    tblocks = doc_map(tiny, tidx)
    t1 = next(b for b in tblocks if b["kind"] == "t")
    err = validate(tblocks, {"op": "delete_row", "id": t1["id"], "row": 0}, tiny)
    assert isinstance(err, str) and "строку" in err, err

    # round-trip save/reopen — проверка, что XML реально валиден для Word,
    # не только для живых lxml-объектов в памяти
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    idx2 = index(doc2)
    blocks2 = doc_map(doc2, idx2)
    table2 = next(b for b in blocks2 if b["kind"] == "t")
    assert len(table2["rows"]) == nrows_before

    print("patch_demo: insert_row/delete_row меняют реальную структуру таблицы 4x3, откат на последней строке отбит")


def test_insert_delete_col_on_real_table():
    # В2, п.1 (колонка): gridCol в w:tblGrid и w:tc в КАЖДОЙ строке обязаны
    # прирастать/убывать синхронно — иначе таблица не откроется в Word.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)
    table = next(b for b in blocks if b["kind"] == "t")
    tbl = idx[table["id"]]
    nrows = len(table["rows"])
    ncols_before = len(table["rows"][0])

    op = {"op": "insert_col", "id": table["id"], "at": 1, "cells": ["x", "y", "z", "w"]}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    table = next(b for b in blocks if b["kind"] == "t")
    assert len(table["rows"][0]) == ncols_before + 1
    assert [row[1] for row in table["rows"]] == ["x", "y", "z", "w"]
    assert table["rows"][0][0] == "Роль" and table["rows"][0][2] == "Зона ответственности", (
        "соседние колонки не должны были сдвинуться содержимым")
    assert len(tbl.tblGrid.findall(qn("w:gridCol"))) == ncols_before + 1
    assert all(len(tr.findall(qn("w:tc"))) == ncols_before + 1 for tr in tbl.findall(qn("w:tr")))

    op_del = {"op": "delete_col", "id": table["id"], "col": 1}
    assert validate(blocks, op_del, doc) is None
    apply(doc, idx, op_del)
    blocks = doc_map(doc, idx)
    table = next(b for b in blocks if b["kind"] == "t")
    assert len(table["rows"][0]) == ncols_before
    assert table["rows"][0] == ["Роль", "Зона ответственности", "Замена"]
    assert len(tbl.tblGrid.findall(qn("w:gridCol"))) == ncols_before

    # нельзя удалить последнюю оставшуюся колонку
    tiny = Document()
    tiny.add_table(rows=nrows, cols=1)
    tidx = index(tiny)
    tblocks = doc_map(tiny, tidx)
    t1 = next(b for b in tblocks if b["kind"] == "t")
    err = validate(tblocks, {"op": "delete_col", "id": t1["id"], "col": 0}, tiny)
    assert isinstance(err, str) and "колонку" in err, err

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    idx2 = index(doc2)
    blocks2 = doc_map(doc2, idx2)
    table2 = next(b for b in blocks2 if b["kind"] == "t")
    assert len(table2["rows"][0]) == ncols_before

    print("patch_demo: insert_col/delete_col держат gridCol и tc каждой строки в синхроне, откат на последней колонке отбит")


def test_set_format_on_table_cell():
    # В2, п.3: «выдели первый столбец жирным» раньше было невыразимо — у
    # set_format был только id абзаца. Теперь id таблицы + row + col находят
    # ячейку тем же способом, что и set_cell.
    doc = Document(REGLAMENT_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)
    table = next(b for b in blocks if b["kind"] == "t")

    op = {"op": "set_format", "id": table["id"], "row": 1, "col": 0, "old": "Релиз-менеджер", "b": True}
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    table = next(b for b in blocks if b["kind"] == "t")
    assert table["rows"][1][0] == "Релиз-менеджер", "текст ячейки не должен был измениться"
    cell_p = Table(idx[table["id"]], doc).cell(1, 0).paragraphs[0]
    assert all(r.bold for r in cell_p.runs if r.text)
    # соседняя ячейка той же строки не тронута
    neighbor_p = Table(idx[table["id"]], doc).cell(1, 1).paragraphs[0]
    assert not any(r.bold for r in neighbor_p.runs if r.text)

    # плохие row/col — честный отказ, тем же классом ошибки, что у set_cell
    err = validate(blocks, {"op": "set_format", "id": table["id"], "row": 9, "col": 0, "old": "x", "b": True}, doc)
    assert isinstance(err, str) and "вне таблицы" in err, err

    print("patch_demo: set_format адресует ячейку таблицы (row+col), соседние ячейки не тронуты")


def test_insert_paragraphs_adds_section():
    # В2, п.2: «раздел» — заголовок плюс текст, минимум два абзаца двух
    # стилей одной операцией. Раньше модель подделывала это одним абзацем
    # Heading 1 с "6. Пересмотр\nРегламент пересматривается…" внутри.
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)
    last_id = blocks[-1]["id"]

    op = {
        "op": "insert_paragraphs", "id": last_id,
        "items": [
            {"text": "6. Пересмотр", "style": "Heading 1"},
            {"text": "Регламент пересматривается ежегодно приказом директора.", "style": "Normal"},
        ],
    }
    assert validate(blocks, op, doc) is None
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)

    new = blocks[len(blocks) - 2:]
    assert new[0]["style"] == "Heading 1" and new[0]["text"] == "6. Пересмотр"
    assert new[1]["style"] == "Normal" and new[1]["text"] == "Регламент пересматривается ежегодно приказом директора."

    print("patch_demo: insert_paragraphs вставляет заголовок и текст раздела как два абзаца двух стилей")


def test_field_simple_and_compound_round_trip():
    # В2-бис: PAGE — простое поле (w:fldSimple), REF — составное (w:fldChar
    # begin/separate/end + w:instrText), одна операция на оба механизма —
    # код выбирает форму по первому слову instr, не вызывающий. Кэш-значение
    # внутри поля сознательно не кладём (см. решение честности в
    # BUILD_PLAN_V2.md, В2-бис) — раны fldChar/instrText дают r.text=="",
    # поэтому вставка поля не меняет видимый текст абзаца ни на символ.
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    op_simple = {"op": "field", "id": "p0", "instr": "PAGE"}
    assert validate(blocks, op_simple, doc) is None
    apply(doc, idx, op_simple)
    blocks = doc_map(doc, idx)
    assert blocks[0]["text"] == "Первый абзац.", "поле не должно менять текст абзаца"
    assert blocks[0]["fields"] == 1
    assert idx["p0"].find(qn("w:fldSimple")) is not None
    assert idx["p0"].find(qn("w:fldSimple")).get(qn("w:instr")) == "PAGE"

    op_compound = {"op": "field", "id": "p2", "old": "Третий", "instr": 'REF bookmark1 \\h'}
    assert validate(blocks, op_compound, doc) is None
    apply(doc, idx, op_compound)
    blocks = doc_map(doc, idx)
    assert blocks[2]["text"] == "Третий абзац.", "поле не должно менять текст абзаца"
    assert blocks[2]["fields"] == 1
    fldchars = idx["p2"].findall(".//" + qn("w:fldChar"))
    types = [f.get(qn("w:fldCharType")) for f in fldchars]
    assert types == ["begin", "separate", "end"], types
    instr_els = idx["p2"].findall(".//" + qn("w:instrText"))
    assert len(instr_els) == 1 and "REF bookmark1" in instr_els[0].text, instr_els

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    idx2 = index(doc2)
    blocks2 = doc_map(doc2, idx2)
    assert blocks2[0]["fields"] == 1 and blocks2[2]["fields"] == 1, "поля обязаны пережить save/reopen"
    assert idx2["p0"].find(qn("w:fldSimple")).get(qn("w:instr")) == "PAGE"

    print("patch_demo: field вставляет простое (PAGE) и составное (REF) поле, оба переживают save/reopen")


def test_field_validate_rejects_unknown_verb_and_table_id():
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    err = validate(blocks, {"op": "field", "id": "p0", "instr": "AUTHOR"}, doc)
    assert err and "не начинается" in err, err

    err = validate(blocks, {"op": "field", "id": "p0", "old": "Такого текста нет", "instr": "PAGE"}, doc)
    assert err and "нет текста" in err, err

    op = {"op": "create_table", "after": "p0", "rows": [["a"]]}
    apply(doc, idx, op)
    blocks = doc_map(doc, idx)
    tbl_id = next(b["id"] for b in blocks if b["kind"] == "t")
    err = validate(blocks, {"op": "field", "id": tbl_id, "instr": "PAGE"}, doc)
    assert err and "field работает только с абзацами" in err, err

    print("patch_demo: field отбивает неизвестный код поля, ненайденный якорь и id таблицы")


def test_set_header_footer_creates_parts_and_field():
    # В2-бис: колонтитул — отдельная часть пакета (word/header1.xml,
    # footer1.xml), связанная через w:sectPr; python-docx заводит её сам
    # через is_linked_to_previous=False (см. проверку в _op_set_header_footer).
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    op_footer = {"op": "set_header_footer", "which": "footer", "text": "Страница ", "field": "PAGE"}
    assert validate(blocks, op_footer, doc) is None
    report = apply(doc, idx, op_footer)
    assert "не отображается" in report, report

    op_header = {"op": "set_header_footer", "which": "header", "text": "Регламент компании"}
    assert validate(blocks, op_header, doc) is None
    apply(doc, idx, op_header)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    names = zipfile.ZipFile(buf).namelist()
    assert "word/header1.xml" in names and "word/footer1.xml" in names, names

    buf.seek(0)
    doc2 = Document(buf)
    sec = doc2.sections[-1]
    assert sec.header.is_linked_to_previous is False
    assert sec.footer.is_linked_to_previous is False
    assert sec.header.paragraphs[0].text == "Регламент компании"
    assert sec.footer.paragraphs[0].text == "Страница "
    assert sec.footer.paragraphs[0]._p.find(qn("w:fldSimple")) is not None

    print("patch_demo: set_header_footer заводит header1.xml/footer1.xml, сектор ссылается на них, поле переживает save/reopen")


def test_set_header_footer_validate_rejects_bad_input():
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    err = validate(blocks, {"op": "set_header_footer", "which": "left", "text": "x"}, doc)
    assert err and "which должен быть" in err, err

    err = validate(blocks, {"op": "set_header_footer", "which": "footer", "text": ""}, doc)
    assert err and "ничего не меняет" in err, err

    err = validate(blocks, {"op": "set_header_footer", "which": "footer", "text": "", "field": "AUTHOR"}, doc)
    assert err and "не начинается" in err, err

    print("patch_demo: set_header_footer отбивает неизвестный which, пустую пару text/field и неизвестный код поля")


def test_newline_in_op_text_rejected():
    # В2, п.4: буквальный перевод строки в тексте операции — не новый абзац,
    # а подделка того, чего нет в контракте (архитектор поймал модель на
    # "6. Пересмотр\nРегламент пересматривается…" внутри одного Heading 1).
    doc = _build()
    idx = index(doc)
    blocks = doc_map(doc, idx)

    bad_ops = [
        {"op": "set_text", "id": "p0", "text": "Первая строка\nВторая строка"},
        {"op": "insert_after", "id": "p0", "text": "раз\nдва"},
        {"op": "replace_text", "id": "p0", "old": "Первый", "new": "Пер\nвый"},
        {"op": "insert_row", "id": "p0", "at": 0, "cells": ["a", "b\nc"]},
    ]
    for op in bad_ops:
        err = validate(blocks, op, doc)
        assert isinstance(err, str) and "перевод строки" in err, (op, err)

    # легитимная операция без \n не отбивается этой защитой
    assert validate(blocks, {"op": "set_text", "id": "p0", "text": "Обычный текст без переводов строк"}, doc) is None

    print("patch_demo: перевод строки в тексте любой операции отбит валидатором с понятной причиной")


def test_all_ops_have_handlers():
    # Находка BUILD_PLAN: коммит однажды завёл set_format в _OPS без записи
    # в _HANDLERS — apply() падал бы KeyError на первом же применении.
    missing = _OPS - set(_HANDLERS)
    assert not missing, f"операции без обработчика в _HANDLERS: {missing}"
    print("patch_demo: у каждой операции из _OPS есть обработчик в _HANDLERS")


def test_missing_required_field_is_rejected_not_crash():
    # Одна операция без old роняла ВЕСЬ прогон: _flex_span(text, None) —
    # TypeError, а не отказ. validate обязан отвечать текстом, а не падать.
    doc = Document(REAL_DOC)
    idx = index(doc)
    blocks = doc_map(doc, idx)
    for op in (
        {"op": "replace_text", "id": "p10", "new": "x"},
        {"op": "replace_all", "new": "x"},
        {"op": "set_text", "id": "p10"},
    ):
        err = validate(blocks, op, doc)
        assert err and "нет поля" in err, (op, err)
    print("patch_demo: операция без обязательного поля отбита текстом, а не падением")


def test_bookmark_gives_ref_field_a_target():
    # Прогон за рулём 2026-07-30: field умел REF, но создать закладку было
    # нечем — «пронумеруй ссылки полями» (ColBERT 20) невыполнимо по
    # контракту. Пара bookmark + field REF обязана дать в пакете и цель, и
    # ссылку на неё, а второе имя той же закладки — честный отказ.
    doc = Document()
    doc.add_paragraph("Khattab и Zaharia, ColBERT: эффективный поиск.")
    doc.add_paragraph("Как показано в работе, поздняя интеракция выигрывает.")
    idx = index(doc)

    assert validate(doc_map(doc, idx), {"op": "bookmark", "id": "p0", "name": "istochnik1"}, doc) is None
    apply(doc, idx, {"op": "bookmark", "id": "p0", "name": "istochnik1"})
    apply(doc, idx, {"op": "field", "id": "p1", "instr": "REF istochnik1 \\h", "old": "в работе"})

    xml = doc.element.body.xml
    assert 'w:name="istochnik1"' in xml, "закладка не попала в документ"
    assert "bookmarkEnd" in xml, "закладка не закрыта"
    assert "REF istochnik1" in xml, "поле REF не попало в документ"
    assert doc_map(doc, idx)[0]["bookmarks"] == 1, doc_map(doc, idx)[0]

    dup = validate(doc_map(doc, idx), {"op": "bookmark", "id": "p1", "name": "istochnik1"}, doc)
    assert dup and "уже есть" in dup, dup
    bad = validate(doc_map(doc, idx), {"op": "bookmark", "id": "p1", "name": "источник 1"}, doc)
    assert bad and "латинские" in bad, bad
    print("patch_demo: bookmark даёт REF-ссылке цель, повтор имени и кириллица отклонены")


if __name__ == "__main__":
    test_all_ops_have_handlers()
    test_missing_required_field_is_rejected_not_crash()
    test_ops()
    test_invalid()
    test_hyperlink()
    test_create_table_has_borders()
    test_real_doc_style_error()
    test_replace_text_tolerates_nbsp()
    test_set_text_rejects_ellipsis_truncation()
    test_replace_text_rejects_mid_word_cut()
    test_set_format()
    test_set_format_splits_run_boundary()
    test_set_format_rejects_no_flags()
    test_set_list_level()
    test_set_list_turns_plain_paragraph_into_list_item()
    test_set_list_rejects_without_numbering_definitions()
    test_replace_all_reports_only_real_changes()
    test_footnote_new_and_existing_part()
    test_set_text_preserves_run_formatting_via_diff()
    test_bookmark_gives_ref_field_a_target()
    test_set_text_preserves_structure_loses_run_formatting()
    test_insert_delete_row_on_real_table()
    test_insert_delete_col_on_real_table()
    test_set_format_on_table_cell()
    test_insert_paragraphs_adds_section()
    test_field_simple_and_compound_round_trip()
    test_field_validate_rejects_unknown_verb_and_table_id()
    test_set_header_footer_creates_parts_and_field()
    test_set_header_footer_validate_rejects_bad_input()
    test_newline_in_op_text_rejected()
