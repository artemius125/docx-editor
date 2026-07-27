"""Приёмка find.py: 20 якорей (по одному на правку из Правки_ColBERT_20.md)
на настоящем документе Архитектура_ColBERT.docx, плюс outline и fragment.
"""

import re
from collections import Counter

from docx import Document

from docx_editor.parse import doc_map, index
from docx_editor import find

REAL_DOC = "/home/artem/Загрузки/Архитектура_ColBERT.docx"

# Якорь — самая специфичная ДОСЛОВНАЯ цитата из текста самой правки (не термин
# своими словами). Для правок 6, 9, 15 такой цитаты в тексте правки нет вообще —
# это не промах find.py, а находка (см. отчёт строителя): правка 6 — глобальное
# типографическое правило (идёт мимо Редактора), правка 9 — правка по позиции
# («заголовок» = первый блок, а не текст), правка 15 — уже известный невыполнимый
# случай (BUILD_PLAN «Известные риски»: в документе нет списков).
ANCHORS = {
    1: "структурные маркеры как",
    2: "исчисляемым секундами",
    3: "меняет этот процесс.Исследователи",
    4: "0.4436",
    5: "MS-MARCO",
    6: None,
    7: "Single-Vector Bi-encoders",
    8: "2025/2026 годов",
    9: None,
    10: "Модернизация агрегации",
    11: "Кросс-энкодер",
    12: "на порядок (в 6–10 раз)",
    13: "коллапс производительности стандартного ColBERT (падение до 86–97%)",
    14: "MRR@10",
    15: None,
    16: "матричных умножений MaxSim",
    17: "представляет собой скалярное произведение",
    18: "недифференцируемой природе операции",
    19: "FLOPs",
    20: "MetaEmbed",
}


def _load():
    doc = Document(REAL_DOC)
    idx = index(doc)
    return doc_map(doc, idx)


def test_anchors_hit_every_edit():
    blocks = _load()
    no_anchor = []
    for n, needle in ANCHORS.items():
        if needle is None:
            no_anchor.append(n)
            continue
        hits = find.by_text(blocks, needle)
        assert hits, f"правка {n}: якорь {needle!r} не нашёл ни одного блока"
        assert len(hits) <= 20, f"правка {n}: якорь {needle!r} слишком общий: {len(hits)} блоков"

    # список правок без цитаты зафиксирован явно: если он изменится — это
    # сигнал, что мой выбор якоря был плохим, а не новая находка
    assert no_anchor == [6, 9, 15], f"неожиданный набор правок без якоря: {no_anchor}"
    print(f"find_demo: {20 - len(no_anchor)}/20 якорей нашли блоки (≤20 каждый), "
          f"{len(no_anchor)} без цитаты в тексте правки")


def test_first_mention_picks_earliest():
    blocks = _load()
    # MRR@10 упоминается дважды (p48 и p77) — first_mention обязан вернуть первый
    hits = find.by_text(blocks, "MRR@10")
    assert len(hits) >= 2, "ожидали повторный термин для проверки first_mention"
    assert find.first_mention(blocks, "MRR@10") == "p48"
    assert find.first_mention(blocks, "текста, которого точно нет") is None
    print("find_demo: first_mention возвращает самое раннее упоминание")


def test_by_regex():
    blocks = _load()
    hits = find.by_regex(blocks, r"\d\.\d")
    assert "p67" in hits, "в p67 три десятичных числа (4.2%, 0.4448, 0.4436)"
    print("find_demo: by_regex находит десятичные числа")


def test_by_text_tolerates_nbsp():
    # p31 в реальном документе: "Оператор\xa0$\top$\xa0представляет..." — цитата
    # обычными пробелами не матчится как подстрока (61 из 116 абзацев содержат
    # U+00A0 посреди предложения, 147 вхождений всего), но by_text обязан её найти
    blocks = _load()
    assert "p31" in find.by_text(blocks, "Оператор $\\top$ представляет")
    print("find_demo: by_text находит цитату сквозь U+00A0 в середине предложения")


def test_outline_is_compact_and_valid():
    # Ф9: outline теперь строится по level (w:outlineLvl), а не по имени стиля —
    # в документе 24 реальных заголовка (1 H0 + 9 H1 + 14 H2), печатаются целиком,
    # это подняло размер с 2972 знаков (плоский список без единого заголовка,
    # старая проверка по Heading*/Title находила ноль) до ~4000 — бюджет всё ещё
    # мелкий на фоне документа в 35 тыс. знаков и укладывается в промпт Навигатора.
    blocks = _load()
    o = find.outline(blocks)
    assert len(o) < 4200, f"оглавление слишком длинное для промпта Навигатора: {len(o)} знаков"
    headings = [l for l in o.splitlines() if re.search(r"^\S+ \[H\d\]", l)]
    assert len(headings) == 24, f"ожидали 24 заголовка (level != None), получили {len(headings)}"
    by_level = Counter(int(re.search(r"\[H(\d)\]", l).group(1)) for l in headings)
    assert by_level == {0: 1, 1: 9, 2: 14}, f"неожиданное распределение по уровням: {by_level}"
    ids = {line.split(" ", 1)[0] for line in o.splitlines()}
    real_ids = {b["id"] for b in blocks}
    assert ids <= real_ids, f"в оглавлении лишние id: {ids - real_ids}"
    assert len(ids) == len(blocks), "оглавление пропустило блоки"
    print(f"find_demo: outline — {len(o)} знаков на {len(blocks)} блоков "
          f"(24 реальных заголовка по level, было 2972 знака и 0 заголовков по стилю)")


def test_outline_marks_fully_bold_paragraph():
    # Находка Ф10: заголовок без стиля/level — просто жирный абзац. outline()
    # обязан пометить его [B], иначе Навигатор не отличит его от тела текста.
    doc = Document()
    doc.add_paragraph("Обычный абзац с текстом подлиннее двадцати символов.")
    p2 = doc.add_paragraph()
    run = p2.add_run("Заголовок жирным без стиля")
    run.bold = True
    idx = index(doc)
    lines = find.outline(doc_map(doc, idx)).splitlines()
    assert lines[0].startswith("p0 ") and "[B]" not in lines[0], lines[0]
    assert lines[1].startswith("p1 [B] "), lines[1]
    print("find_demo: outline помечает [B] сплошь жирный абзац без стиля заголовка")


def test_fragment_dedupes_overlap():
    blocks = _load()
    # p10 и p12 с around=1 дают пересекающиеся диапазоны [p9,p11] и [p11,p13]
    frag = find.fragment(blocks, ["p10", "p12"], around=1)
    ids = [b["id"] for b in frag]
    assert ids == ["p9", "p10", "p11", "p12", "p13"], ids
    print("find_demo: fragment схлопывает пересекающиеся диапазоны без дублей")


MATH_DOC = "/home/artem/Загрузки/Математика как основа.docx"


def _load_math():
    doc = Document(MATH_DOC)
    idx = index(doc)
    return doc_map(doc, idx)


def test_locate_recovers_drifted_quotes():
    # Реальный дрейф цитирования из приёмочного корпуса: в правке — запятая
    # и «нет» без тире, в документе — тире (см. отчёт строителя). Ни точный,
    # ни пробельно-гибкий поиск такое не находит — только locate.
    blocks = _load_math()
    q1 = "Братья имели крайне низкий IQ, уровень развития, сравнимый с маленьким ребёнком"
    q2 = "Там, где число было простым, они улыбались. Там, где нет, хмурились"
    assert find.by_text(blocks, q1) == [], "ожидали, что by_text промахнётся мимо дрейфа"
    assert find.by_text(blocks, q2) == [], "ожидали, что by_text промахнётся мимо дрейфа"
    assert find.locate(blocks, q1)[0] == "p21", find.locate(blocks, q1)
    assert find.locate(blocks, q2)[0] == "p25", find.locate(blocks, q2)
    print("find_demo: locate восстанавливает 2 реальных дрейфа цитирования (p21, p25)")


def test_locate_rejects_short_and_unrelated():
    blocks = _load_math()
    assert find.locate(blocks, "слон бегемот") == [], "2 токена — короче порога, должно быть пусто"
    assert find.locate(blocks, "совершенно случайная фраза не из документа никак") == []
    print("find_demo: locate возвращает [] на короткой и на нерелевантной фразе")


def test_locate_never_returns_more_than_one():
    # Находка (см. docstring locate): жадный вариант измерен вживую на 40 правках
    # и стоил больше правок, чем принёс (3 выигрыша против 6 регрессий) — якоря
    # Навигатора короткие и общие, и старый locate возвращал их все, раздувая
    # фрагмент. "это не просто" — фраза из трёх общих слов документа: жадное
    # сканирование (совпадение по токенам-подпоследовательности, без учёта
    # порядка блоков) находит её в 9 блоках документа ['p14', 'p15', 'p17',
    # 'p18', 'p19', 'p23', 'p26', 'p30', 'p43'], так что проверка неслучайна.
    # Если locate регрессирует к жадному варианту, это упадёт.
    blocks = _load_math()
    hits = find.locate(blocks, "это не просто")
    assert len(hits) <= 1, f"locate обязан возвращать не больше одного блока: {hits}"
    print(f"find_demo: locate на заведомо многозначной фразе «это не просто» вернул {hits}, не список из 9")


def test_by_text_unchanged_for_exact_substring():
    blocks = _load_math()
    assert "p21" in find.by_text(blocks, "Братья имели крайне низкий IQ")
    print("find_demo: by_text по-прежнему находит точную подстроку без изменений")


if __name__ == "__main__":
    test_anchors_hit_every_edit()
    test_first_mention_picks_earliest()
    test_by_regex()
    test_by_text_tolerates_nbsp()
    test_outline_is_compact_and_valid()
    test_outline_marks_fully_bold_paragraph()
    test_fragment_dedupes_overlap()
    test_locate_recovers_drifted_quotes()
    test_locate_rejects_short_and_unrelated()
    test_locate_never_returns_more_than_one()
    test_by_text_unchanged_for_exact_substring()
